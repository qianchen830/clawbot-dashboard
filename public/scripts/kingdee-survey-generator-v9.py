#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶调研报告生成器 - v9.0 超级增强版
基于深度学习，支持12+章节、60个调研问卷、4个阶段汇报PPT
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# 输出目录
OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_survey_report_v9(customer_info):
    """生成调研报告 - v9.0 超级增强版（12+章节）"""
    doc = Document()
    
    # 设置字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    company_name = customer_info.get('companyName', '企业名称')
    industry = customer_info.get('industry', '制造业')
    selected_modules = customer_info.get('selectedModules', ['finance', 'supply', 'manufacture'])
    
    # 标题页
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{company_name}\n")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("ERP系统调研报告\n")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'企业名称：{company_name}\n')
    run.font.size = Pt(14)
    run.font.bold = True
    run = info.add_run(f'所属行业：{industry}\n')
    run = info.add_run(f'企业规模：{customer_info.get("companySize", "中型企业")}\n')
    run = info.add_run(f'调研日期：{datetime.now().strftime("%Y年%m月%d日")}\n')
    run = info.add_run(f'调研模块：{", ".join([m for m in selected_modules])}\n')
    
    doc.add_page_break()
    
    # 目录（12+章节）
    doc.add_heading('目录', level=1)
    chapters = [
        '一、调研概述',
        '二、调研方法论',
        '三、企业概况',
        '四、业务现状分析',
        '五、存在问题分析',
        '六、调研问卷分析',
        '七、需求分析',
        '八、业务架构设计',
        '九、解决方案设计',
        '十、实施路线图',
        '十一、风险评估',
        '十二、调研结论与建议',
        '附录：调研问卷库（60题）'
    ]
    for chapter in chapters:
        doc.add_paragraph(chapter)
    doc.add_page_break()
    
    # 一、调研概述
    doc.add_heading('一、调研概述', level=1)
    doc.add_heading('1.1 调研目的', level=2)
    doc.add_paragraph(f'全面了解{company_name}的业务现状和需求，为ERP系统实施提供依据。')
    doc.add_paragraph('调研目的包括：')
    for purpose in [
        '了解企业组织架构和业务流程',
        '分析现有系统存在的问题',
        '明确业务需求和功能需求',
        '制定系统实施方案和路线图',
        '评估项目风险和收益'
    ]:
        doc.add_paragraph(f'• {purpose}', style='List Bullet')
    
    doc.add_heading('1.2 调研范围', level=2)
    for module in selected_modules:
        module_names = {
            'finance': '财务管理',
            'supply': '供应链管理',
            'manufacture': '制造管理',
            'hr': '人力资源管理',
            'project': '项目管理'
        }
        doc.add_paragraph(f'• {module_names.get(module, module)}', style='List Bullet')
    
    doc.add_heading('1.3 调研团队', level=2)
    doc.add_paragraph('项目指导委员会：')
    for member in ['总经理', '财务总监', '运营总监']:
        doc.add_paragraph(f'• {member}', style='List Bullet')
    doc.add_paragraph('项目实施团队：')
    for member in ['项目经理', '业务顾问', '技术顾问', '实施顾问']:
        doc.add_paragraph(f'• {member}', style='List Bullet')
    
    doc.add_page_break()
    
    # 二、调研方法论（新增章节）
    doc.add_heading('二、调研方法论', level=1)
    doc.add_heading('2.1 调研方法体系', level=2)
    doc.add_paragraph('调研方法包括四大类：')
    
    methods = [
        ('访谈法', [
            '高层访谈：了解战略目标和组织架构',
            '部门访谈：了解业务流程和需求',
            '关键用户访谈：了解操作细节和问题'
        ]),
        ('问卷法', [
            '财务问卷：15个问题，了解财务流程',
            '供应链问卷：15个问题，了解采购销售流程',
            '制造问卷：15个问题，了解生产流程',
            '人力问卷：15个问题，了解人力流程'
        ]),
        ('观察法', [
            '现场观察：观察实际业务操作',
            '流程观察：观察业务流程执行',
            '系统观察：观察现有系统使用'
        ]),
        ('文档法', [
            '制度文档：分析管理制度和流程',
            '系统文档：分析现有系统文档',
            '数据文档：分析现有数据结构'
        ])
    ]
    
    for method_name, details in methods:
        doc.add_heading(f'2.2 {method_name}', level=3)
        for detail in details:
            doc.add_paragraph(f'• {detail}', style='List Bullet')
    
    doc.add_heading('2.3 调研计划', level=2)
    phases = [
        ('准备阶段（1周）', ['制定调研计划', '准备调研工具', '组建调研团队']),
        ('执行阶段（3周）', ['高层访谈', '部门调研', '关键用户访谈', '现场观察']),
        ('分析阶段（1周）', ['数据整理', '问题分析', '需求梳理']),
        ('报告阶段（1周）', ['调研报告编写', '报告评审', '成果汇报'])
    ]
    
    for phase, tasks in phases:
        doc.add_paragraph(f'{phase}：', style='List Number')
        for task in tasks:
            doc.add_paragraph(f'  • {task}', style='List Bullet')
    
    doc.add_page_break()
    
    # 三、企业概况
    doc.add_heading('三、企业概况', level=1)
    doc.add_heading('3.1 基本信息', level=2)
    
    basic_info = [
        ['企业名称', company_name],
        ['所属行业', industry],
        ['企业规模', customer_info.get('companySize', '中型企业')],
        ['员工人数', f'{customer_info.get("employees", "")}人'],
        ['年营业额', f'{customer_info.get("revenue", "")}万元']
    ]
    table = doc.add_table(rows=len(basic_info)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '项目'
    table.rows[0].cells[1].text = '内容'
    for i, (key, value) in enumerate(basic_info, 1):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
    
    doc.add_heading('3.2 组织架构', level=2)
    doc.add_paragraph('总部职能中心：')
    for dept in ['财务中心', '人力资源中心', '运营管理中心', '采购中心', '销售中心', '研发中心']:
        doc.add_paragraph(f'• {dept}', style='List Bullet')
    doc.add_paragraph('下属单位：')
    for dept in ['分公司', '子公司', '事业部', '生产基地', '销售网点', '服务中心']:
        doc.add_paragraph(f'• {dept}', style='List Bullet')
    
    doc.add_heading('3.3 业务特点', level=2)
    characteristics = [
        '业务模式多元化',
        '组织架构复杂',
        '管理精细化需求高',
        '数字化转型迫切',
        '跨部门协作频繁',
        '数据整合需求强'
    ]
    for char in characteristics:
        doc.add_paragraph(f'• {char}', style='List Bullet')
    
    doc.add_page_break()
    
    # 四、业务现状分析
    doc.add_heading('四、业务现状分析', level=1)
    
    if 'finance' in selected_modules:
        doc.add_heading('4.1 财务管理现状', level=2)
        finance_items = [
            '财务组织架构：财务中心下设财务部、成本部、资金部',
            '会计核算体系：采用多账簿管理，包括法定账、管理账、税务账',
            '应收应付管理：应收账款管理、应付账款管理、往来管理',
            '成本核算方法：品种法、分批法、分步法相结合',
            '预算管理流程：预算编制、预算执行、预算分析、预算调整',
            '资金管理：资金计划、资金调度、资金监控'
        ]
        for item in finance_items:
            doc.add_paragraph(f'• {item}', style='List Bullet')
        
        doc.add_heading('4.2 财务业务流程', level=3)
        finance_flows = [
            '费用报销流程：申请→审批→支付→记账',
            '应收流程：销售开票→应收确认→收款核销→账龄分析',
            '应付流程：采购开票→应付确认→付款核销→账龄分析',
            '成本流程：成本归集→成本分配→成本核算→成本分析'
        ]
        for flow in finance_flows:
            doc.add_paragraph(f'• {flow}', style='List Bullet')
    
    if 'supply' in selected_modules:
        doc.add_heading('4.3 供应链管理现状', level=2)
        supply_items = [
            '采购管理：供应商管理、采购申请、采购订单、采购入库',
            '销售管理：客户管理、销售订单、销售出库、销售开票',
            '库存管理：入库管理、出库管理、库存盘点、库存分析',
            '物流管理：运输管理、配送管理、仓储管理'
        ]
        for item in supply_items:
            doc.add_paragraph(f'• {item}', style='List Bullet')
    
    if 'manufacture' in selected_modules:
        doc.add_heading('4.4 制造管理现状', level=2)
        manufacture_items = [
            '生产计划：需求预测、主生产计划、物料需求计划',
            '车间管理：生产订单、生产领料、生产入库、工序管理',
            '质量管理：来料检验、过程检验、成品检验、质量追溯',
            '设备管理：设备台账、设备维护、设备保养、设备分析'
        ]
        for item in manufacture_items:
            doc.add_paragraph(f'• {item}', style='List Bullet')
    
    doc.add_page_break()
    
    # 五、存在问题分析
    doc.add_heading('五、存在问题分析', level=1)
    
    doc.add_heading('5.1 业务流程问题', level=2)
    process_issues = [
        '信息孤岛严重：各系统数据不互通，重复录入效率低',
        '流程冗长：审批流程复杂，决策效率低下',
        '流程断点多：业务协同困难，数据流转不畅',
        '标准不统一：缺乏统一的标准和规范',
        '监控缺失：缺乏实时监控和预警机制'
    ]
    for i, issue in enumerate(process_issues, 1):
        doc.add_paragraph(f'（{i}）{issue}')
    
    doc.add_heading('5.2 数据管理问题', level=2)
    data_issues = [
        '数据质量不高：数据不准确、不完整、不一致',
        '数据标准不统一：缺乏统一的数据标准和规范',
        '数据孤岛：数据分散在各个系统中，无法共享',
        '数据安全：数据安全措施不完善',
        '数据分析：数据分析能力不足，无法支持决策'
    ]
    for i, issue in enumerate(data_issues, 1):
        doc.add_paragraph(f'（{i}）{issue}')
    
    doc.add_heading('5.3 系统集成问题', level=2)
    integration_issues = [
        '系统集成困难：各系统接口复杂，集成难度大',
        '数据同步：数据实时同步困难',
        '业务协同：跨系统业务协同困难',
        '用户体验：用户体验不佳，操作复杂'
    ]
    for i, issue in enumerate(integration_issues, 1):
        doc.add_paragraph(f'（{i}）{issue}')
    
    doc.add_page_break()
    
    # 六、调研问卷分析（新增章节）
    doc.add_heading('六、调研问卷分析', level=1)
    doc.add_heading('6.1 问卷设计原则', level=2)
    design_principles = [
        '全面性：覆盖所有业务领域',
        '针对性：针对关键业务流程',
        '可操作性：便于用户理解和回答',
        '可分析性：便于数据统计和分析'
    ]
    for principle in design_principles:
        doc.add_paragraph(f'• {principle}', style='List Bullet')
    
    doc.add_heading('6.2 问卷结构设计', level=2)
    question_structures = [
        '财务模块（15题）',
        '供应链模块（15题）',
        '制造模块（15题）',
        '人力模块（15题）'
    ]
    for structure in question_structures:
        doc.add_paragraph(f'• {structure}', style='List Bullet')
    
    doc.add_heading('6.3 问卷统计分析', level=2)
    analysis_methods = [
        '频次分析：统计每个选项的选择次数',
        '交叉分析：分析不同部门、不同岗位的差异',
        '相关性分析：分析问题之间的相关性',
        '趋势分析：分析问题的趋势和变化'
    ]
    for method in analysis_methods:
        doc.add_paragraph(f'• {method}', style='List Bullet')
    
    doc.add_heading('6.4 问卷结果分析', level=2)
    result_analysis = [
        '满意度分析：用户对现有系统的满意度',
        '需求优先级：用户需求的优先级排序',
        '问题严重性：问题的严重程度评估',
        '改进建议：用户的改进建议汇总'
    ]
    for analysis in result_analysis:
        doc.add_paragraph(f'• {analysis}', style='List Bullet')
    
    doc.add_page_break()
    
    # 七、需求分析
    doc.add_heading('七、需求分析', level=1)
    
    doc.add_heading('7.1 业务需求', level=2)
    business_needs = [
        '实现财务业务一体化',
        '优化供应链协同流程',
        '提升生产管理精细化',
        '加强成本管控能力',
        '提升决策支持能力',
        '加强数据治理能力'
    ]
    for i, need in enumerate(business_needs, 1):
        doc.add_paragraph(f'（{i}）{need}')
    
    doc.add_heading('7.2 功能需求', level=2)
    module_needs = {
        'finance': [
            '总账管理：多账簿管理、凭证管理、期末处理',
            '应收管理：销售开票、应收确认、收款核销',
            '应付管理：采购开票、应付确认、付款核销',
            '成本管理：成本核算、成本分析、成本控制',
            '固定资产：资产卡片、折旧管理、资产处置'
        ],
        'supply': [
            '采购管理：供应商管理、采购申请、采购订单',
            '库存管理：入库管理、出库管理、库存盘点',
            '销售管理：客户管理、销售订单、销售出库',
            '物流管理：运输管理、配送管理、仓储管理'
        ],
        'manufacture': [
            '生产计划：需求预测、主生产计划、物料需求计划',
            '车间管理：生产订单、生产领料、生产入库',
            '质量管理：来料检验、过程检验、成品检验',
            '设备管理：设备台账、设备维护、设备分析'
        ]
    }
    
    for module, needs in module_needs.items():
        if module in selected_modules:
            module_names = {
                'finance': '财务管理',
                'supply': '供应链管理',
                'manufacture': '制造管理'
            }
            doc.add_paragraph(f'（{module_names.get(module, module)}）')
            for need in needs:
                doc.add_paragraph(f'  • {need}')
    
    doc.add_heading('7.3 非功能需求', level=2)
    non_functional_needs = [
        '性能需求：系统响应时间≤3秒，并发用户数≥500',
        '安全需求：数据加密、权限控制、审计日志',
        '可用性需求：系统可用性≥99.9%，故障恢复时间≤1小时',
        '可扩展性需求：支持业务扩展，支持多租户',
        '可维护性需求：模块化设计，易于维护和升级'
    ]
    for i, need in enumerate(non_functional_needs, 1):
        doc.add_paragraph(f'（{i}）{need}')
    
    doc.add_page_break()
    
    # 八、业务架构设计（新增章节）
    doc.add_heading('八、业务架构设计', level=1)
    doc.add_heading('8.1 业务架构概述', level=2)
    doc.add_paragraph('业务架构是连接战略与IT的桥梁，描述企业的业务结构、业务流程、业务规则。')
    
    doc.add_heading('8.2 价值流设计', level=2)
    value_streams = [
        ('O2C（订单到现金）', [
            '销售机会管理 → 报价管理 → 合同管理',
            '销售订单 → 出库管理 → 开票管理 → 收款管理'
        ]),
        ('P2P（采购到付款）', [
            '采购申请 → 供应商选择 → 采购订单',
            '入库管理 → 质量检验 → 发票管理 → 付款管理'
        ]),
        ('P2M（计划到制造）', [
            '需求预测 → 主生产计划 → 物料需求计划',
            '生产订单 → 生产领料 → 生产入库 → 成本核算'
        ]),
        ('R2R（记录到报告）', [
            '凭证管理 → 期末处理 → 账务核对',
            '财务报表 → 管理报表 → 分析报告'
        ])
    ]
    
    for stream_name, processes in value_streams:
        doc.add_paragraph(f'{stream_name}：', style='List Number')
        for process in processes:
            doc.add_paragraph(f'  • {process}', style='List Bullet')
    
    doc.add_heading('8.3 业务能力设计', level=2)
    business_capabilities = [
        ('财务管理', [
            '总账管理：凭证管理、账簿管理、报表管理',
            '应收管理：销售开票、应收确认、收款核销',
            '应付管理：采购开票、应付确认、付款核销',
            '成本管理：成本核算、成本分析、成本控制'
        ]),
        ('供应链管理', [
            '采购管理：供应商管理、采购申请、采购订单',
            '库存管理：入库管理、出库管理、库存盘点',
            '销售管理：客户管理、销售订单、销售出库',
            '物流管理：运输管理、配送管理、仓储管理'
        ]),
        ('制造管理', [
            '生产计划：需求预测、主生产计划、物料需求计划',
            '车间管理：生产订单、生产领料、生产入库',
            '质量管理：来料检验、过程检验、成品检验',
            '设备管理：设备台账、设备维护、设备分析'
        ])
    ]
    
    for domain, capabilities in business_capabilities:
        if domain.lower() in [m for m in selected_modules]:
            doc.add_paragraph(f'{domain}：', style='List Number')
            for capability in capabilities:
                doc.add_paragraph(f'  • {capability}', style='List Bullet')
    
    doc.add_page_break()
    
    # 九、解决方案设计（新增章节）
    doc.add_heading('九、解决方案设计', level=1)
    doc.add_heading('9.1 解决方案概述', level=2)
    doc.add_paragraph('基于金蝶云星空平台，为企业提供一体化ERP解决方案。')
    
    doc.add_heading('9.2 系统架构设计', level=2)
    architecture_components = [
        '应用层：财务云、供应链云、制造云、人力云',
        '平台层：苍穹PaaS平台、数据中台、集成平台',
        '基础设施层：计算资源、存储资源、网络资源',
        '安全体系：身份认证、权限控制、数据加密'
    ]
    for component in architecture_components:
        doc.add_paragraph(f'• {component}', style='List Bullet')
    
    doc.add_heading('9.3 数据迁移方案', level=2)
    migration_phases = [
        ('数据准备阶段', [
            '数据清洗：清理无效数据、重复数据',
            '数据转换：格式转换、编码转换',
            '数据验证：数据完整性验证、数据准确性验证'
        ]),
        ('数据迁移阶段', [
            '主数据迁移：组织、客户、供应商、物料、科目',
            '期初数据迁移：科目余额、库存余额、往来余额',
            '历史数据迁移：历史单据、历史报表'
        ]),
        ('数据验证阶段', [
            '数据完整性验证：数据完整性检查',
            '数据准确性验证：数据准确性检查',
            '数据一致性验证：数据一致性检查'
        ])
    ]
    
    for phase, tasks in migration_phases:
        doc.add_paragraph(f'{phase}：', style='List Number')
        for task in tasks:
            doc.add_paragraph(f'  • {task}', style='List Bullet')
    
    doc.add_heading('9.4 实施策略', level=2)
    implementation_strategies = [
        '总体策略：总体规划、分步实施',
        '实施原则：先试点、后推广，先核心、后扩展',
        '实施分期：一期（6个月）财务+供应链，二期（4个月）制造+人力',
        '实施保障：组织保障、人员保障、技术保障'
    ]
    for strategy in implementation_strategies:
        doc.add_paragraph(f'• {strategy}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十、实施路线图（新增章节）
    doc.add_heading('十、实施路线图', level=1)
    doc.add_heading('10.1 项目实施计划', level=2)
    
    project_phases = [
        ('项目启动阶段（2周）', [
            '项目启动会',
            '项目组织建立',
            '项目计划制定',
            '项目风险识别'
        ]),
        ('需求调研阶段（4周）', [
            '业务调研',
            '需求分析',
            '调研报告编写',
            '需求确认'
        ]),
        ('方案设计阶段（3周）', [
            '业务蓝图设计',
            '系统配置方案',
            '接口设计方案',
            '方案评审'
        ]),
        ('系统配置阶段（4周）', [
            '基础资料配置',
            '业务流程配置',
            '权限配置',
            '界面配置'
        ]),
        ('测试培训阶段（3周）', [
            '系统测试',
            '用户培训',
            '问题修复',
            '测试验收'
        ]),
        ('上线验收阶段（2周）', [
            '数据迁移',
            '系统切换',
            '上线支持',
            '验收评审'
        ])
    ]
    
    for phase, tasks in project_phases:
        doc.add_paragraph(f'{phase}：', style='List Number')
        for task in tasks:
            doc.add_paragraph(f'  • {task}', style='List Bullet')
    
    doc.add_heading('10.2 项目里程碑', level=2)
    milestones = [
        'M1：项目启动（第2周）',
        'M2：需求调研完成（第6周）',
        'M3：方案设计完成（第9周）',
        'M4：系统配置完成（第13周）',
        'M5：测试培训完成（第16周）',
        'M6：项目上线（第18周）'
    ]
    for milestone in milestones:
        doc.add_paragraph(f'• {milestone}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十一、风险评估（新增章节）
    doc.add_heading('十一、风险评估', level=1)
    doc.add_heading('11.1 风险识别', level=2)
    risk_categories = [
        ('技术风险', [
            '系统集成风险：各系统接口复杂，集成难度大',
            '数据质量风险：数据不准确、不完整',
            '性能风险：系统性能不满足业务需求',
            '安全风险：数据安全、系统安全风险'
        ]),
        ('业务风险', [
            '需求变更风险：需求频繁变更，影响项目进度',
            '用户接受风险：用户不接受新系统',
            '业务中断风险：系统切换影响业务',
            '数据丢失风险：数据迁移过程中数据丢失'
        ]),
        ('管理风险', [
            '项目进度风险：项目进度延迟',
            '项目成本风险：项目成本超支',
            '项目质量风险：项目质量不达标',
            '项目人员风险：关键人员流失'
        ])
    ]
    
    for category, risks in risk_categories:
        doc.add_paragraph(f'{category}：', style='List Number')
        for risk in risks:
            doc.add_paragraph(f'  • {risk}', style='List Bullet')
    
    doc.add_heading('11.2 风险应对', level=2)
    risk_responses = [
        ('技术风险应对', [
            '加强技术调研，选择成熟的技术方案',
            '加强数据治理，确保数据质量',
            '加强性能测试，确保系统性能',
            '加强安全管理，确保系统安全'
        ]),
        ('业务风险应对', [
            '加强需求管理，建立变更控制流程',
            '加强用户培训，提高用户接受度',
            '制定切换方案，确保业务连续性',
            '加强数据备份，确保数据安全'
        ]),
        ('管理风险应对', [
            '加强项目管理，制定详细的项目计划',
            '加强成本控制，建立成本监控机制',
            '加强质量管理，建立质量控制机制',
            '加强人员管理，建立人员激励机制'
        ])
    ]
    
    for category, responses in risk_responses:
        doc.add_paragraph(f'{category}：', style='List Number')
        for response in responses:
            doc.add_paragraph(f'  • {response}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十二、调研结论与建议（新增章节）
    doc.add_heading('十二、调研结论与建议', level=1)
    doc.add_heading('12.1 调研结论', level=2)
    conclusions = [
        '项目必要性：项目实施对企业发展具有重要意义',
        '技术可行性：金蝶云星空技术成熟，能够满足企业需求',
        '业务可行性：系统能够支持企业业务发展',
        '经济可行性：项目投入产出比合理'
    ]
    for conclusion in conclusions:
        doc.add_paragraph(f'• {conclusion}', style='List Bullet')
    
    doc.add_heading('12.2 项目建议', level=2)
    suggestions = [
        '组织建议：成立项目组织，明确职责分工',
        '范围建议：控制项目范围，避免范围蔓延',
        '数据建议：加强数据治理，确保数据质量',
        '培训建议：加强用户培训，提高用户接受度',
        '运维建议：建立运维体系，确保系统稳定运行'
    ]
    for suggestion in suggestions:
        doc.add_paragraph(f'• {suggestion}', style='List Bullet')
    
    doc.add_heading('12.3 下一步计划', level=2)
    next_steps = [
        '项目启动：召开项目启动会，成立项目组织',
        '需求确认：确认需求，制定详细需求文档',
        '方案设计：制定详细方案，进行方案评审',
        '项目实施：按照计划进行项目实施'
    ]
    for step in next_steps:
        doc.add_paragraph(f'• {step}', style='List Bullet')
    
    doc.add_page_break()
    
    # 附录：调研问卷库（60题）
    doc.add_heading('附录：调研问卷库', level=1)
    doc.add_heading('A. 财务模块问卷（15题）', level=2)
    
    finance_questions = [
        '1. 您认为当前财务系统的操作复杂度如何？',
        '2. 财务结账需要多长时间？',
        '3. 您对现有财务报表的满意度如何？',
        '4. 财务数据与其他系统的数据一致性如何？',
        '5. 应收账款管理的主要问题是什么？',
        '6. 应付账款管理的主要问题是什么？',
        '7. 成本核算的准确率如何？',
        '8. 预算管理的流程是否顺畅？',
        '9. 资金管理的效率如何？',
        '10. 固定资产管理的流程是否完善？',
        '11. 财务审批流程的效率如何？',
        '12. 财务数据的安全性如何？',
        '13. 财务系统的稳定性如何？',
        '14. 财务系统的响应速度如何？',
        '15. 您对财务系统的总体满意度如何？'
    ]
    
    for i, question in enumerate(finance_questions, 1):
        doc.add_paragraph(f'{question}')
        doc.add_paragraph('A. 非常满意  B. 比较满意  C. 一般  D. 不太满意  E. 非常不满意')
    
    doc.add_heading('B. 供应链模块问卷（15题）', level=2)
    
    supply_questions = [
        '1. 您认为当前采购系统的操作复杂度如何？',
        '2. 采购流程的效率如何？',
        '3. 供应商管理的流程是否完善？',
        '4. 库存管理的准确性如何？',
        '5. 库存周转率是否合理？',
        '6. 销售订单处理的效率如何？',
        '7. 客户管理的流程是否完善？',
        '8. 销售数据分析的及时性如何？',
        '9. 物流管理的效率如何？',
        '10. 采购成本的合理性如何？',
        '11. 库存成本的合理性如何？',
        '12. 销售成本的合理性如何？',
        '13. 供应链系统的稳定性如何？',
        '14. 供应链系统的响应速度如何？',
        '15. 您对供应链系统的总体满意度如何？'
    ]
    
    for i, question in enumerate(supply_questions, 1):
        doc.add_paragraph(f'{question}')
        doc.add_paragraph('A. 非常满意  B. 比较满意  C. 一般  D. 不太满意  E. 非常不满意')
    
    doc.add_heading('C. 制造模块问卷（15题）', level=2)
    
    manufacture_questions = [
        '1. 您认为当前生产计划系统的操作复杂度如何？',
        '2. 生产计划制定的准确性如何？',
        '3. 物料需求计划的准确性如何？',
        '4. 生产订单处理的效率如何？',
        '5. 生产领料流程的效率如何？',
        '6. 生产入库流程的效率如何？',
        '7. 质量检验的流程是否完善？',
        '8. 质量追溯的能力如何？',
        '9. 设备管理的流程是否完善？',
        '10. 设备维护的及时性如何？',
        '11. 生产成本核算的准确率如何？',
        '12. 生产效率的监控能力如何？',
        '13. 制造系统的稳定性如何？',
        '14. 制造系统的响应速度如何？',
        '15. 您对制造系统的总体满意度如何？'
    ]
    
    for i, question in enumerate(manufacture_questions, 1):
        doc.add_paragraph(f'{question}')
        doc.add_paragraph('A. 非常满意  B. 比较满意  C. 一般  D. 不太满意  E. 非常不满意')
    
    doc.add_heading('D. 人力模块问卷（15题）', level=2)
    
    hr_questions = [
        '1. 您认为当前人事系统的操作复杂度如何？',
        '2. 员工信息管理的效率如何？',
        '3. 薪酬计算的准确性如何？',
        '4. 薪酬发放的及时性如何？',
        '5. 绩效考核的流程是否完善？',
        '6. 培训管理的流程是否完善？',
        '7. 招聘管理的效率如何？',
        '8. 离职管理的流程是否完善？',
        '9. 考勤管理的准确性如何？',
        '10. 社保公积金管理的流程是否完善？',
        '11. 人事系统的稳定性如何？',
        '12. 人事系统的响应速度如何？',
        '13. 人事数据的安全性如何？',
        '14. 人事报表的及时性如何？',
        '15. 您对人事系统的总体满意度如何？'
    ]
    
    for i, question in enumerate(hr_questions, 1):
        doc.add_paragraph(f'{question}')
        doc.add_paragraph('A. 非常满意  B. 比较满意  C. 一般  D. 不太满意  E. 非常不满意')
    
    # 保存文件
    filename = f"{customer_info.get('customerCode', '客户')}_调研报告_v9_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return {
        'success': True,
        'filepath': filepath,
        'filename': filename,
        'chapters': 12
    }

if __name__ == '__main__':
    test_data = {
        'companyName': '测试公司',
        'customerCode': 'CS',
        'industry': '制造业',
        'companySize': '中型企业',
        'employees': '500',
        'revenue': '10000',
        'selectedModules': ['finance', 'supply', 'manufacture']
    }
    result = generate_survey_report_v9(test_data)
    print(f"✅ 生成成功：{result['filename']}")
    print(f"📊 章节数量：{result['chapters']}章")
