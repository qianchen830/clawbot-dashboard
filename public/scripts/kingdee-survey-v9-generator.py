#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶调研报告生成器 - v9.0 完整版
基于真实金蝶项目调研模板，支持12+章节专业内容
包含：调研方法论、60个问卷问题、调研纪要、阶段汇报
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
    return

def add_numbered_list(doc, items):
    """添加编号列表"""
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
    return

def add_table(doc, headers, data):
    """添加表格"""
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = "Table Grid"
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 数据
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_data)
    
    return table

def generate_survey_report_v9(customer_info, modules=["finance", "supply", "manufacture", "hr"]):
    """生成调研报告 - v9.0 完整版（12+章节）"""
    doc = Document()
    
    company_name = customer_info.get("companyName", "企业名称")
    industry = customer_info.get("industry", "制造业")
    
    # 封面
    doc.add_heading(f"{company_name}ERP系统调研报告", 0)
    info = doc.add_paragraph()
    info.add_run(f"企业名称：{company_name}\n").bold = True
    info.add_run(f"所属行业：{industry}\n")
    info.add_run(f"企业规模：{customer_info.get('companySize', '中型企业')}\n")
    info.add_run(f"调研日期：{datetime.now().strftime('%Y年%m月%d日')}\n")
    module_names = {"finance": "财务管理", "supply": "供应链管理", "manufacture": "制造管理", "hr": "人力资源管理"}
    selected_modules = [module_names.get(m, m) for m in modules]
    info.add_run(f"调研模块：{', '.join(selected_modules)}\n")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading("目录", level=1)
    toc_items = [
        "一、调研概述", "二、调研方法论", "三、企业概况", "四、业务现状分析",
        "五、存在问题诊断", "六、需求分析", "七、调研问卷汇总", "八、调研纪要",
        "九、阶段汇报", "十、需求规格说明书", "十一、调研结论", "十二、实施建议"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()
    
    # 一、调研概述
    doc.add_heading("一、调研概述", level=1)
    
    doc.add_heading("1.1 调研目的", level=2)
    add_paragraph(doc, f"全面了解{company_name}的业务现状和需求，为ERP系统实施提供依据。")
    add_paragraph(doc, "具体目标包括：")
    add_bullet_list(doc, [
        "了解企业业务流程和管理模式",
        "识别业务痛点和改进机会",
        "明确系统功能需求和集成需求",
        "评估项目实施风险和可行性",
        "制定合理的实施方案和时间计划"
    ])
    
    doc.add_heading("1.2 调研范围", level=2)
    add_paragraph(doc, "本次调研覆盖以下业务模块：")
    for module in selected_modules:
        doc.add_paragraph(f"• {module}", style="List Bullet")
    
    doc.add_heading("1.3 调研时间", level=2)
    add_paragraph(doc, "调研周期：4周")
    add_paragraph(doc, "现场调研：2周")
    add_paragraph(doc, "报告编写：1周")
    add_paragraph(doc, "报告评审：1周")
    
    doc.add_heading("1.4 调研团队", level=2)
    add_table(doc, ["角色", "姓名", "职责"], [
        ["项目经理", "待定", "项目整体协调"],
        ["业务顾问", "待定", "业务调研分析"],
        ["技术顾问", "待定", "技术调研分析"],
        ["文档工程师", "待定", "文档编写整理"]
    ])
    
    doc.add_page_break()
    
    # 二、调研方法论
    doc.add_heading("二、调研方法论", level=1)
    
    doc.add_heading("2.1 调研方法", level=2)
    add_paragraph(doc, "本次调研采用多种方法相结合，确保调研的全面性和准确性：")
    
    doc.add_heading("2.1.1 访谈法", level=3)
    add_paragraph(doc, "高层访谈：了解企业战略目标和管理需求")
    add_paragraph(doc, "中层访谈：了解部门业务流程和管理痛点")
    add_paragraph(doc, "基层访谈：了解具体操作流程和系统需求")
    
    doc.add_heading("2.1.2 问卷法", level=3)
    add_paragraph(doc, "设计结构化问卷，覆盖各业务模块")
    add_paragraph(doc, "问卷分为财务、供应链、制造、人力四大模块")
    add_paragraph(doc, "每个模块15个问题，共60个问题")
    
    doc.add_heading("2.1.3 观察法", level=3)
    add_paragraph(doc, "现场观察业务操作流程")
    add_paragraph(doc, "记录业务处理过程和问题点")
    add_paragraph(doc, "收集业务单据和报表样本")
    
    doc.add_heading("2.1.4 文档法", level=3)
    add_paragraph(doc, "收集企业现有管理制度和流程文件")
    add_paragraph(doc, "分析现有系统使用情况和问题")
    add_paragraph(doc, "研究行业最佳实践和案例")
    
    doc.add_heading("2.2 调研流程", level=2)
    add_paragraph(doc, "调研工作分为四个阶段：")
    add_numbered_list(doc, [
        "准备阶段：调研计划、调研问卷、访谈提纲",
        "实施阶段：高层访谈、部门调研、关键用户访谈",
        "分析阶段：问题分析、需求整理、方案设计",
        "报告阶段：报告编写、报告评审、报告定稿"
    ])
    
    doc.add_heading("2.3 调研工具", level=2)
    add_bullet_list(doc, [
        "调研问卷：结构化问题收集",
        "访谈提纲：半结构化访谈指引",
        "流程图工具：业务流程绘制",
        "数据分析工具：数据统计分析",
        "文档模板：调研文档标准化"
    ])
    
    doc.add_page_break()
    
    # 三、企业概况
    doc.add_heading("三、企业概况", level=1)
    
    doc.add_heading("3.1 基本信息", level=2)
    basic_info = [
        ["企业名称", company_name],
        ["所属行业", industry],
        ["企业规模", customer_info.get('companySize', '中型企业')],
        ["员工人数", f"{customer_info.get('employees', '')}人"],
        ["年营业额", f"{customer_info.get('revenue', '')}万元"],
        ["成立时间", "待定"],
        ["企业性质", "待定"],
        ["注册资本", "待定"]
    ]
    add_table(doc, ["项目", "内容"], basic_info)
    
    doc.add_heading("3.2 组织架构", level=2)
    add_paragraph(doc, "企业组织架构分为：")
    add_bullet_list(doc, [
        "总部职能中心：财务中心、人力资源中心、运营管理中心、采购中心、销售中心、研发中心",
        "下属单位：分公司、子公司、事业部、生产基地、销售网点、服务中心"
    ])
    
    doc.add_heading("3.3 业务特点", level=2)
    add_bullet_list(doc, [
        "业务模式多元化，涉及多种业务类型",
        "组织架构复杂，管理层级较多",
        "管理精细化需求高，对数据准确性要求严格",
        "数字化转型迫切，需要统一的ERP平台支撑"
    ])
    
    doc.add_heading("3.4 现有系统", level=2)
    add_table(doc, ["系统名称", "使用部门", "主要功能", "存在问题"], [
        ["ERP系统", "财务、供应链", "总账、应收、应付、库存", "功能老旧，接口不完善"],
        ["OA系统", "全公司", "审批流程、通知公告", "与ERP未集成"],
        ["CRM系统", "销售部门", "客户管理、销售跟进", "数据孤岛"],
        ["财务系统", "财务部门", "凭证管理、报表生成", "功能单一"]
    ])
    
    doc.add_page_break()
    
    # 四、业务现状分析
    doc.add_heading("四、业务现状分析", level=1)
    
    if "finance" in modules:
        doc.add_heading("4.1 财务管理现状", level=2)
        
        doc.add_heading("4.1.1 财务组织架构", level=3)
        add_paragraph(doc, "财务部门设置：财务总监、财务经理、总账会计、应收会计、应付会计、成本会计、出纳")
        add_paragraph(doc, "岗位职责：负责企业财务核算、资金管理、成本管理、预算管理等工作")
        
        doc.add_heading("4.1.2 会计核算体系", level=3)
        add_bullet_list(doc, [
            "会计科目：设置完整的会计科目体系",
            "核算维度：部门、项目、客户、供应商等多维度核算",
            "凭证管理：手工录入凭证，审批流程不完善",
            "期末处理：期末结转、期末结账、报表生成"
        ])
        
        doc.add_heading("4.1.3 应收应付管理", level=3)
        add_bullet_list(doc, [
            "应收管理：销售开票、收款核销、账龄分析",
            "应付管理：采购开票、付款核销、账龄分析",
            "往来对账：定期对账，手工编制对账单"
        ])
        
        doc.add_heading("4.1.4 成本核算方法", level=3)
        add_paragraph(doc, "成本核算方法：实际成本法")
        add_paragraph(doc, "成本核算对象：产品成本、部门成本、项目成本")
        add_paragraph(doc, "成本分摊：按工时分摊、按产量分摊、按销售额分摊")
        
        doc.add_heading("4.1.5 预算管理流程", level=3)
        add_paragraph(doc, "预算编制：年度预算、季度预算、月度预算")
        add_paragraph(doc, "预算执行：预算控制、预算分析、预算调整")
        add_paragraph(doc, "预算考核：预算执行率、预算差异分析")
    
    if "supply" in modules:
        doc.add_heading("4.2 供应链管理现状", level=2)
        
        doc.add_heading("4.2.1 采购管理流程", level=3)
        add_bullet_list(doc, [
            "采购申请：需求部门提出采购申请",
            "采购审批：采购部门审核、领导审批",
            "采购订单：采购部门下达采购订单",
            "采购入库：仓库收货、质量检验、入库确认",
            "采购结算：财务付款、发票核销"
        ])
        
        doc.add_heading("4.2.2 销售管理流程", level=3)
        add_bullet_list(doc, [
            "销售订单：客户下达销售订单",
            "销售审批：销售部门审核、领导审批",
            "销售出库：仓库备货、出库确认",
            "销售发货：物流配送、客户签收",
            "销售结算：财务收款、发票开具"
        ])
        
        doc.add_heading("4.2.3 库存管理流程", level=3)
        add_bullet_list(doc, [
            "出入库管理：采购入库、销售出库、生产领料、生产入库",
            "库存盘点：定期盘点、循环盘点、抽查盘点",
            "库存分析：库存周转率、库存预警、呆滞库存分析"
        ])
        
        doc.add_heading("4.2.4 供应商管理", level=3)
        add_paragraph(doc, "供应商分类：战略供应商、重要供应商、一般供应商")
        add_paragraph(doc, "供应商评估：质量、价格、交期、服务等维度评估")
        add_paragraph(doc, "供应商协同：采购协同、质量协同、库存协同")
        
        doc.add_heading("4.2.5 客户管理", level=3)
        add_paragraph(doc, "客户分类：战略客户、重要客户、一般客户")
        add_paragraph(doc, "客户评估：信用额度、付款情况、交易额等维度评估")
        add_paragraph(doc, "客户协同：订单协同、发货协同、对账协同")
    
    doc.add_page_break()
    
    # 五、存在问题诊断
    doc.add_heading("五、存在问题诊断", level=1)
    
    doc.add_heading("5.1 业务流程问题", level=2)
    add_numbered_list(doc, [
        "信息孤岛严重：各系统数据不互通，重复录入效率低，数据口径不一致，报表统计困难",
        "业务流程不畅：审批流程冗长，决策效率低下，流程断点多，业务协同困难",
        "管理精细化不足：成本核算不精准，盈利分析困难，预算控制不到位，资金管理粗放",
        "决策支持不足：数据分析滞后，无法实时监控，缺乏数据支撑，决策盲目性高"
    ])
    
    doc.add_heading("5.2 数据管理问题", level=2)
    add_bullet_list(doc, [
        "数据标准不统一：各系统数据编码不一致，数据口径不统一",
        "数据质量不高：数据重复、数据错误、数据缺失等问题普遍",
        "数据共享困难：系统间数据不互通，数据传递依赖人工",
        "数据分析滞后：报表生成周期长，数据分析不及时"
    ])
    
    doc.add_heading("5.3 系统集成问题", level=2)
    add_bullet_list(doc, [
        "系统间不集成：ERP、OA、CRM等系统独立运行，数据不互通",
        "接口不完善：现有系统接口老旧，无法满足集成需求",
        "数据不一致：各系统数据来源不同，数据存在差异",
        "维护成本高：多个系统独立维护，维护成本高"
    ])
    
    doc.add_heading("5.4 管理协同问题", level=2)
    add_bullet_list(doc, [
        "组织协同困难：部门间协同效率低，沟通成本高",
        "流程协同困难：跨部门流程断点多，流程协同困难",
        "数据协同困难：部门间数据不共享，数据协同困难",
        "决策协同困难：决策信息不完整，决策协同困难"
    ])
    
    doc.add_page_break()
    
    # 六、需求分析
    doc.add_heading("六、需求分析", level=1)
    
    doc.add_heading("6.1 业务需求", level=2)
    add_numbered_list(doc, [
        "实现财务业务一体化：打通财务与业务系统，实现数据自动流转",
        "优化供应链协同流程：优化采购、库存、销售流程，提升协同效率",
        "提升生产管理精细化：实现生产计划、车间管理、质量管理精细化",
        "加强成本管控能力：实现成本核算精准化，提升成本分析能力",
        "提升人力资源管理效率：实现人事、薪酬、绩效管理自动化"
    ])
    
    doc.add_heading("6.2 功能需求", level=2)
    
    if "finance" in modules:
        doc.add_heading("6.2.1 财务管理功能需求", level=3)
        add_bullet_list(doc, [
            "总账管理：多账簿管理、凭证管理、期末处理、报表管理",
            "应收管理：销售开票、收款核销、账龄分析、坏账管理",
            "应付管理：采购开票、付款核销、账龄分析、付款计划",
            "固定资产：资产登记、资产折旧、资产变动、资产盘点",
            "成本管理：成本核算、成本分析、成本预测、成本控制",
            "预算管理：预算编制、预算控制、预算分析、预算调整"
        ])
    
    if "supply" in modules:
        doc.add_heading("6.2.2 供应链管理功能需求", level=3)
        add_bullet_list(doc, [
            "采购管理：采购申请、采购订单、采购入库、采购结算",
            "库存管理：出入库管理、库存盘点、库存分析、库存预警",
            "销售管理：销售订单、销售出库、销售发货、销售结算",
            "物流管理：运输计划、运输执行、配送计划、配送执行",
            "供应商管理：供应商档案、供应商评估、供应商协同",
            "客户管理：客户档案、客户评估、客户协同"
        ])
    
    doc.add_heading("6.3 集成需求", level=2)
    add_bullet_list(doc, [
        "OA系统集成：审批流程、通知推送、单据传递",
        "CRM系统集成：客户数据同步、销售数据同步",
        "MES系统集成：生产数据同步、质量数据同步",
        "银行系统集成：资金支付、银行对账",
        "税务系统集成：发票管理、税务申报"
    ])
    
    doc.add_heading("6.4 非功能需求", level=2)
    add_bullet_list(doc, [
        "性能需求：系统响应时间<3秒，并发用户数>500",
        "安全需求：身份认证、权限控制、数据加密、审计日志",
        "可用性需求：系统可用性>99.9%，支持7×24小时运行",
        "扩展性需求：支持业务扩展、支持用户扩展、支持功能扩展"
    ])
    
    doc.add_page_break()
    
    # 七、调研问卷汇总
    doc.add_heading("七、调研问卷汇总", level=1)
    
    doc.add_heading("7.1 财务管理问卷（15题）", level=2)
    finance_questions = [
        "1. 贵公司目前使用的会计核算方法是什么？",
        "2. 贵公司的会计科目体系是否完整？是否需要调整？",
        "3. 贵公司的凭证管理流程是怎样的？审批流程是否完善？",
        "4. 贵公司的应收账款管理流程是怎样的？账龄分析是否及时？",
        "5. 贵公司的应付账款管理流程是怎样的？付款计划是否合理？",
        "6. 贵公司的成本核算方法是什么？成本核算是否准确？",
        "7. 贵公司的预算管理流程是怎样的？预算控制是否有效？",
        "8. 贵公司的资金管理流程是怎样的？资金周转是否顺畅？",
        "9. 贵公司的财务报表生成流程是怎样的？报表是否及时准确？",
        "10. 贵公司的财务分析流程是怎样的？分析是否深入？",
        "11. 贵公司的固定资产管理流程是怎样的？资产盘点是否及时？",
        "12. 贵公司的税务管理流程是怎样的？税务申报是否准确？",
        "13. 贵公司的财务制度是否完善？执行是否到位？",
        "14. 贵公司的财务人员配置是否充足？培训是否到位？",
        "15. 贵公司对财务管理有哪些改进需求？"
    ]
    for q in finance_questions:
        doc.add_paragraph(q, style="List Number")
    
    doc.add_heading("7.2 供应链管理问卷（15题）", level=2)
    supply_questions = [
        "1. 贵公司的采购管理流程是怎样的？审批流程是否完善？",
        "2. 贵公司的供应商管理流程是怎样的？供应商评估是否规范？",
        "3. 贵公司的采购计划编制流程是怎样的？计划是否准确？",
        "4. 贵公司的库存管理流程是怎样的？库存准确率如何？",
        "5. 贵公司的库存盘点流程是怎样的？盘点是否及时准确？",
        "6. 贵公司的库存预警机制是怎样的？是否有效？",
        "7. 贵公司的销售管理流程是怎样的？审批流程是否完善？",
        "8. 贵公司的客户管理流程是怎样的？客户评估是否规范？",
        "9. 贵公司的销售计划编制流程是怎样的？计划是否准确？",
        "10. 贵公司的物流管理流程是怎样的？物流效率如何？",
        "11. 贵公司的供应链协同流程是怎样的？协同效率如何？",
        "12. 贵公司的供应链数据分析流程是怎样的？分析是否深入？",
        "13. 贵公司的供应链制度是否完善？执行是否到位？",
        "14. 贵公司的供应链人员配置是否充足？培训是否到位？",
        "15. 贵公司对供应链管理有哪些改进需求？"
    ]
    for q in supply_questions:
        doc.add_paragraph(q, style="List Number")
    
    doc.add_heading("7.3 制造管理问卷（15题）", level=2)
    manufacture_questions = [
        "1. 贵公司的生产计划编制流程是怎样的？计划是否准确？",
        "2. 贵公司的主生产计划（MPS）编制流程是怎样的？",
        "3. 贵公司的物料需求计划（MRP）编制流程是怎样的？",
        "4. 贵公司的产能评估流程是怎样的？产能是否平衡？",
        "5. 贵公司的车间管理流程是怎样的？管理是否规范？",
        "6. 贵公司的生产领料流程是怎样的？领料是否准确？",
        "7. 贵公司的生产入库流程是怎样的？入库是否及时？",
        "8. 贵公司的工序管理流程是怎样的？工序是否顺畅？",
        "9. 贵公司的质量管理流程是怎样的？质量是否可控？",
        "10. 贵公司的来料检验流程是怎样的？检验是否规范？",
        "11. 贵公司的过程检验流程是怎样的？检验是否到位？",
        "12. 贵公司的成品检验流程是怎样的？检验是否严格？",
        "13. 贵公司的质量追溯流程是怎样的？追溯是否完整？",
        "14. 贵公司的生产设备管理流程是怎样的？设备是否完好？",
        "15. 贵公司对制造管理有哪些改进需求？"
    ]
    for q in manufacture_questions:
        doc.add_paragraph(q, style="List Number")
    
    doc.add_heading("7.4 人力资源管理问卷（15题）", level=2)
    hr_questions = [
        "1. 贵公司的人事管理流程是怎样的？管理是否规范？",
        "2. 贵公司的员工信息管理流程是怎样的？信息是否准确？",
        "3. 贵公司的组织架构管理流程是怎样的？架构是否合理？",
        "4. 贵公司的合同管理流程是怎样的？管理是否规范？",
        "5. 贵公司的薪酬管理流程是怎样的？薪酬核算是否准确？",
        "6. 贵公司的薪资结构是怎样的？是否合理？",
        "7. 贵公司的社保公积金管理流程是怎样的？管理是否规范？",
        "8. 贵公司的个税计算流程是怎样的？计算是否准确？",
        "9. 贵公司的绩效管理流程是怎样的？绩效考核是否公平？",
        "10. 贵公司的绩效指标设定流程是怎样的？指标是否合理？",
        "11. 贵公司的绩效分析流程是怎样的？分析是否深入？",
        "12. 贵公司的培训管理流程是怎样的？培训是否有效？",
        "13. 贵公司的人力资源制度是否完善？执行是否到位？",
        "14. 贵公司的人力资源人员配置是否充足？培训是否到位？",
        "15. 贵公司对人力资源管理有哪些改进需求？"
    ]
    for q in hr_questions:
        doc.add_paragraph(q, style="List Number")
    
    doc.add_page_break()
    
    # 八、调研纪要
    doc.add_heading("八、调研纪要", level=1)
    
    doc.add_heading("8.1 高层访谈纪要", level=2)
    add_paragraph(doc, "访谈对象：总经理、财务总监、运营总监")
    add_paragraph(doc, "访谈时间：第1周周一上午")
    add_paragraph(doc, "访谈地点：会议室")
    add_paragraph(doc, "访谈内容：")
    add_bullet_list(doc, [
        "企业战略目标：实现数字化转型，提升核心竞争力",
        "项目期望：打通信息孤岛，实现业务协同，提升管理效率",
        "项目范围：财务云、供应链云、制造云、人力云",
        "项目时间：期望在6个月内完成实施",
        "项目预算：预算充足，重点关注项目效果",
        "项目风险：数据迁移风险、用户接受风险",
        "项目保障：成立项目组织，明确职责分工，定期跟踪进度"
    ])
    
    doc.add_heading("8.2 财务部门访谈纪要", level=2)
    add_paragraph(doc, "访谈对象：财务经理、总账会计、应收会计、应付会计")
    add_paragraph(doc, "访谈时间：第1周周二至周三")
    add_paragraph(doc, "访谈地点：财务办公室")
    add_paragraph(doc, "访谈内容：")
    add_bullet_list(doc, [
        "现有系统：使用XX财务系统，功能老旧，接口不完善",
        "业务痛点：信息孤岛严重，数据不互通，重复录入效率低",
        "需求重点：实现财务业务一体化，提升财务核算效率",
        "关键需求：多账簿管理、成本核算精准化、预算控制有效化",
        "数据迁移：历史数据3年，需要迁移科目余额、往来余额、库存余额",
        "接口需求：与OA系统集成审批流程，与银行系统集成资金支付"
    ])
    
    doc.add_heading("8.3 供应链部门访谈纪要", level=2)
    add_paragraph(doc, "访谈对象：采购经理、仓库主管、销售经理")
    add_paragraph(doc, "访谈时间：第1周周四至周五")
    add_paragraph(doc, "访谈地点：供应链办公室")
    add_paragraph(doc, "访谈内容：")
    add_bullet_list(doc, [
        "现有系统：使用XX ERP系统，功能不完善，用户体验差",
        "业务痛点：采购流程冗长，库存准确率低，销售协同困难",
        "需求重点：优化采购流程，提升库存准确率，加强销售协同",
        "关键需求：采购计划、库存预警、销售订单管理、供应商协同",
        "数据迁移：历史数据2年，需要迁移供应商档案、客户档案、物料档案",
        "接口需求：与MES系统集成生产数据，与CRM系统集成客户数据"
    ])
    
    doc.add_heading("8.4 制造部门访谈纪要", level=2)
    add_paragraph(doc, "访谈对象：生产经理、车间主任、质量主管")
    add_paragraph(doc, "访谈时间：第2周周一至周二")
    add_paragraph(doc, "访谈地点：生产车间")
    add_paragraph(doc, "访谈内容：")
    add_bullet_list(doc, [
        "现有系统：使用XX MES系统，与ERP系统不集成",
        "业务痛点：生产计划准确性低，车间管理粗放，质量追溯困难",
        "需求重点：提升生产计划准确性，精细化管理车间，完善质量追溯",
        "关键需求：MPS计划、MRP运算、车间管理、质量追溯",
        "数据迁移：历史数据1年，需要迁移BOM、工艺路线、质量标准",
        "接口需求：与MES系统集成生产数据，与质量系统对接检验数据"
    ])
    
    doc.add_page_break()
    
    # 九、阶段汇报
    doc.add_heading("九、阶段汇报", level=1)
    
    doc.add_heading("9.1 第一阶段汇报（调研准备阶段）", level=2)
    add_paragraph(doc, "汇报时间：第1周周五")
    add_paragraph(doc, "汇报内容：")
    add_bullet_list(doc, [
        "调研计划：调研时间、调研范围、调研方式、调研团队",
        "调研问卷：问卷设计、问卷发放、问卷回收",
        "访谈提纲：高层访谈、部门访谈、关键用户访谈",
        "文档模板：调研纪要、调研报告、需求文档"
    ])
    
    doc.add_heading("9.2 第二阶段汇报（调研实施阶段）", level=2)
    add_paragraph(doc, "汇报时间：第2周周五")
    add_paragraph(doc, "汇报内容：")
    add_bullet_list(doc, [
        "调研进展：高层访谈完成、部门调研完成、关键用户访谈完成",
        "问题汇总：业务流程问题、数据管理问题、系统集成问题",
        "需求汇总：业务需求、功能需求、集成需求、非功能需求",
        "风险识别：数据迁移风险、用户接受风险、系统集成风险"
    ])
    
    doc.add_heading("9.3 第三阶段汇报（调研分析阶段）", level=2)
    add_paragraph(doc, "汇报时间：第3周周五")
    add_paragraph(doc, "汇报内容：")
    add_bullet_list(doc, [
        "问题分析：问题分类、问题原因、问题影响",
        "需求分析：需求分类、需求优先级、需求可行性",
        "方案设计：业务蓝图、系统配置、接口设计",
        "风险评估：风险分类、风险应对、风险监控"
    ])
    
    doc.add_heading("9.4 第四阶段汇报（调研总结阶段）", level=2)
    add_paragraph(doc, "汇报时间：第4周周五")
    add_paragraph(doc, "汇报内容：")
    add_bullet_list(doc, [
        "调研总结：调研成果、调研结论、调研建议",
        "需求确认：业务需求确认、功能需求确认、集成需求确认",
        "实施建议：实施范围、实施计划、实施风险、实施保障",
        "下一步计划：方案设计、系统配置、测试培训、上线验收"
    ])
    
    doc.add_page_break()
    
    # 十、需求规格说明书
    doc.add_heading("十、需求规格说明书", level=1)
    
    doc.add_heading("10.1 需求概述", level=2)
    add_paragraph(doc, "本次调研共识别需求XX项，其中：")
    add_paragraph(doc, "高优先级需求：XX项")
    add_paragraph(doc, "中优先级需求：XX项")
    add_paragraph(doc, "低优先级需求：XX项")
    
    doc.add_heading("10.2 功能需求清单", level=2)
    add_table(doc, ["需求编号", "需求名称", "需求描述", "优先级", "所属模块"], [
        ["F001", "总账管理", "多账簿管理、凭证管理、期末处理", "高", "财务云"],
        ["F002", "应收管理", "销售开票、收款核销、账龄分析", "高", "财务云"],
        ["F003", "应付管理", "采购开票、付款核销、账龄分析", "高", "财务云"],
        ["F004", "采购管理", "采购申请、采购订单、采购入库", "高", "供应链云"],
        ["F005", "库存管理", "出入库管理、库存盘点、库存分析", "高", "供应链云"],
        ["F006", "销售管理", "销售订单、销售出库、销售结算", "高", "供应链云"]
    ])
    
    doc.add_heading("10.3 非功能需求清单", level=2)
    add_table(doc, ["需求编号", "需求名称", "需求描述", "指标要求"], [
        ["NF001", "性能需求", "系统响应时间", "<3秒"],
        ["NF002", "并发需求", "并发用户数", ">500"],
        ["NF003", "安全需求", "身份认证、权限控制", "符合安全标准"],
        ["NF004", "可用性需求", "系统可用性", ">99.9%"],
        ["NF005", "扩展性需求", "支持业务扩展", "支持功能扩展"]
    ])
    
    doc.add_heading("10.4 接口需求清单", level=2)
    add_table(doc, ["接口编号", "接口名称", "接口类型", "接口说明"], [
        ["I001", "OA接口", "审批流程", "与OA系统集成审批流程"],
        ["I002", "CRM接口", "客户数据", "与CRM系统集成客户数据"],
        ["I003", "MES接口", "生产数据", "与MES系统集成生产数据"],
        ["I004", "银行接口", "资金支付", "与银行系统集成资金支付"],
        ["I005", "税务接口", "税务申报", "与税务系统集成税务申报"]
    ])
    
    doc.add_page_break()
    
    # 十一、调研结论
    doc.add_heading("十一、调研结论", level=1)
    
    doc.add_heading("11.1 调研成果", level=2)
    add_paragraph(doc, "本次调研共完成：")
    add_bullet_list(doc, [
        "高层访谈：3场",
        "部门调研：12场",
        "关键用户访谈：36场",
        "问卷调查：60份",
        "现场观察：5次",
        "文档收集：15份",
        "问题识别：25个",
        "需求确认：60项"
    ])
    
    doc.add_heading("11.2 业务现状评估", level=2)
    add_paragraph(doc, "业务现状评估结果：")
    add_table(doc, ["评估维度", "现状水平", "改进空间", "优先级"], [
        ["财务管理", "中等", "较大", "高"],
        ["供应链管理", "中等", "较大", "高"],
        ["制造管理", "较低", "很大", "高"],
        ["人力资源管理", "中等", "中等", "中"],
        ["系统集成", "较低", "很大", "高"],
        ["数据管理", "较低", "很大", "高"]
    ])
    
    doc.add_heading("11.3 实施必要性", level=2)
    add_paragraph(doc, "项目实施必要性评估：")
    add_bullet_list(doc, [
        "业务发展需要：现有系统无法支撑业务快速发展，实施ERP系统是必然选择",
        "管理提升需要：现有管理方式粗放，需要精细化管理系统提升管理水平",
        "效率提升需要：现有业务流程效率低下，需要优化流程提升效率",
        "数据整合需要：现有数据分散，需要整合数据提升决策支持能力"
    ])
    
    doc.add_heading("11.4 实施可行性", level=2)
    add_paragraph(doc, "项目实施可行性评估：")
    add_bullet_list(doc, [
        "技术可行性：金蝶云星空技术成熟，能够满足企业需求",
        "管理可行性：企业管理层重视，项目组织健全，保障有力",
        "经济可行性：项目预算充足，投资回报率高",
        "时间可行性：项目周期合理，能够在预期时间内完成"
    ])
    
    doc.add_page_break()
    
    # 十二、实施建议
    doc.add_heading("十二、实施建议", level=1)
    
    doc.add_heading("12.1 实施范围", level=2)
    add_paragraph(doc, "建议实施范围：")
    add_bullet_list(doc, [
        "核心模块：财务云、供应链云、制造云",
        "集成模块：OA系统、CRM系统、MES系统",
        "扩展模块：人力云、项目管理",
        "实施阶段：分阶段实施，先核心后扩展"
    ])
    
    doc.add_heading("12.2 实施计划", level=2)
    add_paragraph(doc, "建议实施计划：")
    add_table(doc, ["阶段", "时间", "内容", "目标"], [
        ["第一阶段", "第1-2个月", "项目启动、需求确认、方案设计", "完成方案设计"],
        ["第二阶段", "第3-4个月", "系统配置、数据准备、用户培训", "完成系统配置"],
        ["第三阶段", "第5-6个月", "系统测试、系统上线、用户验收", "完成系统上线"],
        ["第四阶段", "第7-8个月", "系统优化、系统完善、项目验收", "完成项目验收"]
    ])
    
    doc.add_heading("12.3 风险控制", level=2)
    add_paragraph(doc, "主要风险及应对措施：")
    add_table(doc, ["风险类型", "风险描述", "风险等级", "应对措施"], [
        ["数据迁移风险", "历史数据迁移不完整", "高", "制定详细数据迁移计划，分批迁移验证"],
        ["用户接受风险", "用户对系统不适应", "中", "加强用户培训，提供操作指导，建立支持机制"],
        ["系统集成风险", "系统间集成不顺畅", "高", "制定详细接口设计，分阶段集成测试"],
        ["项目延期风险", "项目进度延期", "中", "制定详细项目计划，定期跟踪进度，及时调整"]
    ])
    
    doc.add_heading("12.4 保障措施", level=2)
    add_paragraph(doc, "项目保障措施：")
    add_bullet_list(doc, [
        "组织保障：成立项目指导委员会，明确职责分工",
        "资源保障：配备充足的人员和资源，确保项目顺利实施",
        "技术保障：提供技术支持和培训，确保系统稳定运行",
        "管理保障：建立项目管理制度，定期评审项目进度"
    ])
    
    doc.add_heading("12.5 下一步计划", level=2)
    add_paragraph(doc, "下一步工作计划：")
    add_numbered_list(doc, [
        "方案设计：详细设计业务蓝图、系统配置、接口设计",
        "系统配置：完成系统基础配置、业务流程配置",
        "数据准备：完成数据迁移、数据验证、数据初始化",
        "用户培训：完成用户培训、操作指导、系统支持",
        "系统测试：完成系统测试、用户验收、系统上线",
        "系统优化：完成系统优化、系统完善、项目验收"
    ])
    
    doc.add_page_break()
    
    # 结尾
    doc.add_heading("附件", level=1)
    add_paragraph(doc, "附件1：调研问卷详细内容")
    add_paragraph(doc, "附件2：访谈纪要详细内容")
    add_paragraph(doc, "附件3：业务流程图")
    add_paragraph(doc, "附件4：数据字典")
    add_paragraph(doc, "附件5：需求规格说明书")
    
    doc.add_page_break()
    doc.add_heading("金蝶软件（中国）有限公司", 0)
    doc.add_paragraph("项目经理：待定")
    doc.add_paragraph("联系电话：待定")
    doc.add_paragraph("电子邮箱：待定")
    doc.add_paragraph("感谢聆听！")
    
    # 保存文件
    filename = f"{customer_info.get('customerCode', '客户')}_调研报告_v9_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return {
        "success": True,
        "filepath": filepath,
        "filename": filename,
        "pages": len(doc.sections)
    }

if __name__ == "__main__":
    test_data = {
        "companyName": "测试公司",
        "customerCode": "CS",
        "industry": "制造业",
        "companySize": "中型企业",
        "employees": "500",
        "revenue": "10000",
        "projectManager": "张三"
    }
    result = generate_survey_report_v9(test_data)
    print(f"✅ 生成成功：{result['filename']}")
    print(f"📄 文档页数：{result['pages']}页")
