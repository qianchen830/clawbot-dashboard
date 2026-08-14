#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调研纪要生成器 v1.0
基于中煤科工项目调研纪要模板
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import sys

class SurveyMinutesGenerator:
    """调研纪要生成器"""
    
    def __init__(self, customer_name, module, survey_date=None):
        self.customer_name = customer_name
        self.module = module
        self.survey_date = survey_date or datetime.now().strftime("%Y%m%d")
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(11)
    
    def _add_title(self, text, level=1):
        """添加标题"""
        if level == 1:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        elif level == 3:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _add_paragraph(self, text):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
    
    def _add_section(self, title, content):
        """添加章节"""
        self._add_title(title, 3)
        self._add_paragraph(content)
        self.doc.add_paragraph()
    
    def generate(self, survey_info):
        """生成调研纪要"""
        # 封面
        self._add_title(f"{self.customer_name}")
        self._add_title("新ERP管理系统项目")
        self._add_title("调研会议纪要")
        self.doc.add_paragraph()
        
        # 基本信息
        self._add_paragraph(f"调研时间：{survey_info.get('survey_date', '')}")
        self._add_paragraph(f"调研地点：{survey_info.get('survey_location', '')}")
        self._add_paragraph(f"调研主题：{self.module}")
        self.doc.add_paragraph()
        
        # 参会人员
        self._add_title("参会人员", 2)
        self._add_paragraph(f"客户方：{survey_info.get('customer_persons', '')}")
        self._add_paragraph(f"金蝶方：{survey_info.get('kingdee_persons', '')}")
        self.doc.add_paragraph()
        
        # 调研内容
        self._add_title("调研内容", 1)
        
        # 一、客户概况
        self._add_title("一、客户概况", 2)
        self._add_title("1.1 客户基本情况", 3)
        self._add_paragraph(survey_info.get('customer_basic', '（待填写）'))
        self.doc.add_paragraph()
        
        self._add_title("1.2 客户组织架构", 3)
        self._add_paragraph(survey_info.get('customer_org', '（待填写）'))
        self.doc.add_paragraph()
        
        # 二、业务现状
        self._add_title("二、业务现状", 2)
        self._add_paragraph("现状与流程说明：")
        self._add_paragraph(survey_info.get('business_status', '（待填写）'))
        self.doc.add_paragraph()
        
        # 三、主要问题
        self._add_title("三、主要问题", 2)
        problems = survey_info.get('problems', [])
        if problems:
            for i, problem in enumerate(problems, 1):
                self._add_paragraph(f"{i}. {problem}")
        else:
            self._add_paragraph("（待填写）")
        self.doc.add_paragraph()
        
        # 四、需求分析与建议
        self._add_title("四、需求分析与建议", 2)
        requirements = survey_info.get('requirements', [])
        if requirements:
            for i, req in enumerate(requirements, 1):
                self._add_paragraph(f"{i}. {req}")
        else:
            self._add_paragraph("（待填写）")
        self.doc.add_paragraph()
        
        # 五、涉及基础数据
        self._add_title("五、涉及基础数据", 2)
        self._add_paragraph(survey_info.get('basic_data', '（待填写）'))
        self.doc.add_paragraph()
        
        # 六、涉及外围系统对接
        self._add_title("六、涉及外围系统对接", 2)
        self._add_paragraph(survey_info.get('external_systems', '（待填写）'))
        self.doc.add_paragraph()
        
        # 七、后续行动
        self._add_title("七、后续行动", 2)
        actions = survey_info.get('actions', [])
        if actions:
            # 创建表格
            table = self.doc.add_table(rows=len(actions)+1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['序号', '行动项', '责任人', '完成日期']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for i, action in enumerate(actions, 1):
                table.cell(i, 0).text = str(i)
                table.cell(i, 1).text = action.get('item', '')
                table.cell(i, 2).text = action.get('owner', '')
                table.cell(i, 3).text = action.get('due_date', '')
        else:
            self._add_paragraph("（待填写）")
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 签字确认
        self._add_title("签字确认", 2)
        self._add_paragraph("客户方签字：________________  日期：________________")
        self._add_paragraph("金蝶方签字：________________  日期：________________")
        
        # 保存文档
        filename = f"{self.customer_name}_{self.module}_调研纪要_{self.survey_date}_V1.0.docx"
        self.doc.save(filename)
        return filename


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='调研纪要生成器 v1.0')
    parser.add_argument('--customer', required=True, help='客户名称')
    parser.add_argument('--module', required=True, help='调研模块')
    parser.add_argument('--date', help='调研日期 (YYYYMMDD)')
    parser.add_argument('--location', help='调研地点')
    parser.add_argument('--customer-persons', help='客户方参会人员')
    parser.add_argument('--kingdee-persons', help='金蝶方参会人员')
    
    args = parser.parse_args()
    
    generator = SurveyMinutesGenerator(
        customer_name=args.customer,
        module=args.module,
        survey_date=args.date
    )
    
    survey_info = {
        'survey_date': args.date or datetime.now().strftime("%Y年%m月%d日"),
        'survey_location': args.location or '',
        'customer_persons': args.customer_persons or '',
        'kingdee_persons': args.kingdee_persons or '',
        'customer_basic': '',
        'customer_org': '',
        'business_status': '',
        'problems': [],
        'requirements': [],
        'basic_data': '',
        'external_systems': '',
        'actions': []
    }
    
    filename = generator.generate(survey_info)
    print(f"调研纪要已生成: {filename}")


if __name__ == "__main__":
    main()
