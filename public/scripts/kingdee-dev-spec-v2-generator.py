#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶客户化开发设计说明书生成器 v2.0
专业版 - 完整的10大板块结构

特点：
- 专业Word格式排版（封面、文档控制、页眉页脚）
- 完整的10大板块结构（概述、需求、设计、数据库、接口、功能、安全、测试、部署、附录）
- 支持命令行参数（--company, --project, --module, --requirements, --output）
- 符合金蝶项目交付标准

使用方法：
python3 kingdee-dev-spec-v2-generator.py --company "公司名称" --project "项目名称" --module "模块名称" --requirements "需求描述" --output "输出路径"
"""

import os
import sys
import json
import argparse
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 输出目录
OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 版本信息
__version__ = "2.0.0"
__author__ = "ClawBot"


class KingdeeDevSpecGenerator:
    """金蝶客户化开发设计说明书生成器"""
    
    def __init__(self, company: str, project: str, module: str, requirements: str = ""):
        """
        初始化生成器
        
        Args:
            company: 公司名称
            project: 项目名称
            module: 模块名称
            requirements: 需求描述
        """
        self.company = company
        self.project = project
        self.module = module
        self.requirements = requirements
        self.doc = Document()
        self.timestamp = datetime.now()
        
        # 文档版本信息
        self.doc_version = "1.0"
        self.doc_status = "初稿"
        
        # 设置文档基础样式
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """设置文档基础样式"""
        # 设置页面
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)  # A4宽度
        section.page_height = Cm(29.7)  # A4高度
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置标题样式
        for i in range(1, 4):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = '黑体'
            heading_style.font.bold = True
            heading_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            if i == 1:
                heading_style.font.size = Pt(16)
            elif i == 2:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)
    
    def _set_cell_shading(self, cell, color: str):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _add_paragraph(self, text: str = "", bold: bool = False, font_size: int = 10.5, 
                       indent: bool = False, space_before: int = 0, space_after: int = 0):
        """添加段落"""
        p = self.doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        return p
    
    def _add_table(self, headers: List[str], rows: List[List[str]], col_widths: List[float] = None):
        """添加表格"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            cell.paragraphs[0].runs[0].font.name = '宋体'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            self._set_cell_shading(cell, 'D9E2F3')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置内容
        for i, row_data in enumerate(rows):
            row = table.rows[i + 1]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].runs[0].font.name = '宋体'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(width)
        
        return table
    
    def _add_bullet_list(self, items: List[str], font_size: int = 10.5):
        """添加项目符号列表"""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            for run in p.runs:
                run.font.size = Pt(font_size)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return self.doc
    
    def _add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def generate_cover(self):
        """生成封面"""
        # 添加空行
        for _ in range(6):
            self.doc.add_paragraph()
        
        # 文档标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("客户化开发设计说明书")
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 项目名称
        self.doc.add_paragraph()
        project_para = self.doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project_para.add_run(f"（{self.project}）")
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加空行
        for _ in range(6):
            self.doc.add_paragraph()
        
        # 项目信息表格
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ('项目名称', self.project),
            ('客户名称', self.company),
            ('文档版本', self.doc_version),
            ('文档状态', self.doc_status),
            ('编制日期', self.timestamp.strftime('%Y年%m月%d日'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.name = '宋体'
            info_table.rows[i].cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            info_table.rows[i].cells[0].width = Cm(4)
            
            info_table.rows[i].cells[1].text = value
            info_table.rows[i].cells[1].paragraphs[0].runs[0].font.name = '宋体'
            info_table.rows[i].cells[1].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            info_table.rows[i].cells[1].width = Cm(10)
        
        self._add_page_break()
    
    def generate_document_control(self):
        """生成文档控制页"""
        # 标题
        self.doc.add_heading('文档控制', 1)
        
        # 文档修订记录
        self.doc.add_heading('文档修订记录', 2)
        
        self._add_table(
            ['版本号', '修订日期', '修订人', '修订内容', '审核人'],
            [
                ['1.0', self.timestamp.strftime('%Y-%m-%d'), '项目经理', '初稿', '项目总监'],
            ],
            col_widths=[2, 3, 3, 5, 3]
        )
        
        self.doc.add_paragraph()
        
        # 文档分发记录
        self.doc.add_heading('文档分发记录', 2)
        
        self._add_table(
            ['接收人', '部门', '接收日期', '签收人'],
            [
                ['项目经理', '项目管理部', self.timestamp.strftime('%Y-%m-%d'), '________'],
                ['技术负责人', '技术部', self.timestamp.strftime('%Y-%m-%d'), '________'],
                ['业务负责人', '业务部', self.timestamp.strftime('%Y-%m-%d'), '________'],
                ['客户代表', self.company, self.timestamp.strftime('%Y-%m-%d'), '________'],
            ],
            col_widths=[3, 4, 3, 3]
        )
        
        self._add_page_break()
    
    def generate_toc(self):
        """生成目录"""
        self.doc.add_heading('目录', 1)
        
        toc_items = [
            '一、概述',
            '    1.1 项目背景',
            '    1.2 项目目标',
            '    1.3 项目范围',
            '二、需求分析',
            '    2.1 功能需求',
            '    2.2 非功能需求',
            '三、系统设计',
            '    3.1 架构设计',
            '    3.2 模块设计',
            '四、数据库设计',
            '    4.1 表结构设计',
            '    4.2 字段说明',
            '五、接口设计',
            '    5.1 输入输出设计',
            '    5.2 数据格式定义',
            '六、功能设计',
            '    6.1 功能清单',
            '    6.2 流程设计',
            '七、安全设计',
            '    7.1 权限控制',
            '    7.2 数据安全',
            '八、测试方案',
            '    8.1 测试用例',
            '    8.2 测试数据',
            '九、部署方案',
            '    9.1 环境要求',
            '    9.2 部署步骤',
            '十、附录',
            '    10.1 代码示例',
            '    10.2 参考文档',
        ]
        
        for item in toc_items:
            self._add_paragraph(item, font_size=11)
        
        self._add_page_break()
    
    def generate_overview(self):
        """生成第一章：概述"""
        self.doc.add_heading('一、概述', 1)
        
        # 1.1 项目背景
        self.doc.add_heading('1.1 项目背景', 2)
        self._add_paragraph(
            f'{self.company}{self.project}客户化开发项目是在金蝶云·星空ERP系统基础上，'
            f'根据企业实际业务需求进行的定制化开发。'
            f'本次开发主要针对{self.module}模块，旨在解决企业在{self.module}业务管理中的实际问题，'
            '提升业务处理效率，优化管理流程。',
            indent=True
        )
        
        if self.requirements:
            self._add_paragraph(f'具体需求：{self.requirements}', indent=True)
        
        self._add_paragraph(
            '通过本次客户化开发，企业将实现业务流程的数字化、自动化，'
            '提高数据准确性和实时性，为管理层决策提供有力支持。',
            indent=True
        )
        
        # 1.2 项目目标
        self.doc.add_heading('1.2 项目目标', 2)
        self._add_paragraph('本次客户化开发项目的主要目标包括：', indent=True)
        
        self._add_bullet_list([
            '实现业务流程自动化：通过系统自动化处理，减少人工操作，提高工作效率',
            '提升数据准确性：通过系统校验和控制，确保数据准确、完整、一致',
            '优化管理流程：通过系统固化业务流程，规范操作步骤，提升管理水平',
            '增强系统功能：通过二次开发，扩展系统功能，满足企业个性化需求',
            '提高用户体验：优化操作界面，简化操作步骤，提升用户使用体验',
        ])
        
        # 1.3 项目范围
        self.doc.add_heading('1.3 项目范围', 2)
        self._add_paragraph('本次客户化开发项目的范围包括：', indent=True)
        
        self._add_table(
            ['范围类别', '具体内容'],
            [
                ['功能模块', self.module],
                ['开发类型', '客户化开发、接口开发、报表开发'],
                ['涉及系统', '金蝶云·星空ERP系统'],
                ['业务部门', '相关业务部门'],
                ['用户范围', '系统操作人员、管理人员'],
            ],
            col_widths=[4, 12]
        )
        
        self._add_paragraph()
        self._add_paragraph('项目边界说明：', bold=True)
        self._add_bullet_list([
            '本次开发仅针对指定功能模块，不涉及其他模块的改动',
            '开发工作基于现有系统架构，不改变系统基础架构',
            '数据迁移工作不在本次开发范围内',
            '用户培训工作另行安排',
        ])
        
        self._add_page_break()
    
    def generate_requirements(self):
        """生成第二章：需求分析"""
        self.doc.add_heading('二、需求分析', 1)
        
        # 2.1 功能需求
        self.doc.add_heading('2.1 功能需求', 2)
        self._add_paragraph('根据业务调研和需求分析，本次开发需要实现以下功能需求：', indent=True)
        
        self._add_table(
            ['需求编号', '需求名称', '需求描述', '优先级', '来源'],
            [
                ['F001', '基础数据管理', '实现基础数据的增删改查功能', '高', '业务部门'],
                ['F002', '业务单据处理', '实现业务单据的录入、审核、查询功能', '高', '业务部门'],
                ['F003', '报表查询', '实现业务报表的生成和查询功能', '中', '管理层'],
                ['F004', '数据导入导出', '实现数据的批量导入导出功能', '中', '业务部门'],
                ['F005', '权限控制', '实现按角色和用户的功能权限控制', '高', '信息部'],
            ],
            col_widths=[2, 3, 5, 2, 2]
        )
        
        self.doc.add_paragraph()
        
        # 详细功能需求描述
        self._add_paragraph('功能需求详细说明：', bold=True)
        
        self._add_paragraph('F001 基础数据管理：', bold=True)
        self._add_bullet_list([
            '支持基础数据的录入、修改、删除、查询操作',
            '支持数据校验，确保数据完整性和准确性',
            '支持数据批量操作，提高操作效率',
            '支持数据版本管理，记录数据变更历史',
        ])
        
        self._add_paragraph('F002 业务单据处理：', bold=True)
        self._add_bullet_list([
            '支持业务单据的录入、审核、反审核操作',
            '支持单据流程审批，实现业务流程控制',
            '支持单据查询和统计分析',
            '支持单据打印和导出功能',
        ])
        
        # 2.2 非功能需求
        self.doc.add_heading('2.2 非功能需求', 2)
        self._add_paragraph('除功能需求外，系统还需满足以下非功能需求：', indent=True)
        
        self._add_table(
            ['需求类别', '需求描述', '指标要求'],
            [
                ['性能需求', '系统响应时间要求', '页面加载≤3秒，查询响应≤5秒'],
                ['并发需求', '系统并发用户数要求', '支持≥100并发用户'],
                ['可用性需求', '系统可用性要求', '系统可用率≥99.5%'],
                ['安全性需求', '数据安全要求', '数据加密存储，权限严格控制'],
                ['易用性需求', '用户操作要求', '界面友好，操作简便'],
                ['扩展性需求', '系统扩展要求', '支持功能扩展和二次开发'],
            ],
            col_widths=[3, 5, 6]
        )
        
        self._add_paragraph()
        
        self._add_paragraph('性能指标详细说明：', bold=True)
        self._add_bullet_list([
            '响应时间：简单查询≤3秒，复杂查询≤5秒，报表生成≤10秒',
            '并发用户：系统支持≥100并发用户同时在线操作',
            '数据处理：支持单次处理≥10000条数据',
            '系统稳定性：7×24小时稳定运行，故障恢复时间≤30分钟',
        ])
        
        self._add_page_break()
    
    def generate_system_design(self):
        """生成第三章：系统设计"""
        self.doc.add_heading('三、系统设计', 1)
        
        # 3.1 架构设计
        self.doc.add_heading('3.1 架构设计', 2)
        self._add_paragraph('系统架构设计遵循金蝶云·星空技术架构标准，采用云原生架构设计。', indent=True)
        
        self._add_paragraph('3.1.1 总体架构', bold=True)
        self._add_paragraph(
            '系统采用分层架构设计，包括表现层、业务逻辑层、数据访问层和数据层。'
            '各层之间通过标准接口进行通信，实现松耦合设计。',
            indent=True
        )
        
        self._add_table(
            ['架构层次', '主要功能', '技术实现'],
            [
                ['表现层', '用户界面展示、用户交互', '金蝶云·星空前端框架'],
                ['业务逻辑层', '业务规则处理、流程控制', '金蝶云·星空业务中台'],
                ['数据访问层', '数据访问、数据缓存', '金蝶云·星空数据服务'],
                ['数据层', '数据存储、数据管理', '关系型数据库'],
            ],
            col_widths=[3, 5, 6]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('3.1.2 技术选型', bold=True)
        self._add_bullet_list([
            '前端技术：HTML5、JavaScript、CSS3、Vue.js',
            '后端技术：.NET Core、C#、WebAPI',
            '数据库：SQL Server（云端）',
            '开发平台：金蝶云·星空BOS平台',
            '开发工具：金蝶云·星空开发工具、Visual Studio',
        ])
        
        # 3.2 模块设计
        self.doc.add_heading('3.2 模块设计', 2)
        self._add_paragraph(f'本次开发针对{self.module}模块，模块设计如下：', indent=True)
        
        self._add_table(
            ['模块编号', '模块名称', '模块功能', '依赖模块'],
            [
                ['M001', '基础数据管理', '基础数据的增删改查', '系统基础模块'],
                ['M002', '业务单据处理', '业务单据的录入、审核', '基础数据管理'],
                ['M003', '报表查询', '业务报表的生成查询', '业务单据处理'],
                ['M004', '数据接口', '对外数据接口服务', '业务单据处理'],
                ['M005', '权限管理', '功能权限和数据权限', '系统权限模块'],
            ],
            col_widths=[2, 4, 6, 3]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('模块间调用关系：', bold=True)
        self._add_bullet_list([
            '基础数据管理模块为其他模块提供基础数据支持',
            '业务单据处理模块依赖基础数据管理模块',
            '报表查询模块从业务单据处理模块获取数据',
            '数据接口模块调用业务单据处理模块提供对外服务',
            '权限管理模块控制所有模块的访问权限',
        ])
        
        self._add_page_break()
    
    def generate_database_design(self):
        """生成第四章：数据库设计"""
        self.doc.add_heading('四、数据库设计', 1)
        
        # 4.1 表结构设计
        self.doc.add_heading('4.1 表结构设计', 2)
        self._add_paragraph('根据功能需求，设计以下数据库表：', indent=True)
        
        self._add_table(
            ['表名', '中文名称', '说明', '记录数预估'],
            [
                ['T_BD_BASEDATA', '基础数据表', '存储基础数据信息', '10000'],
                ['T_BD_BILL', '业务单据表', '存储业务单据信息', '100000'],
                ['T_BD_BILLDTL', '单据明细表', '存储单据明细信息', '500000'],
                ['T_BD_LOG', '操作日志表', '记录操作日志', '1000000'],
                ['T_BD_CONFIG', '配置参数表', '存储系统配置参数', '100'],
            ],
            col_widths=[4, 3, 6, 3]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('表设计原则：', bold=True)
        self._add_bullet_list([
            '遵循金蝶云·星空数据库设计规范',
            '使用统一的表命名规则（T_BD_前缀）',
            '主键使用bigint类型自增ID',
            '所有表必须包含创建时间、修改时间、创建人、修改人字段',
            '建立合理的索引，提高查询效率',
        ])
        
        # 4.2 字段说明
        self.doc.add_heading('4.2 字段说明', 2)
        self._add_paragraph('以基础数据表（T_BD_BASEDATA）为例，字段说明如下：', indent=True)
        
        self._add_table(
            ['字段名', '中文名', '类型', '长度', '必填', '说明'],
            [
                ['FID', '主键ID', 'bigint', '8', '是', '自增主键'],
                ['FNUMBER', '编码', 'nvarchar', '50', '是', '数据编码，唯一'],
                ['FNAME', '名称', 'nvarchar', '200', '是', '数据名称'],
                ['FDESCRIPTION', '描述', 'nvarchar', '500', '否', '数据描述'],
                ['FSTATUS', '状态', 'int', '4', '是', '数据状态：1-有效，0-无效'],
                ['FCREATEDATE', '创建时间', 'datetime', '8', '是', '数据创建时间'],
                ['FMODIFYDATE', '修改时间', 'datetime', '8', '否', '数据修改时间'],
                ['FCREATORID', '创建人', 'bigint', '8', '是', '创建人ID'],
                ['FMODIFIERID', '修改人', 'bigint', '8', '否', '修改人ID'],
            ],
            col_widths=[3, 2, 2, 2, 1, 6]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('索引设计：', bold=True)
        self._add_table(
            ['索引名', '索引类型', '索引字段', '说明'],
            [
                ['PK_T_BD_BASEDATA', '主键索引', 'FID', '主键聚集索引'],
                ['IX_T_BD_BASEDATA_NUMBER', '唯一索引', 'FNUMBER', '编码唯一索引'],
                ['IX_T_BD_BASEDATA_NAME', '普通索引', 'FNAME', '名称普通索引'],
                ['IX_T_BD_BASEDATA_STATUS', '普通索引', 'FSTATUS', '状态普通索引'],
            ],
            col_widths=[5, 2, 3, 4]
        )
        
        self._add_page_break()
    
    def generate_interface_design(self):
        """生成第五章：接口设计"""
        self.doc.add_heading('五、接口设计', 1)
        
        # 5.1 输入输出设计
        self.doc.add_heading('5.1 输入输出设计', 2)
        self._add_paragraph('系统对外提供标准REST API接口，支持JSON数据格式。', indent=True)
        
        self._add_paragraph('5.1.1 接口清单', bold=True)
        self._add_table(
            ['接口编号', '接口名称', '接口地址', '请求方式', '说明'],
            [
                ['API001', '基础数据查询', '/api/basedata/query', 'POST', '查询基础数据'],
                ['API002', '基础数据新增', '/api/basedata/create', 'POST', '新增基础数据'],
                ['API003', '基础数据修改', '/api/basedata/update', 'POST', '修改基础数据'],
                ['API004', '基础数据删除', '/api/basedata/delete', 'POST', '删除基础数据'],
                ['API005', '业务单据查询', '/api/bill/query', 'POST', '查询业务单据'],
                ['API006', '业务单据提交', '/api/bill/submit', 'POST', '提交业务单据'],
            ],
            col_widths=[2, 3, 4, 2, 3]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('5.1.2 输入参数示例（基础数据查询）', bold=True)
        self._add_paragraph(
            '{\n'
            '  "pageNum": 1,\n'
            '  "pageSize": 20,\n'
            '  "filter": {\n'
            '    "number": "BD001",\n'
            '    "name": "基础数据",\n'
            '    "status": 1\n'
            '  }\n'
            '}',
            font_size=9
        )
        
        self._add_paragraph()
        
        self._add_paragraph('5.1.3 输出参数示例', bold=True)
        self._add_paragraph(
            '{\n'
            '  "code": 200,\n'
            '  "message": "操作成功",\n'
            '  "data": {\n'
            '    "total": 100,\n'
            '    "rows": [\n'
            '      {\n'
            '        "id": 1,\n'
            '        "number": "BD001",\n'
            '        "name": "基础数据1",\n'
            '        "status": 1\n'
            '      }\n'
            '    ]\n'
            '  }\n'
            '}',
            font_size=9
        )
        
        # 5.2 数据格式定义
        self.doc.add_heading('5.2 数据格式定义', 2)
        self._add_paragraph('接口数据格式遵循以下规范：', indent=True)
        
        self._add_paragraph('5.2.1 公共参数', bold=True)
        self._add_table(
            ['参数名', '类型', '必填', '说明'],
            [
                ['token', 'string', '是', '用户认证令牌'],
                ['timestamp', 'long', '是', '请求时间戳（毫秒）'],
                ['sign', 'string', '是', '签名（MD5加密）'],
            ],
            col_widths=[3, 2, 2, 9]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('5.2.2 返回码定义', bold=True)
        self._add_table(
            ['返回码', '说明', '处理建议'],
            [
                ['200', '操作成功', '正常处理'],
                ['400', '参数错误', '检查参数是否正确'],
                ['401', '认证失败', '检查token是否有效'],
                ['403', '权限不足', '检查用户权限'],
                ['500', '系统错误', '联系技术支持'],
            ],
            col_widths=[2, 4, 8]
        )
        
        self._add_paragraph()
        
        self._add_paragraph('5.2.3 数据字典', bold=True)
        self._add_table(
            ['字典类型', '字典编码', '字典名称', '说明'],
            [
                ['状态', '1', '有效', '数据有效状态'],
                ['状态', '0', '无效', '数据无效状态'],
                ['审核状态', 'A', '已审核', '单据已审核'],
                ['审核状态', 'B', '未审核', '单据未审核'],
                ['审核状态', 'C', '已关闭', '单据已关闭'],
            ],
            col_widths=[3, 3, 3, 7]
        )
        
        self._add_page_break()
    
    def generate_function_design(self):
        """生成第六章：功能设计"""
        self.doc.add_heading('六、功能设计', 1)
        
        # 6.1 功能清单
        self.doc.add_heading('6.1 功能清单', 2)
        self._add_paragraph('本次开发包含以下功能：', indent=True)
        
        self._add_table(
            ['功能编号', '功能名称', '功能描述', '开发类型', '优先级'],
            [
                ['FN001', '基础数据录入', '录入基础数据信息', '表单开发', '高'],
                ['FN002', '基础数据查询', '查询基础数据信息', '查询开发', '高'],
                ['FN003', '基础数据修改', '修改基础数据信息', '表单开发', '高'],
                ['FN004', '基础数据删除', '删除基础数据信息', '表单开发', '中'],
                ['FN005', '业务单据录入', '录入业务单据信息', '表单开发', '高'],
                ['FN006', '业务单据审核', '审核业务单据', '流程开发', '高'],
                ['FN007', '业务报表生成', '生成业务报表', '报表开发', '中'],
                ['FN008', '数据导入', '批量导入数据', '接口开发', '中'],
                ['FN009', '数据导出', '批量导出数据', '接口开发', '中'],
                ['FN010', '权限配置', '配置功能权限', '权限开发', '高'],
            ],
            col_widths=[2, 3, 4, 2, 2]
        )
        
        self.doc.add_paragraph()
        
        # 6.2 流程设计
        self.doc.add_heading('6.2 流程设计', 2)
        self._add_paragraph('主要业务流程设计如下：', indent=True)
        
        self._add_paragraph('6.2.1 基础数据管理流程', bold=True)
        self._add_table(
            ['步骤', '操作', '操作人', '说明'],
            [
                ['1', '数据录入', '业务员', '录入基础数据信息'],
                ['2', '数据校验', '系统', '系统自动校验数据'],
                ['3', '数据保存', '系统', '保存数据到数据库'],
                ['4', '数据生效', '系统', '数据生效，可供使用'],
            ],
            col_widths=[2, 3, 2, 9]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('6.2.2 业务单据处理流程', bold=True)
        self._add_table(
            ['步骤', '操作', '操作人', '说明'],
            [
                ['1', '单据录入', '业务员', '录入业务单据信息'],
                ['2', '数据校验', '系统', '系统自动校验数据'],
                ['3', '提交审批', '业务员', '提交单据进行审批'],
                ['4', '审批处理', '审批人', '审批单据'],
                ['5', '审批通过', '审批人', '单据审核通过'],
                ['6', '单据生效', '系统', '单据生效，执行业务'],
            ],
            col_widths=[2, 3, 2, 9]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('流程说明：', bold=True)
        self._add_bullet_list([
            '基础数据管理流程：简单流程，录入后直接生效',
            '业务单据处理流程：需要审批流程，审批通过后生效',
            '所有流程都记录操作日志，便于追溯',
            '支持流程回退和重新提交',
        ])
        
        self._add_page_break()
    
    def generate_security_design(self):
        """生成第七章：安全设计"""
        self.doc.add_heading('七、安全设计', 1)
        
        # 7.1 权限控制
        self.doc.add_heading('7.1 权限控制', 2)
        self._add_paragraph('系统采用基于角色的权限控制（RBAC）模型，实现细粒度的权限管理。', indent=True)
        
        self._add_paragraph('7.1.1 角色设计', bold=True)
        self._add_table(
            ['角色编码', '角色名称', '角色说明', '权限范围'],
            [
                ['ROLE_ADMIN', '系统管理员', '系统管理角色', '所有功能权限'],
                ['ROLE_MANAGER', '业务经理', '业务管理角色', '业务管理和审批权限'],
                ['ROLE_USER', '业务员', '业务操作角色', '业务录入和查询权限'],
                ['ROLE_VIEWER', '查看者', '只读角色', '仅查询权限'],
            ],
            col_widths=[3, 3, 3, 7]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('7.1.2 权限设计', bold=True)
        self._add_table(
            ['权限编码', '权限名称', '权限说明', '所属角色'],
            [
                ['PERM_BASEDATA_ADD', '基础数据新增', '新增基础数据权限', 'ROLE_ADMIN, ROLE_USER'],
                ['PERM_BASEDATA_EDIT', '基础数据修改', '修改基础数据权限', 'ROLE_ADMIN, ROLE_USER'],
                ['PERM_BASEDATA_DELETE', '基础数据删除', '删除基础数据权限', 'ROLE_ADMIN'],
                ['PERM_BILL_APPROVE', '单据审批', '审批业务单据权限', 'ROLE_MANAGER'],
                ['PERM_REPORT_VIEW', '报表查看', '查看报表权限', 'ROLE_ADMIN, ROLE_MANAGER'],
            ],
            col_widths=[4, 3, 4, 5]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('7.1.3 数据权限', bold=True)
        self._add_bullet_list([
            '组织权限：按组织机构控制数据访问范围',
            '部门权限：按部门控制数据访问范围',
            '人员权限：按人员控制数据访问范围',
            '自定义权限：支持自定义数据权限规则',
        ])
        
        # 7.2 数据安全
        self.doc.add_heading('7.2 数据安全', 2)
        self._add_paragraph('系统从多个维度保障数据安全：', indent=True)
        
        self._add_paragraph('7.2.1 数据加密', bold=True)
        self._add_bullet_list([
            '传输加密：使用HTTPS协议传输数据，确保数据传输安全',
            '存储加密：敏感数据采用AES加密存储',
            '密码加密：用户密码采用SHA256+盐值加密',
        ])
        
        self._add_paragraph('7.2.2 数据备份', bold=True)
        self._add_bullet_list([
            '自动备份：每天自动备份数据库',
            '增量备份：每小时增量备份关键数据',
            '异地备份：数据备份到异地机房',
            '备份恢复：支持数据快速恢复',
        ])
        
        self._add_paragraph('7.2.3 数据审计', bold=True)
        self._add_bullet_list([
            '操作日志：记录所有用户操作日志',
            '数据变更：记录数据变更历史',
            '登录日志：记录用户登录日志',
            '异常监控：监控系统异常行为',
        ])
        
        self._add_page_break()
    
    def generate_test_plan(self):
        """生成第八章：测试方案"""
        self.doc.add_heading('八、测试方案', 1)
        
        # 8.1 测试用例
        self.doc.add_heading('8.1 测试用例', 2)
        self._add_paragraph('根据功能需求，设计以下测试用例：', indent=True)
        
        self._add_paragraph('8.1.1 功能测试用例', bold=True)
        self._add_table(
            ['用例编号', '用例名称', '测试步骤', '预期结果', '优先级'],
            [
                ['TC001', '基础数据新增', '1.打开新增界面\n2.录入数据\n3.点击保存', '数据保存成功', '高'],
                ['TC002', '基础数据查询', '1.打开查询界面\n2.输入查询条件\n3.点击查询', '显示查询结果', '高'],
                ['TC003', '基础数据修改', '1.选择数据\n2.点击修改\n3.修改数据\n4.保存', '数据修改成功', '高'],
                ['TC004', '基础数据删除', '1.选择数据\n2.点击删除\n3.确认删除', '数据删除成功', '中'],
                ['TC005', '业务单据提交', '1.录入单据\n2.点击提交', '单据提交成功', '高'],
                ['TC006', '业务单据审批', '1.打开审批界面\n2.选择单据\n3.点击审批', '审批成功', '高'],
            ],
            col_widths=[2, 3, 4, 3, 2]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('8.1.2 性能测试用例', bold=True)
        self._add_table(
            ['用例编号', '测试项目', '测试方法', '测试指标', '预期结果'],
            [
                ['PT001', '响应时间', '并发测试', '页面加载时间', '≤3秒'],
                ['PT002', '并发用户', '并发测试', '并发用户数', '≥100用户'],
                ['PT003', '数据处理', '压力测试', '数据处理量', '≥10000条'],
                ['PT004', '系统稳定性', '稳定性测试', '运行时间', '≥72小时'],
            ],
            col_widths=[2, 3, 3, 3, 3]
        )
        
        self.doc.add_paragraph()
        
        # 8.2 测试数据
        self.doc.add_heading('8.2 测试数据', 2)
        self._add_paragraph('测试所需数据准备如下：', indent=True)
        
        self._add_table(
            ['数据类型', '数据量', '数据说明', '用途'],
            [
                ['基础数据', '1000条', '包含各种状态的基础数据', '功能测试'],
                ['业务单据', '10000条', '包含各种状态的业务单据', '功能测试、性能测试'],
                ['用户数据', '100个', '包含各种角色的用户', '权限测试'],
                ['历史数据', '100000条', '历史业务数据', '性能测试、压力测试'],
            ],
            col_widths=[3, 2, 6, 5]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('数据准备要求：', bold=True)
        self._add_bullet_list([
            '基础数据：覆盖各种业务场景和边界情况',
            '业务单据：包含各种状态的单据（未审核、已审核、已关闭等）',
            '用户数据：包含各种角色的用户，测试权限控制',
            '历史数据：大量历史数据用于性能和压力测试',
        ])
        
        self._add_page_break()
    
    def generate_deployment_plan(self):
        """生成第九章：部署方案"""
        self.doc.add_heading('九、部署方案', 1)
        
        # 9.1 环境要求
        self.doc.add_heading('9.1 环境要求', 2)
        self._add_paragraph('系统部署所需环境要求如下：', indent=True)
        
        self._add_paragraph('9.1.1 硬件要求', bold=True)
        self._add_table(
            ['项目', '最低配置', '推荐配置', '说明'],
            [
                ['CPU', '4核', '8核', '服务器处理器'],
                ['内存', '8GB', '16GB', '服务器内存'],
                ['硬盘', '100GB', '500GB', '服务器硬盘空间'],
                ['网络', '100Mbps', '1000Mbps', '网络带宽'],
            ],
            col_widths=[3, 3, 3, 7]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('9.1.2 软件要求', bold=True)
        self._add_table(
            ['软件类型', '软件名称', '版本要求', '说明'],
            [
                ['操作系统', 'Windows Server', '2016及以上', '服务器操作系统'],
                ['数据库', 'SQL Server', '2016及以上', '数据库服务器'],
                ['应用服务器', 'IIS', '10.0及以上', 'Web应用服务器'],
                ['运行环境', '.NET Framework', '4.7及以上', '.NET运行环境'],
                ['浏览器', 'Chrome/Edge', '最新版本', '客户端浏览器'],
            ],
            col_widths=[3, 3, 3, 7]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('9.1.3 网络要求', bold=True)
        self._add_bullet_list([
            '内网环境：确保内网网络通畅，各服务器之间能够正常通信',
            '外网环境：如需外网访问，需配置防火墙和安全策略',
            '端口开放：开放必要的端口（HTTP:80, HTTPS:443, 数据库端口等）',
            '域名配置：如需域名访问，需配置DNS解析',
        ])
        
        # 9.2 部署步骤
        self.doc.add_heading('9.2 部署步骤', 2)
        self._add_paragraph('系统部署按照以下步骤进行：', indent=True)
        
        self._add_table(
            ['步骤', '工作内容', '责任人', '预计时间', '说明'],
            [
                ['1', '环境准备', '系统管理员', '1天', '准备服务器、安装软件'],
                ['2', '数据库部署', '数据库管理员', '1天', '创建数据库、执行脚本'],
                ['3', '应用部署', '开发工程师', '1天', '部署应用、配置参数'],
                ['4', '功能测试', '测试工程师', '2天', '执行测试用例'],
                ['5', '性能测试', '测试工程师', '1天', '执行性能测试'],
                ['6', '问题修复', '开发工程师', '1天', '修复测试问题'],
                ['7', '用户培训', '业务顾问', '1天', '培训用户操作'],
                ['8', '上线准备', '项目经理', '1天', '准备上线'],
                ['9', '系统上线', '项目经理', '1天', '正式上线'],
            ],
            col_widths=[1, 3, 3, 2, 7]
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('部署注意事项：', bold=True)
        self._add_bullet_list([
            '备份数据：部署前备份现有数据，以防意外',
            '测试环境：先在测试环境验证，再部署到生产环境',
            '监控日志：部署后监控系统日志，及时发现问题',
            '应急预案：准备应急预案，出现问题能够快速回退',
        ])
        
        self._add_page_break()
    
    def generate_appendix(self):
        """生成第十章：附录"""
        self.doc.add_heading('十、附录', 1)
        
        # 10.1 代码示例
        self.doc.add_heading('10.1 代码示例', 2)
        
        self._add_paragraph('10.1.1 基础数据查询接口示例', bold=True)
        self._add_paragraph(
            '// C# 示例代码\n'
            'public class BaseDataController : ApiController\n'
            '{\n'
            '    [HttpPost]\n'
            '    public IHttpActionResult Query(BaseDataQueryRequest request)\n'
            '    {\n'
            '        try\n'
            '        {\n'
            '            // 参数校验\n'
            '            if (request == null)\n'
            '                return BadRequest("请求参数不能为空");\n'
            '            \n'
            '            // 查询数据\n'
            '            var result = _baseDataService.Query(request);\n'
            '            \n'
            '            return Ok(new\n'
            '            {\n'
            '                code = 200,\n'
            '                message = "查询成功",\n'
            '                data = result\n'
            '            });\n'
            '        }\n'
            '        catch (Exception ex)\n'
            '        {\n'
            '            return InternalServerError(ex);\n'
            '        }\n'
            '    }\n'
            '}',
            font_size=9
        )
        
        self.doc.add_paragraph()
        
        self._add_paragraph('10.1.2 数据库脚本示例', bold=True)
        self._add_paragraph(
            '-- SQL 示例代码\n'
            '-- 创建基础数据表\n'
            'CREATE TABLE T_BD_BASEDATA (\n'
            '    FID BIGINT IDENTITY(1,1) PRIMARY KEY,\n'
            '    FNUMBER NVARCHAR(50) NOT NULL,\n'
            '    FNAME NVARCHAR(200) NOT NULL,\n'
            '    FDESCRIPTION NVARCHAR(500),\n'
            '    FSTATUS INT NOT NULL DEFAULT(1),\n'
            '    FCREATEDATE DATETIME NOT NULL DEFAULT(GETDATE()),\n'
            '    FMODIFYDATE DATETIME,\n'
            '    FCREATORID BIGINT NOT NULL,\n'
            '    FMODIFIERID BIGINT\n'
            ');\n'
            '\n'
            '-- 创建索引\n'
            'CREATE UNIQUE INDEX IX_T_BD_BASEDATA_NUMBER \n'
            'ON T_BD_BASEDATA(FNUMBER);\n'
            '\n'
            'CREATE INDEX IX_T_BD_BASEDATA_NAME \n'
            'ON T_BD_BASEDATA(FNAME);',
            font_size=9
        )
        
        self.doc.add_paragraph()
        
        # 10.2 参考文档
        self.doc.add_heading('10.2 参考文档', 2)
        self._add_paragraph('本设计说明书参考以下文档：', indent=True)
        
        self._add_table(
            ['文档编号', '文档名称', '版本', '说明'],
            [
                ['REF001', '金蝶云·星空开发手册', 'V7.5', '金蝶云·星空开发指南'],
                ['REF002', '金蝶云·星空数据库设计规范', 'V7.5', '数据库设计标准'],
                ['REF003', '金蝶云·星空接口开发规范', 'V7.5', '接口开发标准'],
                ['REF004', '金蝶云·星空权限管理手册', 'V7.5', '权限管理指南'],
                ['REF005', '金蝶云·星空部署手册', 'V7.5', '系统部署指南'],
            ],
            col_widths=[2, 5, 2, 7]
        )
        
        self.doc.add_paragraph()
        
        # 文档结束
        self._add_paragraph()
        self._add_paragraph()
        
        end_para = self.doc.add_paragraph()
        end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = end_para.add_run('——文档结束——')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def generate(self, output_path: str = None) -> str:
        """
        生成完整的客户化开发设计说明书
        
        Args:
            output_path: 输出文件路径（可选）
        
        Returns:
            生成的文件路径
        """
        print(f"开始生成客户化开发设计说明书...")
        print(f"公司：{self.company}")
        print(f"项目：{self.project}")
        print(f"模块：{self.module}")
        print(f"需求：{self.requirements}")
        print()
        
        # 生成各章节
        print("生成封面...")
        self.generate_cover()
        
        print("生成文档控制页...")
        self.generate_document_control()
        
        print("生成目录...")
        self.generate_toc()
        
        print("生成第一章：概述...")
        self.generate_overview()
        
        print("生成第二章：需求分析...")
        self.generate_requirements()
        
        print("生成第三章：系统设计...")
        self.generate_system_design()
        
        print("生成第四章：数据库设计...")
        self.generate_database_design()
        
        print("生成第五章：接口设计...")
        self.generate_interface_design()
        
        print("生成第六章：功能设计...")
        self.generate_function_design()
        
        print("生成第七章：安全设计...")
        self.generate_security_design()
        
        print("生成第八章：测试方案...")
        self.generate_test_plan()
        
        print("生成第九章：部署方案...")
        self.generate_deployment_plan()
        
        print("生成第十章：附录...")
        self.generate_appendix()
        
        # 保存文档
        if not output_path:
            timestamp = self.timestamp.strftime('%Y%m%d%H%M%S')
            filename = f"{self.company}_{self.project}_客户化开发设计说明书_{timestamp}.docx"
            output_path = os.path.join(OUTPUT_DIR, filename)
        
        self.doc.save(output_path)
        
        print()
        print(f"✅ 文档生成完成！")
        print(f"📄 文件路径：{output_path}")
        print(f"📊 文档版本：{self.doc_version}")
        print(f"📅 生成时间：{self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        
        return output_path


def main():
    """主函数"""
    # 解析命令行参数
    parser = argparse.ArgumentParser(
        description='金蝶客户化开发设计说明书生成器 v2.0',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  python3 kingdee-dev-spec-v2-generator.py --company "XX公司" --project "ERP项目" --module "财务管理"
  python3 kingdee-dev-spec-v2-generator.py -c "XX公司" -p "ERP项目" -m "供应链" -r "实现采购、销售、库存管理"
  python3 kingdee-dev-spec-v2-generator.py --company "XX公司" --project "ERP项目" --module "生产管理" --output "/path/to/output.docx"
        '''
    )
    
    parser.add_argument('-c', '--company', type=str, required=True,
                       help='公司名称（必填）')
    parser.add_argument('-p', '--project', type=str, required=True,
                       help='项目名称（必填）')
    parser.add_argument('-m', '--module', type=str, required=True,
                       help='模块名称（必填）')
    parser.add_argument('-r', '--requirements', type=str, default='',
                       help='需求描述（可选）')
    parser.add_argument('-o', '--output', type=str, default='',
                       help='输出文件路径（可选）')
    parser.add_argument('-v', '--version', action='version',
                       version=f'金蝶客户化开发设计说明书生成器 v{__version__}')
    
    args = parser.parse_args()
    
    # 创建生成器
    generator = KingdeeDevSpecGenerator(
        company=args.company,
        project=args.project,
        module=args.module,
        requirements=args.requirements
    )
    
    # 生成文档
    output_path = args.output if args.output else None
    try:
        file_path = generator.generate(output_path)
        print(f"\n文档已生成：{file_path}")
        return 0
    except Exception as e:
        print(f"\n❌ 文档生成失败：{str(e)}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())