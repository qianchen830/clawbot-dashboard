#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶系统集成方案生成器 v2.0 增强版
- 详细接口设计（入参/出参/映射）
- 集成架构图描述
- 数据映射规则
- 异常处理机制
- 接口监控方案
"""

import os, sys, json, argparse
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

INTEGRATION_KNOWLEDGE = {
    '制造业': {
        'interfaces': [
            {'system': 'MES系统', 'direction': '双向', 'type': '实时',
             'desc': '生产工单下发、完工汇报接收、质量数据同步',
             'data_objects': [
                 {'name': '生产工单', 'from': '金蝶', 'to': 'MES', 'fields': '工单号、物料、数量、计划日期', 'frequency': '实时'},
                 {'name': '完工汇报', 'from': 'MES', 'to': '金蝶', 'fields': '工单号、工序、完工数量、工时', 'frequency': '实时'},
                 {'name': '质量数据', 'from': 'MES', 'to': '金蝶', 'fields': '检验单号、检验结果、不良原因', 'frequency': '实时'}
             ]},
            {'system': 'WMS系统', 'direction': '双向', 'type': '实时',
             'desc': '入库通知、出库通知、库存同步',
             'data_objects': [
                 {'name': '入库通知', 'from': '金蝶', 'to': 'WMS', 'fields': '单据号、物料、数量、仓库', 'frequency': '实时'},
                 {'name': '出库通知', 'from': '金蝶', 'to': 'WMS', 'fields': '单据号、物料、数量、库位', 'frequency': '实时'},
                 {'name': '库存同步', 'from': 'WMS', 'to': '金蝶', 'fields': '物料、仓库、可用量、冻结量', 'frequency': '每15分钟'}
             ]},
            {'system': 'SRM系统', 'direction': '双向', 'type': '定时',
             'desc': '供应商数据、采购订单、收货确认同步',
             'data_objects': [
                 {'name': '采购订单', 'from': '金蝶', 'to': 'SRM', 'fields': '订单号、物料、数量、交期', 'frequency': '实时'},
                 {'name': '收货确认', 'from': 'SRM', 'to': '金蝶', 'fields': '订单号、收货数量、收货日期', 'frequency': '实时'}
             ]},
            {'system': 'OA系统', 'direction': '单向', 'type': '实时',
             'desc': '审批流程推送、消息通知',
             'data_objects': [
                 {'name': '审批通知', 'from': '金蝶', 'to': 'OA', 'fields': '单据类型、单据号、金额、申请人', 'frequency': '实时'}
             ]},
            {'system': '银行系统', 'direction': '双向', 'type': '定时',
             'desc': '付款指令、银行回单、余额查询',
             'data_objects': [
                 {'name': '付款指令', 'from': '金蝶', 'to': '银行', 'fields': '付款单号、收款方、金额、银行账号', 'frequency': '定时'},
                 {'name': '银行回单', 'from': '银行', 'to': '金蝶', 'fields': '回单号、金额、交易日期、对方账户', 'frequency': '每日'}
             ]}
        ]
    },
    '零售业': {
        'interfaces': [
            {'system': '电商平台', 'direction': '双向', 'type': '实时',
             'desc': '订单同步、库存同步、会员同步',
             'data_objects': [
                 {'name': '销售订单', 'from': '电商', 'to': '金蝶', 'fields': '订单号、商品、数量、金额、收货地址', 'frequency': '实时'},
                 {'name': '库存同步', 'from': '金蝶', 'to': '电商', 'fields': '商品编码、可用库存', 'frequency': '每30分钟'}
             ]},
            {'system': 'POS系统', 'direction': '双向', 'type': '实时',
             'desc': '销售数据、会员数据同步',
             'data_objects': [
                 {'name': '销售数据', 'from': 'POS', 'to': '金蝶', 'fields': '门店、交易号、商品、数量、金额', 'frequency': '实时'}
             ]},
            {'system': 'CRM系统', 'direction': '双向', 'type': '定时',
             'desc': '会员数据、营销数据同步',
             'data_objects': [
                 {'name': '会员数据', 'from': 'CRM', 'to': '金蝶', 'fields': '会员号、姓名、积分、等级', 'frequency': '每日'}
             ]}
        ]
    }
}

DEFAULT_KNOWLEDGE = {
    'interfaces': [
        {'system': 'OA系统', 'direction': '单向', 'type': '实时',
         'desc': '审批流程、消息通知',
         'data_objects': [
             {'name': '审批通知', 'from': '金蝶', 'to': 'OA', 'fields': '单据类型、单据号、金额', 'frequency': '实时'}
         ]},
        {'system': '银行系统', 'direction': '双向', 'type': '定时',
         'desc': '付款指令、银行回单',
         'data_objects': [
             {'name': '付款指令', 'from': '金蝶', 'to': '银行', 'fields': '付款单号、金额、银行账号', 'frequency': '定时'},
             {'name': '银行回单', 'from': '银行', 'to': '金蝶', 'fields': '回单号、金额、交易日期', 'frequency': '每日'}
         ]}
    ]
}

def get_knowledge(industry):
    for k in INTEGRATION_KNOWLEDGE:
        if k in industry or industry in k:
            return INTEGRATION_KNOWLEDGE[k]
    return DEFAULT_KNOWLEDGE

def set_cell_shading(cell, color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def styled_table(doc, headers, rows, color='006699'):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(c, color)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
            if i % 2 == 1:
                set_cell_shading(t.rows[i+1].cells[j], 'F5F5F5')
    return t

def generate_integration_v2(customer_info):
    doc = Document()
    company = customer_info.get('companyName', '企业名称')
    industry = customer_info.get('industry', '制造业')
    knowledge = get_knowledge(industry)

    # 封面
    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run(company)
    r.font.size = Pt(28); r.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    r2 = p2.add_run('系统集成方案')
    r2.font.size = Pt(22); r2.font.bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info = doc.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ('企业名称', company), ('所属行业', industry),
        ('版本', 'V1.0'),
        ('编制日期', datetime.now().strftime('%Y年%m月%d日')),
        ('编制单位', '金蝶软件（中国）有限公司')
    ]):
        info.rows[i].cells[0].text = k
        info.rows[i].cells[1].text = v

    doc.add_page_break()

    # 目录
    doc.add_heading('目  录', level=1)
    for item in ['一、集成概述', '二、集成架构', '三、接口设计', '四、数据映射', '五、异常处理', '六、接口监控', '七、实施计划', '附录：接口清单']:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 一、集成概述
    doc.add_heading('一、集成概述', level=1)
    doc.add_heading('1.1 集成目标', level=2)
    for item in ['打通信息孤岛，实现业务系统间数据共享', '消除手工重复录入，提升工作效率', '确保数据一致性，提升数据质量', '实现业务协同，支撑端到端业务流程']:
        doc.add_paragraph(f'• {item}')

    doc.add_heading('1.2 集成原则', level=2)
    for item in ['标准化：统一接口规范，使用RESTful API', '可靠性：确保数据传输的完整性和准确性', '实时性：关键业务数据实时同步', '安全性：接口认证加密，数据传输加密', '可维护性：接口日志完整，问题可追溯']:
        doc.add_paragraph(f'• {item}')

    doc.add_heading('1.3 集成范围', level=2)
    total = sum(len(i['data_objects']) for i in knowledge['interfaces'])
    doc.add_paragraph(f'本次集成涉及 {len(knowledge["interfaces"])} 个外部系统，{total} 个数据接口。')
    rows = [[i['system'], i['direction'], i['type'], i['desc'], len(i['data_objects'])] for i in knowledge['interfaces']]
    styled_table(doc, ['系统名称', '数据方向', '集成类型', '说明', '接口数'], rows)
    doc.add_paragraph('')

    # 二、集成架构
    doc.add_heading('二、集成架构', level=1)
    doc.add_heading('2.1 总体架构', level=2)
    doc.add_paragraph(f'{company}系统集成采用企业服务总线（ESB）模式，金蝶云·星空作为核心业务平台，通过标准API接口与各外部系统进行数据交互。')
    doc.add_paragraph('')
    doc.add_paragraph('集成架构层次：')
    layers = [
        ('应用层', '金蝶云·星空、MES、WMS、OA、银行等业务系统'),
        ('集成层', 'API网关、消息队列、数据转换、路由分发'),
        ('数据层', '主数据管理、数据映射、数据同步、数据质量'),
        ('监控层', '接口监控、日志管理、异常告警、性能分析')
    ]
    styled_table(doc, ['架构层次', '说明'], layers, '336699')
    doc.add_paragraph('')

    doc.add_heading('2.2 技术方案', level=2)
    tech = [
        ('接口协议', 'RESTful API / WebService'),
        ('数据格式', 'JSON / XML'),
        ('认证方式', 'OAuth2.0 / API Key'),
        ('传输加密', 'HTTPS / SSL/TLS'),
        ('消息队列', 'RabbitMQ / Kafka（异步场景）'),
        ('日志记录', 'ELK Stack'),
        ('监控告警', 'Prometheus + Grafana')
    ]
    styled_table(doc, ['技术项', '方案'], tech, '336699')
    doc.add_paragraph('')

    # 三、接口设计
    doc.add_heading('三、接口设计', level=1)
    for idx, interface in enumerate(knowledge['interfaces'], 1):
        doc.add_heading(f'3.{idx} {interface["system"]}集成', level=2)
        doc.add_paragraph(f'集成方向：{interface["direction"]} | 集成类型：{interface["type"]}')
        doc.add_paragraph(f'说明：{interface["desc"]}')
        doc.add_paragraph('')

        doc.add_heading(f'3.{idx}.1 接口清单', level=3)
        rows = [[d['name'], d['from'], d['to'], d['fields'], d['frequency']] for d in interface['data_objects']]
        styled_table(doc, ['数据对象', '数据来源', '数据目标', '关键字段', '同步频率'], rows, '336699')
        doc.add_paragraph('')

        for di, dobj in enumerate(interface['data_objects'], 1):
            doc.add_heading(f'3.{idx}.{di+1} {dobj["name"]}接口', level=3)
            doc.add_paragraph(f'• 数据方向：{dobj["from"]} → {dobj["to"]}')
            doc.add_paragraph(f'• 关键字段：{dobj["fields"]}')
            doc.add_paragraph(f'• 同步频率：{dobj["frequency"]}')
            doc.add_paragraph(f'• 接口地址：/api/integration/{dobj["name"].lower().replace(" ", "-")}')
            doc.add_paragraph(f'• 请求方式：POST')
            doc.add_paragraph(f'• 认证方式：Bearer Token')
            doc.add_paragraph('')

    # 四、数据映射
    doc.add_heading('四、数据映射', level=1)
    doc.add_heading('4.1 主数据映射规则', level=2)
    mappings = [
        ('物料编码', '金蝶物料编码', '外部系统物料编码', '编码映射表'),
        ('客户编码', '金蝶客户编码', '外部系统客户编码', '编码映射表'),
        ('供应商编码', '金蝶供应商编码', '外部系统供应商编码', '编码映射表'),
        ('仓库编码', '金蝶仓库编码', 'WMS仓库编码', '仓库映射表'),
        ('组织编码', '金蝶组织编码', '外部系统组织编码', '组织映射表'),
        ('币种编码', '金蝶币种编码', '银行币种编码', '币种映射表')
    ]
    styled_table(doc, ['数据类型', '金蝶编码', '外部编码', '映射方式'], mappings, '336633')
    doc.add_paragraph('')

    doc.add_heading('4.2 数据质量规则', level=2)
    rules = [
        ('完整性', '必填字段不能为空', '接口层校验'),
        ('准确性', '数据值在有效范围内', '业务规则校验'),
        ('一致性', '同一数据在多个系统中保持一致', '主数据管理'),
        ('及时性', '数据在规定时间内同步完成', '监控告警'),
        ('唯一性', '编码全局唯一', '编码规则校验')
    ]
    styled_table(doc, ['规则类型', '说明', '保障措施'], rules, '666633')
    doc.add_paragraph('')

    # 五、异常处理
    doc.add_heading('五、异常处理', level=1)
    doc.add_heading('5.1 异常类型', level=2)
    errors = [
        ('网络异常', '网络中断或超时', '自动重试3次，间隔递增', '告警通知运维'),
        ('数据异常', '数据格式错误或校验失败', '记录错误日志，人工处理', '告警通知业务'),
        ('系统异常', '目标系统不可用', '消息队列暂存，系统恢复后重发', '告警通知运维'),
        ('业务异常', '业务规则冲突', '记录冲突日志，人工裁决', '告警通知业务'),
        ('性能异常', '响应超时或吞吐量不足', '限流降级，异步处理', '告警通知运维')
    ]
    styled_table(doc, ['异常类型', '描述', '处理方式', '通知方式'], errors, '993333')
    doc.add_paragraph('')

    doc.add_heading('5.2 异常处理流程', level=2)
    doc.add_paragraph('1. 接口调用 → 2. 结果校验 → 3. 成功：更新状态 → 4. 失败：记录日志 → 5. 自动重试 → 6. 超过阈值：人工处理')

    # 六、接口监控
    doc.add_heading('六、接口监控', level=1)
    doc.add_heading('6.1 监控指标', level=2)
    metrics = [
        ('接口成功率', '≥99.5%', '每5分钟'),
        ('接口响应时间', '≤3秒（实时）/ ≤60秒（批量）', '每5分钟'),
        ('数据同步延迟', '≤5分钟（实时）/ ≤1小时（批量）', '每15分钟'),
        ('接口调用量', '按小时/天统计', '每小时'),
        ('错误数量', '按错误类型统计', '每5分钟')
    ]
    styled_table(doc, ['监控指标', '阈值', '采集频率'], metrics, '336699')
    doc.add_paragraph('')

    doc.add_heading('6.2 告警机制', level=2)
    alerts = [
        ('严重告警', '接口成功率<95%或完全不可用', '短信+邮件+电话', '5分钟内'),
        ('重要告警', '接口成功率<99%或响应超时', '短信+邮件', '15分钟内'),
        ('一般告警', '数据延迟或调用量异常', '邮件', '1小时内')
    ]
    styled_table(doc, ['告警级别', '触发条件', '通知方式', '响应时间'], alerts, '993333')
    doc.add_paragraph('')

    # 七、实施计划
    doc.add_heading('七、实施计划', level=1)
    plan = [
        ('第1阶段', '接口调研', '2周', '梳理接口需求，确定接口规格'),
        ('第2阶段', '接口开发', '3周', '开发接口程序，单元测试'),
        ('第3阶段', '联调测试', '2周', '与外部系统联调，集成测试'),
        ('第4阶段', '性能测试', '1周', '性能压测，优化调整'),
        ('第5阶段', '上线部署', '1周', '生产环境部署，灰度发布')
    ]
    styled_table(doc, ['阶段', '名称', '周期', '主要内容'], plan, '006699')
    doc.add_paragraph('')

    # 附录
    doc.add_page_break()
    doc.add_heading('附录：接口清单汇总', level=1)
    all_interfaces = []
    for iface in knowledge['interfaces']:
        for dobj in iface['data_objects']:
            all_interfaces.append([
                f'{dobj["from"]}→{dobj["to"]}',
                iface['system'], dobj['name'], iface['type'], dobj['frequency']
            ])
    styled_table(doc, ['数据方向', '系统', '数据对象', '类型', '频率'], all_interfaces, '666666')

    # 保存
    filename = f"{customer_info.get('customerCode', company)}_系统集成方案_v2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    return {
        'success': True,
        'filepath': filepath,
        'filename': filename,
        'version': '2.0',
        'interfaces': len(knowledge['interfaces']),
        'data_objects': sum(len(i['data_objects']) for i in knowledge['interfaces'])
    }

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--companyName', required=True)
    parser.add_argument('--companySize', default='中型企业', help='企业规模')
    parser.add_argument('--customerCode', default='')
    parser.add_argument('--industry', default='制造业')
    parser.add_argument('--employees', default='', help='员工人数')
    parser.add_argument('--revenue', default='', help='年营业额')
    parser.add_argument('--modules', default='', help='实施模块')
    args = parser.parse_args()

    customer_info = {
        'companyName': args.companyName,
        'customerCode': args.customerCode or args.companyName,
        'industry': args.industry
    }

    result = generate_integration_v2(customer_info)
    print(json.dumps(result, ensure_ascii=False, indent=2))
