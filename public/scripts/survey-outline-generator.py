#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调研提纲生成器 v1.0
基于中煤科工项目调研提纲模板
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import sys

class SurveyOutlineGenerator:
    """调研提纲生成器"""
    
    def __init__(self, customer_name, project_name, module, survey_date=None):
        self.customer_name = customer_name
        self.project_name = project_name
        self.module = module
        self.survey_date = survey_date or datetime.now().strftime("%Y%m%d")
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(11)
    
    def _add_title(self, text, level=1):
        """添加标题"""
        if level == 1:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        elif level == 3:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _add_paragraph(self, text):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
    
    def _add_question(self, question):
        """添加问题"""
        p = self.doc.add_paragraph()
        run = p.add_run(question)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.5)
    
    def generate(self):
        """生成调研提纲"""
        # 封面
        self._add_title(f"{self.customer_name}新ERP管理系统项目")
        self._add_title(f"{self.module}调研提纲")
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 文档信息
        self._add_paragraph(f"{self.customer_name}")
        self._add_paragraph("金蝶软件（中国）有限公司")
        self._add_paragraph(datetime.now().strftime("%Y年%m月"))
        self.doc.add_page_break()
        
        # 文档控制
        self._add_title("文档控制", 2)
        
        # 更改记录表
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        headers = ['日期', '作者', '版本', '更改参考']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        table.cell(1, 0).text = datetime.now().strftime("%Y-%m-%d")
        table.cell(1, 1).text = ""
        table.cell(1, 2).text = "V1.0"
        table.cell(1, 3).text = ""
        
        self.doc.add_paragraph()
        
        # 调研提纲内容
        self._add_title(f"{self.module}调研提纲", 1)
        
        # 根据模块生成不同的问题
        questions = self._get_module_questions()
        
        for i, (section, qs) in enumerate(questions.items(), 1):
            self._add_title(f"{i}. {section}", 2)
            for q in qs:
                self._add_question(q)
            self.doc.add_paragraph()
        
        # 保存文档
        filename = f"{self.customer_name}_{self.module}_调研提纲_{self.survey_date}_V1.0.docx"
        self.doc.save(filename)
        return filename
    
    def _get_module_questions(self):
        """根据模块获取问题列表"""
        
        # 财务模块问题
        finance_questions = {
            "业务现状": [
                "请描述贵公司当前的财务核算流程？",
                "贵公司目前使用的财务系统是什么？",
                "财务核算的组织架构是怎样的？",
                "目前财务核算存在哪些主要问题？"
            ],
            "业务流程": [
                "日常凭证处理的流程是怎样的？",
                "月末结账的流程和时间节点？",
                "财务报表编制的流程？",
                "与业务系统的数据传递方式？"
            ],
            "主要问题": [
                "目前系统存在哪些功能不足？",
                "业务流程中存在哪些痛点？",
                "数据准确性方面存在哪些问题？",
                "工作效率方面存在哪些瓶颈？"
            ],
            "需求分析": [
                "对新系统的核心需求是什么？",
                "希望解决哪些业务痛点？",
                "对系统功能有哪些期望？",
                "对实施周期有何要求？"
            ],
            "基础数据": [
                "会计科目体系是怎样的？",
                "辅助核算项目有哪些？",
                "凭证类型有哪些？",
                "核算组织有哪些？"
            ],
            "外围系统": [
                "需要与哪些外围系统集成？",
                "数据交换的频率和方式？",
                "接口格式要求是什么？"
            ]
        }
        
        # 供应链模块问题
        supply_chain_questions = {
            "业务现状": [
                "请描述贵公司当前的采购业务流程？",
                "库存管理的现状如何？",
                "供应商管理的现状如何？",
                "采购计划是如何制定的？"
            ],
            "业务流程": [
                "采购申请到入库的完整流程？",
                "供应商准入和评估流程？",
                "采购退货的处理流程？",
                "库存盘点的方式和频率？"
            ],
            "主要问题": [
                "采购周期是否过长？",
                "库存周转率是否满意？",
                "供应商管理是否规范？",
                "数据准确性如何？"
            ],
            "需求分析": [
                "对采购管理的期望？",
                "对库存管理的期望？",
                "对供应商管理的期望？",
                "希望实现哪些自动化？"
            ],
            "基础数据": [
                "物料分类体系是怎样的？",
                "仓库设置情况？",
                "供应商分类标准？",
                "采购价格管理方式？"
            ],
            "外围系统": [
                "是否需要与WMS系统集成？",
                "是否需要与MES系统集成？",
                "是否需要与财务系统集成？"
            ]
        }
        
        # 制造模块问题
        manufacturing_questions = {
            "业务现状": [
                "请描述贵公司当前的生产计划流程？",
                "生产执行的现状如何？",
                "质量管理的现状如何？",
                "设备管理的现状如何？"
            ],
            "业务流程": [
                "生产计划制定的流程？",
                "生产订单执行流程？",
                "质量检验流程？",
                "设备维护流程？"
            ],
            "主要问题": [
                "生产计划是否准确？",
                "生产进度是否可控？",
                "质量追溯是否完整？",
                "设备效率如何？"
            ],
            "需求分析": [
                "对生产管理的期望？",
                "对质量管理的期望？",
                "对设备管理的期望？",
                "希望实现哪些智能化？"
            ],
            "基础数据": [
                "BOM管理方式？",
                "工艺路线设置？",
                "工作中心设置？",
                "质量标准设置？"
            ],
            "外围系统": [
                "是否需要与MES集成？",
                "是否需要与WMS集成？",
                "是否需要与PLM集成？"
            ]
        }
        
        # 人力模块问题
        hr_questions = {
            "业务现状": [
                "请描述贵公司当前的人力资源管理流程？",
                "薪酬核算的流程如何？",
                "绩效管理的流程如何？",
                "招聘管理的流程如何？"
            ],
            "业务流程": [
                "员工入职流程？",
                "薪酬核算流程？",
                "绩效考核流程？",
                "离职办理流程？"
            ],
            "主要问题": [
                "薪酬核算是否准确及时？",
                "绩效管理是否有效？",
                "招聘流程是否高效？",
                "数据准确性如何？"
            ],
            "需求分析": [
                "对薪酬管理的期望？",
                "对绩效管理的期望？",
                "对招聘管理的期望？",
                "对自助服务的期望？"
            ],
            "基础数据": [
                "组织架构设置？",
                "职位体系设置？",
                "薪酬体系设置？",
                "绩效指标设置？"
            ],
            "外围系统": [
                "是否需要与财务系统集成？",
                "是否需要与OA系统集成？",
                "是否需要与考勤系统集成？"
            ]
        }
        
        # 模块映射
        module_map = {
            "财务管理": finance_questions,
            "全面预算": finance_questions,
            "费用管理": finance_questions,
            "应收管理": finance_questions,
            "应付管理": finance_questions,
            "出纳": finance_questions,
            "总账": finance_questions,
            "成本管理": finance_questions,
            "税务管理": finance_questions,
            "供应链": supply_chain_questions,
            "采购管理": supply_chain_questions,
            "库存管理": supply_chain_questions,
            "供应商管理": supply_chain_questions,
            "生产制造": manufacturing_questions,
            "质量管理": manufacturing_questions,
            "设备管理": manufacturing_questions,
            "人力资源": hr_questions,
            "薪酬管理": hr_questions,
            "绩效管理": hr_questions,
            "招聘管理": hr_questions,
        }
        
        return module_map.get(self.module, finance_questions)


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='调研提纲生成器 v1.0')
    parser.add_argument('--customer', required=True, help='客户名称')
    parser.add_argument('--module', required=True, 
                       choices=['财务管理', '全面预算', '费用管理', '应收管理', '应付管理', 
                               '出纳', '总账', '成本管理', '税务管理',
                               '供应链', '采购管理', '库存管理', '供应商管理',
                               '生产制造', '质量管理', '设备管理',
                               '人力资源', '薪酬管理', '绩效管理', '招聘管理'],
                       help='调研模块')
    parser.add_argument('--project', default='新ERP管理系统项目', help='项目名称')
    parser.add_argument('--date', help='调研日期 (YYYYMMDD)')
    
    args = parser.parse_args()
    
    generator = SurveyOutlineGenerator(
        customer_name=args.customer,
        project_name=args.project,
        module=args.module,
        survey_date=args.date
    )
    
    filename = generator.generate()
    print(f"调研提纲已生成: {filename}")


if __name__ == "__main__":
    main()
