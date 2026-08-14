#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶系统集成方案生成器 v3.0 增强版
- 规范Word文档格式（页眉页脚、封面页、目录页、正文格式）
- 扩展接口设计（15+标准接口：财务、供应链、OA、CRM、MES、WMS）
- 完整接口设计格式（请求参数、响应参数、错误码）
- 数据映射表格（源系统、目标系统、字段映射、转换规则）
- 集成架构图描述
- 接口测试用例表格
- 常见问题FAQ
- 命令行参数兼容：--companyName, --industry, --modules
"""

import os, sys, json, argparse
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ===================== 15 标准接口设计 =====================
INTERFACES = {
    '财务系统': [
        {
            'name': '凭证同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'OA → 金蝶', 'frequency': '实时',
            'description': '从OA系统审批通过的凭证数据同步到金蝶生成财务凭证',
            'endpoint': '/api/v1/finance/voucher/sync',
            'request_params': [
                ('voucher_no', 'string', '是', '凭证号', 'OA系统凭证编号'),
                ('voucher_date', 'date', '是', '凭证日期', 'YYYY-MM-DD'),
                ('debit_amount', 'decimal', '是', '借方金额', '精确到分'),
                ('credit_amount', 'decimal', '是', '贷方金额', '精确到分'),
                ('account_code', 'string', '是', '科目编码', '金蝶科目编码'),
                ('summary', 'string', '否', '摘要', '凭证摘要'),
                ('currency', 'string', '否', '币种', '默认CNY'),
                ('org_code', 'string', '是', '组织编码', '核算组织编码'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('message', 'string', '响应消息', '成功或错误描述'),
                ('data.kingdee_voucher_no', 'string', '金蝶凭证号', '金蝶生成'),
                ('data.create_time', 'datetime', '创建时间', '金蝶创建时间'),
            ],
            'error_codes': [
                ('40001', '参数校验失败', '检查请求参数格式和必填项'),
                ('40002', '科目编码不存在', '确认金蝶科目编码是否正确'),
                ('40003', '借贷不平衡', '检查借贷金额是否相等'),
                ('40004', '凭证日期不在期间', '确认会计期间是否开放'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '付款申请接口',
            'type': '业务触发', 'method': 'RESTful API',
            'direction': 'OA → 金蝶', 'frequency': '实时',
            'description': '从OA审批通过的付款申请同步到金蝶生成付款单',
            'endpoint': '/api/v1/finance/payment/apply',
            'request_params': [
                ('apply_no', 'string', '是', '申请单号', 'OA付款申请编号'),
                ('apply_date', 'date', '是', '申请日期', 'YYYY-MM-DD'),
                ('amount', 'decimal', '是', '付款金额', '精确到分'),
                ('payee_name', 'string', '是', '收款方名称', '单位或个人'),
                ('payee_account', 'string', '是', '收款账号', '银行账号'),
                ('payee_bank', 'string', '是', '收款银行', '开户银行'),
                ('purpose', 'string', '否', '付款用途', '付款说明'),
                ('cost_center', 'string', '否', '成本中心', '成本中心编码'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('message', 'string', '响应消息', '成功或错误描述'),
                ('data.payment_no', 'string', '付款单号', '金蝶付款单编号'),
                ('data.status', 'string', '状态', '待审批/已审批'),
            ],
            'error_codes': [
                ('40101', '收款方不存在', '确认供应商或客户档案'),
                ('40102', '金额超限', '检查预算控制'),
                ('40103', '账户信息错误', '确认银行账户信息'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '银行回单同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': '银行 → 金蝶', 'frequency': '每日',
            'description': '从银行系统获取银行回单数据同步到金蝶',
            'endpoint': '/api/v1/finance/bank/receipt',
            'request_params': [
                ('bank_code', 'string', '是', '银行编码', '银行编号'),
                ('account_no', 'string', '是', '银行账号', '查询账号'),
                ('start_date', 'date', '是', '开始日期', 'YYYY-MM-DD'),
                ('end_date', 'date', '是', '结束日期', 'YYYY-MM-DD'),
                ('page_no', 'integer', '否', '页码', '默认1'),
                ('page_size', 'integer', '否', '每页条数', '默认100'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.total', 'integer', '总记录数', '符合条件的总条数'),
                ('data.list', 'array', '回单列表', '银行回单明细数组'),
                ('data.list[].transaction_no', 'string', '流水号', '银行交易流水'),
                ('data.list[].amount', 'decimal', '交易金额', '正收负支'),
            ],
            'error_codes': [
                ('40201', '银行编码无效', '确认银行编码'),
                ('40202', '账号不存在', '确认银行账户档案'),
                ('40203', '日期范围超限', '最多查询90天'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
    ],
    '供应链系统': [
        {
            'name': '采购订单同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'SRM → 金蝶', 'frequency': '实时',
            'description': '从SRM系统同步采购订单到金蝶ERP',
            'endpoint': '/api/v1/scm/purchase/order/sync',
            'request_params': [
                ('order_no', 'string', '是', '采购订单号', 'SRM订单编号'),
                ('order_date', 'date', '是', '订单日期', 'YYYY-MM-DD'),
                ('supplier_code', 'string', '是', '供应商编码', '金蝶供应商编码'),
                ('material_code', 'string', '是', '物料编码', '金蝶物料编码'),
                ('qty', 'decimal', '是', '采购数量', '采购数量'),
                ('unit_price', 'decimal', '是', '单价', '含税单价'),
                ('delivery_date', 'date', '是', '交货日期', '期望交货日期'),
                ('warehouse', 'string', '否', '收货仓库', '默认仓库编码'),
                ('purchase_org', 'string', '否', '采购组织', '采购组织编码'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.kingdee_order_no', 'string', '金蝶订单号', '金蝶采购订单号'),
                ('data.status', 'string', '状态', '已创建/已审核'),
            ],
            'error_codes': [
                ('40301', '供应商不存在', '确认供应商档案'),
                ('40302', '物料不存在', '确认物料档案'),
                ('40303', '价格不一致', '检查采购价目表'),
                ('40304', '库存不足', '确认仓库库存'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '销售订单同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'CRM → 金蝶', 'frequency': '实时',
            'description': '从CRM系统同步销售订单到金蝶ERP',
            'endpoint': '/api/v1/scm/sales/order/sync',
            'request_params': [
                ('order_no', 'string', '是', '销售订单号', 'CRM订单编号'),
                ('order_date', 'date', '是', '订单日期', 'YYYY-MM-DD'),
                ('customer_code', 'string', '是', '客户编码', '金蝶客户编码'),
                ('material_code', 'string', '是', '物料编码', '金蝶物料编码'),
                ('qty', 'decimal', '是', '销售数量', '销售数量'),
                ('unit_price', 'decimal', '是', '单价', '含税单价'),
                ('delivery_date', 'date', '是', '交货日期', '承诺交货日期'),
                ('sales_org', 'string', '否', '销售组织', '销售组织编码'),
                ('shipping_address', 'string', '否', '收货地址', '客户收货地址'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.kingdee_order_no', 'string', '金蝶订单号', '金蝶销售订单号'),
                ('data.available_qty', 'decimal', '可用库存', '当前可用库存量'),
            ],
            'error_codes': [
                ('40401', '客户不存在', '确认客户档案'),
                ('40402', '物料不存在', '确认物料档案'),
                ('40403', '库存不足', '检查可用库存'),
                ('40404', '价格未维护', '维护销售价目表'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '库存同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'WMS → 金蝶', 'frequency': '每15分钟',
            'description': '从WMS系统同步库存数据到金蝶ERP',
            'endpoint': '/api/v1/scm/inventory/sync',
            'request_params': [
                ('warehouse_code', 'string', '是', '仓库编码', 'WMS仓库编码'),
                ('material_code', 'string', '是', '物料编码', '物料编码'),
                ('batch_no', 'string', '否', '批次号', '批次管理物料必填'),
                ('qty_available', 'decimal', '是', '可用数量', '可用库存量'),
                ('qty_frozen', 'decimal', '否', '冻结数量', '冻结库存量'),
                ('location', 'string', '否', '库位', '库位编码'),
                ('sync_time', 'datetime', '是', '同步时间', '库存快照时间'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.sync_id', 'string', '同步批次号', '本次同步批次号'),
                ('data.diff_qty', 'decimal', '差异量', '与上次同步的差异'),
            ],
            'error_codes': [
                ('40501', '仓库不存在', '确认仓库档案'),
                ('40502', '物料不存在', '确认物料档案'),
                ('40503', '数量为负', '检查库存数据'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
    ],
    'OA系统': [
        {
            'name': '审批流程触发接口',
            'type': '业务触发', 'method': 'RESTful API',
            'direction': '金蝶 → OA', 'frequency': '实时',
            'description': '金蝶单据提交审批时触发OA审批流程',
            'endpoint': '/api/v1/oa/approval/trigger',
            'request_params': [
                ('bill_type', 'string', '是', '单据类型', '付款申请/费用报销等'),
                ('bill_no', 'string', '是', '单据编号', '金蝶单据编号'),
                ('applicant', 'string', '是', '申请人', '申请人工号'),
                ('amount', 'decimal', '否', '金额', '涉及金额'),
                ('dept_code', 'string', '是', '部门编码', '申请人部门'),
                ('summary', 'string', '是', '摘要', '审批事项摘要'),
                ('callback_url', 'string', '是', '回调地址', '审批结果回调URL'),
                ('attach_url', 'string', '否', '附件地址', '附件下载地址'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.process_id', 'string', '流程实例ID', 'OA流程编号'),
                ('data.current_node', 'string', '当前节点', '当前审批节点'),
            ],
            'error_codes': [
                ('40601', '流程模板不存在', '确认流程配置'),
                ('40602', '申请人不存在', '确认员工档案'),
                ('40603', '部门不存在', '确认组织架构'),
                ('40604', '回调地址无效', '检查URL格式'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '审批结果回调接口',
            'type': '业务触发', 'method': 'RESTful API',
            'direction': 'OA → 金蝶', 'frequency': '实时',
            'description': 'OA审批完成后回调金蝶更新单据状态',
            'endpoint': '/api/v1/oa/approval/callback',
            'request_params': [
                ('process_id', 'string', '是', '流程实例ID', 'OA流程编号'),
                ('bill_no', 'string', '是', '单据编号', '金蝶单据编号'),
                ('approve_result', 'string', '是', '审批结果', '通过/驳回/撤销'),
                ('approver', 'string', '是', '审批人', '最终审批人工号'),
                ('approve_time', 'datetime', '是', '审批时间', 'YYYY-MM-DD HH:mm:ss'),
                ('approve_opinion', 'string', '否', '审批意见', '审批意见内容'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.bill_status', 'string', '单据状态', '更新后单据状态'),
            ],
            'error_codes': [
                ('40701', '单据不存在', '确认单据编号'),
                ('40702', '单据状态错误', '确认单据当前状态'),
                ('40703', '审批人无权限', '确认审批权限'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
    ],
    'CRM系统': [
        {
            'name': '客户档案同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'CRM → 金蝶', 'frequency': '每日',
            'description': '从CRM系统同步客户档案到金蝶ERP',
            'endpoint': '/api/v1/crm/customer/sync',
            'request_params': [
                ('customer_code', 'string', '是', '客户编码', 'CRM客户编号'),
                ('customer_name', 'string', '是', '客户名称', '客户全称'),
                ('customer_type', 'string', '是', '客户类型', '企业/个人'),
                ('credit_code', 'string', '否', '统一社会信用代码', '企业客户必填'),
                ('contact_person', 'string', '否', '联系人', '主要联系人'),
                ('contact_phone', 'string', '否', '联系电话', '联系人电话'),
                ('address', 'string', '否', '地址', '客户地址'),
                ('credit_limit', 'decimal', '否', '信用额度', '授信额度'),
                ('sales_rep', 'string', '否', '销售员', '负责销售员工号'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.kingdee_code', 'string', '金蝶客户编码', '金蝶客户编号'),
                ('data.create_time', 'datetime', '创建时间', '档案创建时间'),
            ],
            'error_codes': [
                ('40801', '客户已存在', '检查客户编码是否重复'),
                ('40802', '客户名称重复', '确认客户名称唯一性'),
                ('40803', '销售员不存在', '确认员工档案'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '商机转化接口',
            'type': '业务触发', 'method': 'RESTful API',
            'direction': 'CRM → 金蝶', 'frequency': '实时',
            'description': 'CRM商机赢单后同步到金蝶生成销售订单',
            'endpoint': '/api/v1/crm/opportunity/convert',
            'request_params': [
                ('opportunity_id', 'string', '是', '商机ID', 'CRM商机编号'),
                ('customer_code', 'string', '是', '客户编码', '金蝶客户编码'),
                ('expected_revenue', 'decimal', '是', '预期收入', '商机金额'),
                ('products', 'array', '是', '产品列表', '产品明细数组'),
                ('expected_close_date', 'date', '是', '预计成交日期', 'YYYY-MM-DD'),
                ('sales_rep', 'string', '是', '销售员', '负责销售员工号'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.order_no', 'string', '销售订单号', '生成的销售订单号'),
            ],
            'error_codes': [
                ('40901', '商机不存在', '确认商机ID'),
                ('40902', '客户不存在', '确认客户档案'),
                ('40903', '产品不存在', '确认产品档案'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
    ],
    'MES系统': [
        {
            'name': '生产工单下发接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': '金蝶 → MES', 'frequency': '实时',
            'description': '从金蝶下发生产工单到MES系统',
            'endpoint': '/api/v1/mes/production/order/push',
            'request_params': [
                ('mo_no', 'string', '是', '生产工单号', '金蝶生产订单号'),
                ('material_code', 'string', '是', '产品编码', '生产产品编码'),
                ('plan_qty', 'decimal', '是', '计划数量', '计划生产数量'),
                ('plan_start_date', 'date', '是', '计划开始日期', 'YYYY-MM-DD'),
                ('plan_end_date', 'date', '是', '计划完成日期', 'YYYY-MM-DD'),
                ('bom_version', 'string', '否', 'BOM版本', 'BOM版本号'),
                ('routing_code', 'string', '否', '工艺路线', '工艺路线编码'),
                ('workshop', 'string', '是', '车间', '生产车间编码'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.mes_mo_no', 'string', 'MES工单号', 'MES系统工单编号'),
                ('data.status', 'string', '状态', '已接收/已排产'),
            ],
            'error_codes': [
                ('41001', '工单已存在', '检查工单编号'),
                ('41002', '产品不存在', '确认产品档案'),
                ('41003', 'BOM不存在', '确认BOM档案'),
                ('41005', '车间不存在', '确认车间档案'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '完工汇报接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'MES → 金蝶', 'frequency': '实时',
            'description': '从MES系统汇报生产完工数据到金蝶',
            'endpoint': '/api/v1/mes/production/report',
            'request_params': [
                ('mo_no', 'string', '是', '生产工单号', 'MES工单编号'),
                ('report_no', 'string', '是', '汇报单号', 'MES汇报单号'),
                ('material_code', 'string', '是', '产品编码', '完工产品编码'),
                ('complete_qty', 'decimal', '是', '完工数量', '本次完工数量'),
                ('scrap_qty', 'decimal', '否', '报废数量', '本次报废数量'),
                ('work_hours', 'decimal', '否', '工时', '实际工时'),
                ('operator', 'string', '是', '操作员', '操作员工号'),
                ('report_time', 'datetime', '是', '汇报时间', 'YYYY-MM-DD HH:mm:ss'),
                ('warehouse', 'string', '是', '入库仓库', '完工入库仓库'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.kingdee_report_no', 'string', '金蝶汇报单号', '金蝶汇报单编号'),
                ('data.inventory_updated', 'boolean', '库存更新', '是否更新库存'),
            ],
            'error_codes': [
                ('41101', '工单不存在', '确认工单编号'),
                ('41102', '完工数量超限', '检查工单计划数量'),
                ('41103', '仓库不存在', '确认仓库档案'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
    ],
    'WMS系统': [
        {
            'name': '入库通知接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': '金蝶 → WMS', 'frequency': '实时',
            'description': '从金蝶下发入库通知到WMS系统',
            'endpoint': '/api/v1/wms/inbound/notice',
            'request_params': [
                ('notice_no', 'string', '是', '入库通知号', '金蝶入库通知单号'),
                ('notice_type', 'string', '是', '入库类型', '采购入库/生产入库/其他入库'),
                ('warehouse', 'string', '是', '仓库', '入库仓库编码'),
                ('details', 'array', '是', '明细列表', '入库明细数组'),
                ('detail.material_code', 'string', '是', '物料编码', '明细-物料编码'),
                ('detail.qty', 'decimal', '是', '数量', '明细-入库数量'),
                ('detail.batch_no', 'string', '否', '批次号', '明细-批次号'),
                ('source_no', 'string', '否', '来源单号', '采购订单/生产工单号'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.wms_notice_no', 'string', 'WMS通知号', 'WMS入库通知单号'),
                ('data.status', 'string', '状态', '已接收/已入库'),
            ],
            'error_codes': [
                ('41201', '通知单已存在', '检查通知单编号'),
                ('41202', '仓库不存在', '确认仓库档案'),
                ('41203', '物料不存在', '确认物料档案'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '出库通知接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': '金蝶 → WMS', 'frequency': '实时',
            'description': '从金蝶下发出库通知到WMS系统',
            'endpoint': '/api/v1/wms/outbound/notice',
            'request_params': [
                ('notice_no', 'string', '是', '出库通知号', '金蝶出库通知单号'),
                ('notice_type', 'string', '是', '出库类型', '销售出库/生产领料/其他出库'),
                ('warehouse', 'string', '是', '仓库', '出库仓库编码'),
                ('details', 'array', '是', '明细列表', '出库明细数组'),
                ('detail.material_code', 'string', '是', '物料编码', '明细-物料编码'),
                ('detail.qty', 'decimal', '是', '数量', '明细-出库数量'),
                ('detail.batch_no', 'string', '否', '批次号', '明细-批次号'),
                ('target_no', 'string', '否', '目标单号', '销售订单/生产工单号'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.wms_notice_no', 'string', 'WMS通知号', 'WMS出库通知单号'),
                ('data.available_qty', 'decimal', '可用库存', '当前可用库存量'),
            ],
            'error_codes': [
                ('41301', '通知单已存在', '检查通知单编号'),
                ('41302', '库存不足', '确认库存数量'),
                ('41303', '仓库不存在', '确认仓库档案'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
        {
            'name': '库存调整同步接口',
            'type': '数据同步', 'method': 'RESTful API',
            'direction': 'WMS → 金蝶', 'frequency': '实时',
            'description': '从WMS同步库存调整数据到金蝶',
            'endpoint': '/api/v1/wms/inventory/adjust',
            'request_params': [
                ('adjust_no', 'string', '是', '调整单号', 'WMS调整单号'),
                ('adjust_type', 'string', '是', '调整类型', '盘盈/盘亏/报废/其他'),
                ('warehouse', 'string', '是', '仓库', '仓库编码'),
                ('material_code', 'string', '是', '物料编码', '物料编码'),
                ('batch_no', 'string', '否', '批次号', '批次号'),
                ('adjust_qty', 'decimal', '是', '调整数量', '正增负减'),
                ('reason', 'string', '是', '调整原因', '调整原因说明'),
                ('operator', 'string', '是', '操作员', '操作员工号'),
                ('adjust_time', 'datetime', '是', '调整时间', 'YYYY-MM-DD HH:mm:ss'),
            ],
            'response_params': [
                ('code', 'integer', '响应码', '200=成功'),
                ('data.kingdee_adjust_no', 'string', '金蝶调整单号', '金蝶调整单编号'),
                ('data.new_qty', 'decimal', '调整后数量', '调整后库存数量'),
            ],
            'error_codes': [
                ('41401', '仓库不存在', '确认仓库档案'),
                ('41402', '物料不存在', '确认物料档案'),
                ('41403', '调整后库存为负', '检查调整数量'),
                ('50001', '系统异常', '联系运维人员排查'),
            ],
        },
    ],
}

INDUSTRY_SYSTEMS = {
    '制造业': ['财务系统', '供应链系统', 'OA系统', 'CRM系统', 'MES系统', 'WMS系统'],
    '零售业': ['财务系统', '供应链系统', 'OA系统', 'CRM系统'],
    '服务业': ['财务系统', 'OA系统', 'CRM系统'],
    '贸易业': ['财务系统', '供应链系统', 'OA系统', 'CRM系统'],
    '建筑业': ['财务系统', '供应链系统', 'OA系统'],
    'default': ['财务系统', '供应链系统', 'OA系统'],
}

DATA_MAPPINGS = [
    {'data_type': '物料主数据', 'source': '外部系统', 'target': '金蝶云·星空',
     'fields': [
         ('物料编码', 'material_code', 'FNumber', '直接映射'),
         ('物料名称', 'material_name', 'FName', '直接映射'),
         ('规格型号', 'spec', 'FSpecification', '直接映射'),
         ('基本单位', 'base_unit', 'FBaseUnitId', '单位映射表'),
         ('物料分组', 'material_group', 'FMaterialGroup', '分组映射表'),
         ('采购单位', 'purchase_unit', 'FPurchaseUnitId', '单位映射表'),
         ('销售单位', 'sales_unit', 'FSaleUnitId', '单位映射表'),
         ('创建组织', 'create_org', 'FCreateOrgId', '组织编码'),
     ]},
    {'data_type': '客户主数据', 'source': 'CRM', 'target': '金蝶云·星空',
     'fields': [
         ('客户编码', 'customer_code', 'FNumber', '直接映射'),
         ('客户名称', 'customer_name', 'FName', '直接映射'),
         ('客户类型', 'customer_type', 'FCustTypeId', '类型映射表'),
         ('信用额度', 'credit_limit', 'FCreditLimit', '金额转换'),
         ('联系人', 'contact_person', 'FLinkman', '直接映射'),
         ('联系电话', 'contact_phone', 'FTelephone', '直接映射'),
         ('地址', 'address', 'FAddress', '直接映射'),
         ('销售员', 'sales_rep', 'FSalerId', '员工映射表'),
     ]},
    {'data_type': '供应商主数据', 'source': 'SRM', 'target': '金蝶云·星空',
     'fields': [
         ('供应商编码', 'supplier_code', 'FNumber', '直接映射'),
         ('供应商名称', 'supplier_name', 'FName', '直接映射'),
         ('供应商类型', 'supplier_type', 'FSupplierTypeId', '类型映射表'),
         ('联系人', 'contact_person', 'FContact', '直接映射'),
         ('银行账号', 'bank_account', 'FBankCode', '直接映射'),
         ('开户行', 'bank_name', 'FOpenBankName', '直接映射'),
         ('创建组织', 'create_org', 'FCreateOrgId', '组织编码'),
     ]},
    {'data_type': '员工主数据', 'source': 'HR/OA', 'target': '金蝶云·星空',
     'fields': [
         ('工号', 'employee_code', 'FStaffNumber', '直接映射'),
         ('姓名', 'name', 'FName', '直接映射'),
         ('部门', 'department', 'FPostDeptId', '部门映射表'),
         ('岗位', 'position', 'FPostId', '岗位映射表'),
         ('入职日期', 'hire_date', 'FEntryDate', '日期格式转换'),
         ('手机号', 'mobile', 'FMobile', '直接映射'),
         ('邮箱', 'email', 'FEmail', '直接映射'),
         ('状态', 'status', 'FDocumentStatus', '状态映射表'),
     ]},
    {'data_type': '组织主数据', 'source': 'OA/HR', 'target': '金蝶云·星空',
     'fields': [
         ('组织编码', 'org_code', 'FNumber', '直接映射'),
         ('组织名称', 'org_name', 'FName', '直接映射'),
         ('上级组织', 'parent_org', 'FParentId', '组织映射表'),
         ('组织类型', 'org_type', 'FOrgFormId', '类型映射表'),
         ('负责人', 'manager', 'FManagerId', '员工映射表'),
     ]},
]

TEST_CASES = [
    {'interface': '凭证同步接口', 'cases': [
        ('TC-FIN-001', '正常凭证同步', '借贷平衡、科目存在', '成功创建凭证', '高'),
        ('TC-FIN-002', '借贷不平衡', '借方≠贷方', '返回40003错误', '高'),
        ('TC-FIN-003', '科目不存在', '无效科目编码', '返回40002错误', '高'),
        ('TC-FIN-004', '必填字段缺失', '缺少凭证日期', '返回40001错误', '高'),
        ('TC-FIN-005', '日期格式错误', 'YYYY/MM/DD格式', '返回40001错误', '中'),
    ]},
    {'interface': '采购订单同步接口', 'cases': [
        ('TC-SCM-001', '正常采购订单', '完整参数', '成功创建订单', '高'),
        ('TC-SCM-002', '供应商不存在', '无效供应商编码', '返回40301错误', '高'),
        ('TC-SCM-003', '物料不存在', '无效物料编码', '返回40302错误', '高'),
        ('TC-SCM-004', '数量为负', '负数采购量', '返回参数校验错误', '高'),
    ]},
    {'interface': '库存同步接口', 'cases': [
        ('TC-WMS-001', '正常库存同步', '完整库存数据', '同步成功', '高'),
        ('TC-WMS-002', '仓库不存在', '无效仓库编码', '返回40501错误', '高'),
        ('TC-WMS-003', '数量为负', '负数库存量', '返回40503错误', '高'),
        ('TC-WMS-004', '重复同步', '相同同步批次号', '幂等返回成功', '中'),
    ]},
    {'interface': '审批流程触发接口', 'cases': [
        ('TC-OA-001', '正常触发审批', '完整审批参数', '流程创建成功', '高'),
        ('TC-OA-002', '流程模板不存在', '无效单据类型', '返回40601错误', '高'),
        ('TC-OA-003', '申请人不存在', '无效员工工号', '返回40602错误', '高'),
        ('TC-OA-004', '回调地址无效', '非URL格式', '返回40604错误', '中'),
    ]},
    {'interface': '生产工单下发接口', 'cases': [
        ('TC-MES-001', '正常工单下发', '完整工单参数', 'MES接收成功', '高'),
        ('TC-MES-002', '工单已存在', '重复工单号', '返回41001错误', '高'),
        ('TC-MES-003', '产品不存在', '无效产品编码', '返回41002错误', '高'),
        ('TC-MES-004', 'BOM不存在', '无效BOM版本', '返回41003错误', '中'),
    ]},
    {'interface': '客户档案同步接口', 'cases': [
        ('TC-CRM-001', '正常客户同步', '完整客户数据', '同步成功', '高'),
        ('TC-CRM-002', '客户已存在', '重复客户编码', '返回40801错误', '高'),
        ('TC-CRM-003', '名称重复', '重名客户', '返回40802错误', '中'),
    ]},
]

FAQS = [
    ('接口调用频率有限制吗？',
     '单个接口默认限流1000次/分钟，批量接口500次/分钟。如需更高频率，请联系运维申请提升配额。'),
    ('接口数据传输是否加密？',
     '所有接口均使用HTTPS加密传输，敏感字段（银行账号、身份证号等）使用AES-256二次加密。'),
    ('接口调用失败如何处理？',
     '系统会自动重试3次（间隔5秒/15秒/60秒）。超过重试次数后进入死信队列，需人工介入处理。'),
    ('如何查看接口调用日志？',
     '通过Kibana日志平台查看，索引模式：integration-log-*。支持按接口、时间、状态筛选。'),
    ('接口认证Token过期怎么办？',
     'Token有效期24小时，过期后需重新获取。建议在Token有效期剩余30分钟时自动刷新。'),
    ('批量数据同步接口一次最多传多少条？',
     '单次批量请求最多1000条记录。超过1000条请分页传输，每页不超过1000条。'),
    ('如何确认数据已同步成功？',
     '每个接口返回唯一sync_id，可通过查询接口传入sync_id确认同步状态。也可通过监控平台实时查看。'),
    ('系统升级时接口会受影响吗？',
     '金蝶云版本升级前会提前通知，API采用版本号管理（v1/v2），旧版本接口至少保留6个月兼容期。'),
    ('接口响应超时怎么处理？',
     '实时接口超时时间5秒，批量接口超时时间30秒。建议使用异步调用模式，通过回调获取结果。'),
    ('如何申请新增接口？',
     '通过项目管理平台提交接口需求申请，附上接口说明文档，经技术评审后排期开发，一般5个工作日内完成。'),
]


# ===================== Word 文档格式工具 =====================

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_body_text(doc, text, indent=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    if bold:
        run.font.bold = True
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return p

def styled_table(doc, headers, rows, header_color='006699', alt_color='F0F6FC'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            if i % 2 == 1:
                set_cell_shading(cell, alt_color)
    return t

def setup_page(doc, company_name, doc_type='系统集成方案'):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(company_name + '  ' + doc_type)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '006699')
    pBdr.append(bottom)
    hp._p.get_or_add_pPr().append(pBdr)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run('  \u7b2c  ')
    r1.font.name = '宋体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r1.font.size = Pt(9)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page = fp.add_run()
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page.font.name = '宋体'
    run_page._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_page.font.size = Pt(9)
    r_sep = fp.add_run('  \u9875  |  ')
    r_sep.font.name = '宋体'
    r_sep._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r_sep.font.size = Pt(9)
    r_sep.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r_date = fp.add_run(datetime.now().strftime('%Y\u5e74%m\u6708%d\u65e5'))
    r_date.font.name = '宋体'
    r_date._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r_date.font.size = Pt(9)
    r_date.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    pBdr2 = OxmlElement('w:pBdr')
    top2 = OxmlElement('w:top')
    top2.set(qn('w:val'), 'single')
    top2.set(qn('w:sz'), '4')
    top2.set(qn('w:space'), '1')
    top2.set(qn('w:color'), '006699')
    pBdr2.append(top2)
    fp._p.get_or_add_pPr().append(pBdr2)

def get_industry_systems(industry):
    for k, v in INDUSTRY_SYSTEMS.items():
        if k in industry or industry in k:
            return v
    return INDUSTRY_SYSTEMS['default']


# ===================== 主文档生成函数 =====================
def generate_integration_v3(customer_info):
    doc = Document()
    company = customer_info.get('companyName', '企业名称')
    industry = customer_info.get('industry', '制造业')
    systems = get_industry_systems(industry)
    all_interfaces = []
    for sys_name in systems:
        if sys_name in INTERFACES:
            all_interfaces.extend([(sys_name, i) for i in INTERFACES[sys_name]])
    total_ifaces = len(all_interfaces)

    setup_page(doc, company)

    # ====== 封面页 ======
    for _ in range(6):
        doc.add_paragraph('')
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run(company)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    doc.add_paragraph('')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('系统集成方案')
    r2.font.name = '宋体'; r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r2.font.size = Pt(22); r2.font.bold = True; r2.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('（V3.0）')
    r3.font.name = '宋体'; r3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r3.font.size = Pt(14); r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(3):
        doc.add_paragraph('')

    info_t = doc.add_table(rows=6, cols=2)
    info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ('文档类型', '系统集成方案'),
        ('企业名称', company),
        ('所属行业', industry),
        ('版本号', 'V3.0'),
        ('编制日期', datetime.now().strftime('%Y年%m月%d日')),
        ('编制单位', '金蝶软件（中国）有限公司'),
    ]):
        ck = info_t.rows[i].cells[0]
        cv = info_t.rows[i].cells[1]
        ck.text = ''; cv.text = ''
        rk = ck.paragraphs[0].add_run(k)
        rk.font.name = '宋体'; rk._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        rk.font.size = Pt(12); rk.font.bold = True; rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(ck, '006699')
        rv = cv.paragraphs[0].add_run(v)
        rv.font.name = '宋体'; rv._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        rv.font.size = Pt(12)
        set_cell_shading(cv, 'E8F0FE')
        ck.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cv.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ====== 修订记录 ======
    add_heading_styled(doc, '修订记录', 1)
    styled_table(doc, ['版本', '日期', '修订人', '修订内容', '审核人'], [
        ('V1.0', '2024-01-15', '项目经理', '初稿编写', '技术总监'),
        ('V2.0', '2024-03-01', '集成顾问', '补充接口详细设计', '技术总监'),
        ('V3.0', datetime.now().strftime('%Y-%m-%d'), '集成顾问', '全面升级：15+接口/数据映射/测试用例/FAQ', '技术总监'),
    ])
    doc.add_paragraph('')

    # ====== 目录页 ======
    doc.add_page_break()
    add_heading_styled(doc, '目  录', 1)
    toc_items = [
        ('一、集成概述', True),
        ('    1.1 集成目标', False),
        ('    1.2 集成原则', False),
        ('    1.3 集成范围', False),
        ('二、集成架构设计', True),
        ('    2.1 总体架构', False),
        ('    2.2 架构图描述', False),
        ('    2.3 技术方案', False),
        ('三、接口详细设计', True),
        ('    3.x 各系统接口（请求/响应/错误码）', False),
        ('四、数据映射规则', True),
        ('    4.1 主数据映射', False),
        ('    4.2 数据质量规则', False),
        ('五、接口测试方案', True),
        ('    5.1 测试用例', False),
        ('    5.2 测试环境要求', False),
        ('六、异常处理机制', True),
        ('七、接口监控与告警', True),
        ('八、实施计划', True),
        ('九、常见问题FAQ', True),
        ('附录A：接口清单汇总', True),
        ('附录B：术语表', True),
    ]
    for item, is_bold in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in p.runs:
            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
            if is_bold:
                run.font.bold = True

    doc.add_page_break()

    # ====== 一、集成概述 ======
    add_heading_styled(doc, '一、集成概述', 1)
    add_heading_styled(doc, '1.1 集成目标', 2)
    for item in ['打通信息孤岛，实现业务系统间数据实时共享与协同',
                 '消除手工重复录入，提升工作效率60%以上',
                 '确保数据一致性，建立主数据管理体系',
                 '实现业务流程端到端贯通，支撑数字化转型战略']:
        add_bullet(doc, item)

    add_heading_styled(doc, '1.2 集成原则', 2)
    styled_table(doc, ['原则', '说明'], [
        ('标准化', '统一接口规范，使用RESTful API，JSON数据格式'),
        ('可靠性', '确保数据传输完整性，支持断点续传和幂等处理'),
        ('实时性', '关键业务数据实时同步，非关键数据定时同步'),
        ('安全性', 'OAuth2.0认证，HTTPS加密传输，敏感字段AES-256加密'),
        ('可扩展性', '接口版本化管理，支持水平扩展'),
        ('可维护性', '完整接口日志，链路追踪，问题可追溯'),
    ], '006699')

    add_heading_styled(doc, '1.3 集成范围', 2)
    add_body_text(doc, f'本次集成涉及 {len(systems)} 个外部系统，共 {total_ifaces} 个标准接口。')
    doc.add_paragraph('')
    sys_rows = []
    for sys_name in systems:
        if sys_name in INTERFACES:
            count = len(INTERFACES[sys_name])
            sys_rows.append((sys_name, str(count), '双向', '详见接口设计章节'))
    styled_table(doc, ['系统名称', '接口数量', '数据方向', '说明'], sys_rows, '336699')

    doc.add_page_break()

    # ====== 二、集成架构设计 ======
    add_heading_styled(doc, '二、集成架构设计', 1)
    add_heading_styled(doc, '2.1 总体架构', 2)
    add_body_text(doc, f'{company}系统集成采用企业服务总线（ESB）+ API网关模式，金蝶云·星空作为核心业务平台，通过标准API接口与各外部系统进行数据交互。', indent=True)

    add_heading_styled(doc, '2.2 集成架构图描述', 2)
    add_body_text(doc, '系统集成分为四层架构，自上而下依次为：', bold=True)
    arch_layers = [
        ('应用层（业务系统）', '金蝶云·星空 | OA | CRM | MES | WMS | SRM | 银行', '各业务系统通过API与集成层交互'),
        ('集成层（ESB/API网关）', 'API网关 → 认证鉴权 → 路由分发\n消息队列(Kafka) → 数据转换 → 协议适配', '统一入口，支持限流/熔断/日志/监控'),
        ('数据层（数据管理）', '主数据管理(MDM)\n数据映射引擎 | 数据质量校验 | ETL', '统一主数据标准，保障数据一致性'),
        ('监控层（运维保障）', 'Prometheus+Grafana（性能）\nELK Stack（日志） | 告警通知', '实时监控接口健康度，异常自动告警'),
    ]
    styled_table(doc, ['架构层次', '核心组件', '职责说明'], arch_layers, '336699')

    doc.add_paragraph('')
    add_body_text(doc, '数据流向说明：', bold=True)
    add_body_text(doc, '1. 外部系统发起请求，经API网关统一入口进入集成层。')
    add_body_text(doc, '2. API网关进行身份认证（OAuth2.0）、权限校验、限流控制。')
    add_body_text(doc, '3. 根据路由规则将请求分发到消息队列（Kafka）或直接转发。')
    add_body_text(doc, '4. 数据转换引擎对源系统数据进行格式转换、字段映射、编码转换。')
    add_body_text(doc, '5. 转换后的数据通过金蝶Cloud API写入金蝶云·星空。')
    add_body_text(doc, '6. 全链路日志记录到ELK，关键指标采集到Prometheus，异常触发Grafana告警。')

    add_heading_styled(doc, '2.3 技术方案', 2)
    styled_table(doc, ['技术项', '方案选型'], [
        ('接口协议', 'RESTful API / WebService'),
        ('数据格式', 'JSON（默认）/ XML（兼容旧系统）'),
        ('认证方式', 'OAuth2.0 / API Key（内网场景）'),
        ('传输加密', 'HTTPS (TLS 1.2+)'),
        ('消息队列', 'Apache Kafka（异步解耦）'),
        ('日志记录', 'ELK Stack（Elasticsearch + Logstash + Kibana）'),
        ('监控告警', 'Prometheus + Grafana'),
        ('链路追踪', 'SkyWalking / Jaeger'),
        ('API网关', 'Kong / APISIX'),
        ('配置中心', 'Nacos / Apollo'),
    ], '336699')

    doc.add_page_break()

    # ====== 三、接口详细设计 ======
    add_heading_styled(doc, '三、接口详细设计', 1)
    add_body_text(doc, f'本章详细描述本次集成的 {total_ifaces} 个标准接口设计，每个接口包含：接口基本信息、请求参数、响应参数、错误码定义。', indent=True)

    iface_idx = 0
    for sys_name in systems:
        if sys_name not in INTERFACES:
            continue
        for iface in INTERFACES[sys_name]:
            iface_idx += 1
            add_heading_styled(doc, f'3.{iface_idx} {iface["name"]}（{sys_name}）', 2)
            add_body_text(doc, f'接口名称：{iface["name"]}', bold=True)
            add_body_text(doc, f'接口类型：{iface["type"]}')
            add_body_text(doc, f'调用方式：{iface["method"]}')
            add_body_text(doc, f'接口方向：{iface["direction"]}')
            add_body_text(doc, f'同步频率：{iface["frequency"]}')
            add_body_text(doc, f'接口地址：{iface["endpoint"]}')
            add_body_text(doc, f'接口描述：{iface["description"]}', indent=True)
            doc.add_paragraph('')

            add_heading_styled(doc, f'3.{iface_idx}.1 请求参数', 3)
            styled_table(doc, ['参数名', '类型', '必填', '说明', '备注'],
                [(p[0], p[1], p[2], p[3], p[4]) for p in iface['request_params']], '006699')
            doc.add_paragraph('')

            add_heading_styled(doc, f'3.{iface_idx}.2 响应参数', 3)
            styled_table(doc, ['字段名', '类型', '说明', '备注'],
                [(p[0], p[1], p[2], p[3]) for p in iface['response_params']], '006699')
            doc.add_paragraph('')

            add_heading_styled(doc, f'3.{iface_idx}.3 错误码定义', 3)
            styled_table(doc, ['错误码', '错误描述', '处理建议'],
                [(e[0], e[1], e[2]) for e in iface['error_codes']], '993333')
            doc.add_paragraph('')

    doc.add_page_break()

    # ====== 四、数据映射规则 ======
    add_heading_styled(doc, '四、数据映射规则', 1)
    add_heading_styled(doc, '4.1 主数据映射', 2)
    add_body_text(doc, '以下为各主数据在源系统与金蝶云·星空之间的字段映射规则：', indent=True)

    for dm in DATA_MAPPINGS:
        add_body_text(doc, f"\u25ce {dm['data_type']}（{dm['source']} → {dm['target']}）", bold=True)
        styled_table(doc, ['源系统字段', '源字段编码', '金蝶字段', '转换规则'],
            [(f[0], f[1], f[2], f[3]) for f in dm['fields']], '336633')
        doc.add_paragraph('')

    add_heading_styled(doc, '4.2 数据质量规则', 2)
    styled_table(doc, ['规则类型', '说明', '保障措施'], [
        ('完整性', '必填字段不能为空', '接口层参数校验 + 业务层校验'),
        ('准确性', '数据值在有效范围内', '业务规则校验 + 枚举值校验'),
        ('一致性', '同一数据在多个系统中保持一致', '主数据管理（MDM）+ 变更同步'),
        ('及时性', '数据在规定时间内同步完成', '监控告警 + 延迟统计'),
        ('唯一性', '编码全局唯一', '编码规则 + 数据库唯一索引'),
        ('规范性', '数据格式符合标准', '数据清洗 + 格式校验'),
    ], '666633')

    doc.add_page_break()

    # ====== 五、接口测试方案 ======
    add_heading_styled(doc, '五、接口测试方案', 1)
    add_heading_styled(doc, '5.1 测试用例', 2)
    add_body_text(doc, '以下为各接口的核心测试用例，覆盖正常场景和异常场景：', indent=True)

    iface_names = [i['name'] for _, i in all_interfaces]
    for tc_group in TEST_CASES:
        if tc_group['interface'] not in iface_names:
            continue
        add_body_text(doc, f"\u25ce {tc_group['interface']}", bold=True)
        styled_table(doc, ['用例编号', '用例名称', '前置条件/输入', '预期结果', '优先级'],
            [(c[0], c[1], c[2], c[3], c[4]) for c in tc_group['cases']], '336699')
        doc.add_paragraph('')

    add_heading_styled(doc, '5.2 测试环境要求', 2)
    styled_table(doc, ['项目', '要求'], [
        ('金蝶环境', '金蝶云·星空测试环境（最新补丁版本）'),
        ('外部系统', '各外部系统测试环境，已部署对接接口'),
        ('网络环境', '测试环境网络互通，防火墙已放行接口端口'),
        ('测试数据', '准备基础主数据（物料/客户/供应商/组织/员工）'),
        ('工具', 'Postman / JMeter（性能测试）'),
        ('文档', '接口文档V3.0（本文档）'),
    ], '336699')

    doc.add_page_break()

    # ====== 六、异常处理机制 ======
    add_heading_styled(doc, '六、异常处理机制', 1)
    add_heading_styled(doc, '6.1 异常分类与处理', 2)
    styled_table(doc, ['异常类型', '描述', '处理方式', '告警方式'], [
        ('网络异常', '网络中断、超时、DNS解析失败', '自动重试3次（5s/15s/60s），超过后进入死信队列', '运维告警'),
        ('数据异常', '格式错误、校验失败、编码不存在', '记录错误日志，返回明确错误码，人工处理', '业务告警'),
        ('系统异常', '目标系统不可用、数据库故障', '消息队列暂存，系统恢复后自动重发', '运维告警'),
        ('业务异常', '规则冲突、状态不匹配、审批驳回', '记录冲突日志，推送通知相关处理人', '业务告警'),
        ('性能异常', '响应超时、吞吐量不足、队列堆积', '自动限流降级，异步处理，扩容通知', '运维告警'),
        ('安全异常', '认证失败、越权访问、数据篡改', '记录安全日志，封禁异常IP，通知安全团队', '安全告警'),
    ], '993333')

    add_heading_styled(doc, '6.2 异常处理流程', 2)
    add_body_text(doc, '接口调用 → 参数校验 → 认证鉴权 → 业务处理 → 结果校验 → 成功：更新状态并返回 → 失败：记录日志 → 自动重试 → 超过阈值：进入死信队列 → 人工处理 → 闭环确认', indent=True)

    doc.add_page_break()

    # ====== 七、接口监控与告警 ======
    add_heading_styled(doc, '七、接口监控与告警', 1)
    add_heading_styled(doc, '7.1 监控指标', 2)
    styled_table(doc, ['监控指标', '阈值', '采集频率'], [
        ('接口成功率', '>=99.5%', '每5分钟'),
        ('接口响应时间', '<=3秒（实时）/ <=60秒（批量）', '每5分钟'),
        ('数据同步延迟', '<=5分钟（实时）/ <=1小时（批量）', '每15分钟'),
        ('接口调用量', '按小时/天统计', '每小时'),
        ('错误数量', '按错误类型/接口统计', '每5分钟'),
        ('队列积压', '<=1000条（Kafka积压量）', '每分钟'),
    ], '336699')

    add_heading_styled(doc, '7.2 告警机制', 2)
    styled_table(doc, ['告警级别', '触发条件', '通知方式', '响应时间'], [
        ('P0-严重', '接口完全不可用或成功率<95%', '短信+电话+邮件', '5分钟内'),
        ('P1-重要', '成功率<99%或响应超时>10秒', '短信+邮件', '15分钟内'),
        ('P2-一般', '数据延迟或调用量波动>50%', '邮件+钉钉', '1小时内'),
        ('P3-提示', '磁盘空间/内存使用率>80%', '邮件', '工作时间内'),
    ], '993333')

    doc.add_page_break()

    # ====== 八、实施计划 ======
    add_heading_styled(doc, '八、实施计划', 1)
    add_heading_styled(doc, '8.1 阶段计划', 2)
    styled_table(doc, ['阶段', '名称', '周期', '主要内容'], [
        ('第1阶段', '需求调研', '2周', '梳理集成需求，确定接口规格，输出接口设计文档'),
        ('第2阶段', '环境搭建', '1周', '部署API网关、消息队列、监控平台等基础环境'),
        ('第3阶段', '接口开发', '3周', '开发接口程序，编写单元测试，Code Review'),
        ('第4阶段', '联调测试', '2周', '与各外部系统联调，执行集成测试，缺陷修复'),
        ('第5阶段', '性能测试', '1周', '压力测试，性能调优，安全测试'),
        ('第6阶段', 'UAT验收', '1周', '业务用户验收测试，培训操作人员'),
        ('第7阶段', '上线部署', '1周', '生产环境部署，灰度发布，全量切换'),
    ], '006699')

    add_heading_styled(doc, '8.2 里程碑节点', 2)
    styled_table(doc, ['里程碑', '节点名称', '时间', '参与方'], [
        ('M1', '需求评审通过', '第2周末', '项目经理、技术负责人、业务方'),
        ('M2', '接口开发完成', '第6周末', '开发团队'),
        ('M3', '联调测试通过', '第8周末', '测试团队、外部系统方'),
        ('M4', 'UAT验收通过', '第10周末', '业务用户、项目经理'),
        ('M5', '正式上线', '第11周末', '全体项目成员'),
    ], '006699')

    doc.add_page_break()

    # ====== 九、常见问题FAQ ======
    add_heading_styled(doc, '九、常见问题FAQ', 1)
    for i, (q, a) in enumerate(FAQS, 1):
        add_body_text(doc, f'Q{i}. {q}', bold=True)
        add_body_text(doc, f'    A：{a}')
        doc.add_paragraph('')

    doc.add_page_break()

    # ====== 附录A ======
    add_heading_styled(doc, '附录A：接口清单汇总', 1)
    summary_rows = []
    for sys_name in systems:
        if sys_name not in INTERFACES:
            continue
        for iface in INTERFACES[sys_name]:
            summary_rows.append((iface['name'], sys_name, iface['type'], iface['direction'],
                                 iface['method'], iface['frequency'], iface['endpoint']))
    styled_table(doc, ['接口名称', '所属系统', '接口类型', '数据方向', '调用方式', '频率', '接口地址'],
                 summary_rows, '666666')
    doc.add_paragraph('')
    add_body_text(doc, f'合计：{len(systems)} 个系统，{total_ifaces} 个标准接口。', bold=True)

    # ====== 附录B ======
    doc.add_page_break()
    add_heading_styled(doc, '附录B：术语表', 1)
    styled_table(doc, ['术语', '说明'], [
        ('ESB', 'Enterprise Service Bus，企业服务总线'),
        ('API', 'Application Programming Interface，应用程序编程接口'),
        ('OAuth2.0', '开放授权协议，用于接口认证'),
        ('JSON', 'JavaScript Object Notation，轻量级数据交换格式'),
        ('Kafka', 'Apache Kafka，分布式消息队列'),
        ('ELK', 'Elasticsearch + Logstash + Kibana，日志管理平台'),
        ('MDM', 'Master Data Management，主数据管理'),
        ('UAT', 'User Acceptance Testing，用户验收测试'),
        ('REST', 'Representational State Transfer，表述性状态转移'),
        ('HTTPS', 'Hypertext Transfer Protocol Secure，安全超文本传输协议'),
        ('MES', 'Manufacturing Execution System，制造执行系统'),
        ('WMS', 'Warehouse Management System，仓库管理系统'),
        ('SRM', 'Supplier Relationship Management，供应商关系管理'),
        ('CRM', 'Customer Relationship Management，客户关系管理'),
        ('OA', 'Office Automation，办公自动化'),
    ], '006699')

    # 保存
    filename = f"{customer_info.get('customerCode', company)}_系统集成方案_v3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return {
        'success': True,
        'filepath': filepath,
        'filename': filename,
        'version': '3.0',
        'systems': len(systems),
        'interfaces': total_ifaces,
    }

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='金蝶系统集成方案生成器 V3.0')
    parser.add_argument('--companyName', required=True, help='企业名称')
    parser.add_argument('--companySize', default='中型企业', help='企业规模')
    parser.add_argument('--customerCode', default='', help='客户编码')
    parser.add_argument('--industry', default='制造业', help='所属行业（制造业/零售业/服务业/贸易业/建筑业）')
    parser.add_argument('--employees', default='', help='员工人数')
    parser.add_argument('--revenue', default='', help='年营业额')
    parser.add_argument('--modules', default='', help='实施模块（逗号分隔）')
    args = parser.parse_args()
    customer_info = {
        'companyName': args.companyName,
        'customerCode': args.customerCode or args.companyName,
        'industry': args.industry,
    }
    result = generate_integration_v3(customer_info)
    print(json.dumps(result, ensure_ascii=False, indent=2))
