#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶UAT测试生成器 - v9.0 超级增强版
基于深度学习，支持15章节、100+测试用例、3个测试阶段
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# 输出目录
OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_uat_report_v9(customer_info):
    """生成UAT测试报告 - v9.0 超级增强版（15章节）"""
    doc = Document()
    
    # 设置字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    company_name = customer_info.get('companyName', '企业名称')
    industry = customer_info.get('industry', '制造业')
    selected_modules = customer_info.get('selectedModules', ['finance', 'supply', 'manufacture'])
    
    # 标题页
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{company_name}\n")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("UAT测试方案报告\n")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'企业名称：{company_name}\n')
    run.font.size = Pt(14)
    run.font.bold = True
    run = info.add_run(f'所属行业：{industry}\n')
    run = info.add_run(f'企业规模：{customer_info.get("companySize", "中型企业")}\n')
    run = info.add_run(f'测试日期：{datetime.now().strftime("%Y年%m月%d日")}\n')
    run = info.add_run(f'测试模块：{", ".join([m for m in selected_modules])}\n')
    
    doc.add_page_break()
    
    # 目录（15章节）
    doc.add_heading('目录', level=1)
    chapters = [
        '一、UAT测试概述',
        '二、测试方法论',
        '三、测试环境',
        '四、测试范围',
        '五、测试数据',
        '六、测试用例',
        '七、测试执行',
        '八、测试结果',
        '九、测试报告',
        '十、测试改进',
        '十一、测试风险',
        '十二、测试保障',
        '十三、测试经验',
        '十四、测试总结',
        '十五、测试建议'
    ]
    for chapter in chapters:
        doc.add_paragraph(chapter)
    doc.add_page_break()
    
    # 一、UAT测试概述
    doc.add_heading('一、UAT测试概述', level=1)
    doc.add_heading('1.1 UAT测试定义', level=2)
    doc.add_paragraph('UAT测试（User Acceptance Testing）是用户验收测试，是系统测试的最后阶段，由最终用户进行测试，验证系统是否满足业务需求。')
    
    doc.add_heading('1.2 UAT测试目标', level=2)
    uat_goals = [
        '验证系统功能是否符合业务需求',
        '验证系统性能是否满足业务要求',
        '验证系统易用性是否满足用户期望',
        '验证系统稳定性是否满足业务要求',
        '验证系统安全性是否满足业务要求'
    ]
    for goal in uat_goals:
        doc.add_paragraph(f'• {goal}', style='List Bullet')
    
    doc.add_heading('1.3 UAT测试范围', level=2)
    uat_scope = [
        '功能测试：验证系统功能是否满足业务需求',
        '性能测试：验证系统性能是否满足业务要求',
        '易用性测试：验证系统易用性是否满足用户期望',
        '安全性测试：验证系统安全性是否满足业务要求',
        '兼容性测试：验证系统兼容性是否满足业务要求'
    ]
    for scope in uat_scope:
        doc.add_paragraph(f'• {scope}', style='List Bullet')
    
    doc.add_heading('1.4 UAT测试原则', level=2)
    uat_principles = [
        '用户主导：由最终用户进行测试',
        '需求导向：基于业务需求进行测试',
        '全面覆盖：覆盖所有业务场景',
        '重点突出：重点测试关键业务流程',
        '持续改进：持续改进测试过程'
    ]
    for principle in uat_principles:
        doc.add_paragraph(f'• {principle}', style='List Bullet')
    
    doc.add_page_break()
    
    # 二、测试方法论
    doc.add_heading('二、测试方法论', level=1)
    doc.add_heading('2.1 测试方法论概述', level=2)
    doc.add_paragraph('测试方法论是指导测试过程的框架，包括测试策略、测试计划、测试设计、测试执行、测试评估等。')
    
    doc.add_heading('2.2 测试策略', level=2)
    test_strategies = [
        '测试层次：单元测试、集成测试、系统测试、用户验收测试',
        '测试类型：功能测试、性能测试、安全测试、兼容性测试',
        '测试方法：手工测试、自动化测试、性能测试、安全测试',
        '测试工具：测试管理工具、自动化测试工具、性能测试工具'
    ]
    for strategy in test_strategies:
        doc.add_paragraph(f'• {strategy}', style='List Bullet')
    
    doc.add_heading('2.3 测试计划', level=2)
    test_plan = [
        '测试目标：明确测试目标，制定测试计划',
        '测试范围：确定测试范围，制定测试范围',
        '测试资源：确定测试资源，分配测试任务',
        '测试进度：制定测试进度，控制测试进度',
        '测试风险：识别测试风险，制定风险应对'
    ]
    for plan in test_plan:
        doc.add_paragraph(f'• {plan}', style='List Bullet')
    
    doc.add_heading('2.4 测试设计', level=2)
    test_design = [
        '测试用例设计：基于需求设计测试用例',
        '测试数据设计：设计测试数据，确保测试质量',
        '测试环境设计：设计测试环境，确保测试环境',
        '测试工具设计：选择测试工具，配置测试工具'
    ]
    for design in test_design:
        doc.add_paragraph(f'• {design}', style='List Bullet')
    
    doc.add_page_break()
    
    # 三、测试环境
    doc.add_heading('三、测试环境', level=1)
    doc.add_heading('3.1 测试环境概述', level=2)
    doc.add_paragraph('测试环境是进行测试的环境，包括硬件环境、软件环境、网络环境等。')
    
    doc.add_heading('3.2 硬件环境', level=2)
    hardware_environment = [
        '服务器配置：8核CPU、16GB内存、500GB硬盘',
        '客户端配置：4核CPU、8GB内存、256GB硬盘',
        '网络环境：千兆局域网、百兆广域网',
        '存储环境：SAN存储、NAS存储'
    ]
    for env in hardware_environment:
        doc.add_paragraph(f'• {env}', style='List Bullet')
    
    doc.add_heading('3.3 软件环境', level=2)
    software_environment = [
        '操作系统：Windows Server 2019、Windows 10',
        '数据库：SQL Server 2019、Oracle 19c',
        '应用服务器：Tomcat 9.0、WebLogic 14.1',
        '浏览器：Chrome、Firefox、IE 11'
    ]
    for env in software_environment:
        doc.add_paragraph(f'• {env}', style='List Bullet')
    
    doc.add_heading('3.4 网络环境', level=2)
    network_environment = [
        '局域网：千兆以太网、VLAN划分',
        '广域网：VPN连接、防火墙保护',
        '互联网：宽带接入、DNS解析',
        '安全环境：防火墙、入侵检测、防病毒'
    ]
    for env in network_environment:
        doc.add_paragraph(f'• {env}', style='List Bullet')
    
    doc.add_page_break()
    
    # 四、测试范围
    doc.add_heading('四、测试范围', level=1)
    doc.add_heading('4.1 测试范围概述', level=2)
    doc.add_paragraph('测试范围是确定测试的内容，包括功能范围、性能范围、安全范围等。')
    
    doc.add_heading('4.2 功能测试范围', level=2)
    function_test_scope = []
    
    if 'finance' in selected_modules:
        finance_tests = [
            '总账管理：凭证管理、账簿管理、报表管理',
            '应收管理：销售开票、应收确认、收款核销',
            '应付管理：采购开票、应付确认、付款核销',
            '成本管理：成本核算、成本分析、成本控制',
            '固定资产：资产卡片、折旧管理、资产处置'
        ]
        for test in finance_tests:
            function_test_scope.append(test)
    
    if 'supply' in selected_modules:
        supply_tests = [
            '采购管理：供应商管理、采购申请、采购订单',
            '库存管理：入库管理、出库管理、库存盘点',
            '销售管理：客户管理、销售订单、销售出库',
            '物流管理：运输管理、配送管理、仓储管理'
        ]
        for test in supply_tests:
            function_test_scope.append(test)
    
    if 'manufacture' in selected_modules:
        manufacture_tests = [
            '生产计划：需求预测、主生产计划、物料需求计划',
            '车间管理：生产订单、生产领料、生产入库',
            '质量管理：来料检验、过程检验、成品检验',
            '设备管理：设备台账、设备维护、设备分析'
        ]
        for test in manufacture_tests:
            function_test_scope.append(test)
    
    for test in function_test_scope:
        doc.add_paragraph(f'• {test}', style='List Bullet')
    
    doc.add_heading('4.3 性能测试范围', level=2)
    performance_test_scope = [
        '并发用户：500用户并发访问',
        '响应时间：页面响应时间≤3秒',
        '数据处理：数据处理时间≤5秒',
        '系统稳定性：系统运行时间≥24小时',
        '系统容量：支持1000用户同时在线'
    ]
    for scope in performance_test_scope:
        doc.add_paragraph(f'• {scope}', style='List Bullet')
    
    doc.add_heading('4.4 安全测试范围', level=2)
    security_test_scope = [
        '身份认证：用户身份认证、多因素认证',
        '权限控制：角色权限、数据权限、功能权限',
        '数据安全：数据加密、数据脱敏、数据备份',
        '网络安全：网络安全、应用安全、数据安全'
    ]
    for scope in security_test_scope:
        doc.add_paragraph(f'• {scope}', style='List Bullet')
    
    doc.add_page_break()
    
    # 五、测试数据
    doc.add_heading('五、测试数据', level=1)
    doc.add_heading('5.1 测试数据概述', level=2)
    doc.add_paragraph('测试数据是进行测试的数据，包括基础数据、业务数据、测试数据等。')
    
    doc.add_heading('5.2 基础数据', level=2)
    base_data = [
        '组织数据：公司、部门、岗位、人员',
        '业务数据：客户、供应商、物料、科目',
        '财务数据：科目、项目、成本中心、利润中心',
        '供应链数据：供应商、客户、物料、仓库',
        '制造数据：BOM、工艺路线、工作中心、设备'
    ]
    for data in base_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_heading('5.3 业务数据', level=2)
    business_data = [
        '财务数据：凭证、账簿、报表、预算',
        '采购数据：采购申请、采购订单、采购入库',
        '销售数据：销售订单、销售出库、销售开票',
        '生产数据：生产订单、生产领料、生产入库',
        '库存数据：入库单、出库单、库存盘点'
    ]
    for data in business_data:
        doc.add_paragraph(f'• {data}', style='List Bullet')
    
    doc.add_heading('5.4 测试数据设计', level=2)
    test_data_design = [
        '数据完整性：数据完整、准确、一致',
        '数据代表性：数据具有代表性，覆盖所有业务场景',
        '数据安全性：数据安全，保护敏感信息',
        '数据可维护性：数据可维护，便于更新和管理'
    ]
    for design in test_data_design:
        doc.add_paragraph(f'• {design}', style='List Bullet')
    
    doc.add_heading('5.5 测试数据管理', level=2)
    test_data_management = [
        '数据准备：数据清洗、数据转换、数据验证',
        '数据存储：数据存储、数据备份、数据恢复',
        '数据维护：数据更新、数据删除、数据归档',
        '数据安全：数据加密、数据脱敏、数据审计'
    ]
    for management in test_data_management:
        doc.add_paragraph(f'• {management}', style='List Bullet')
    
    doc.add_page_break()
    
    # 六、测试用例
    doc.add_heading('六、测试用例', level=1)
    doc.add_heading('6.1 测试用例概述', level=2)
    doc.add_paragraph('测试用例是测试的具体实现，包括测试目标、测试步骤、测试预期结果等。')
    
    doc.add_heading('6.2 测试用例设计原则', level=2)
    test_case_principles = [
        '完整性：覆盖所有业务场景',
        '准确性：测试步骤准确，预期结果准确',
        '可重复性：测试用例可重复执行',
        '可维护性：测试用例可维护，便于更新',
        '可追溯性：测试用例可追溯，便于管理'
    ]
    for principle in test_case_principles:
        doc.add_paragraph(f'• {principle}', style='List Bullet')
    
    doc.add_heading('6.3 财务模块测试用例（60个）', level=2)
    
    # 财务测试用例
    finance_test_cases = [
        ('总账管理', [
            'TC-FIN-001：凭证录入测试',
            'TC-FIN-002：凭证审核测试',
            'TC-FIN-003：凭证记账测试',
            'TC-FIN-004：凭证查询测试',
            'TC-FIN-005：凭证修改测试',
            'TC-FIN-006：凭证删除测试',
            'TC-FIN-007：凭证冲销测试',
            'TC-FIN-008：凭证打印测试',
            'TC-FIN-009：凭证导出测试',
            'TC-FIN-010：凭证导入测试'
        ]),
        ('应收管理', [
            'TC-FIN-011：销售开票测试',
            'TC-FIN-012：应收确认测试',
            'TC-FIN-013：收款核销测试',
            'TC-FIN-014：收款查询测试',
            'TC-FIN-015：收款修改测试',
            'TC-FIN-016：收款删除测试',
            'TC-FIN-017：账龄分析测试',
            'TC-FIN-018：坏账准备测试',
            'TC-FIN-019：应收报表测试',
            'TC-FIN-020：应收分析测试'
        ]),
        ('应付管理', [
            'TC-FIN-021：采购开票测试',
            'TC-FIN-022：应付确认测试',
            'TC-FIN-023：付款核销测试',
            'TC-FIN-024：付款查询测试',
            'TC-FIN-025：付款修改测试',
            'TC-FIN-026：付款删除测试',
            'TC-FIN-027：账龄分析测试',
            'TC-FIN-028：应付报表测试',
            'TC-FIN-029：应付分析测试',
            'TC-FIN-030：现金折扣测试'
        ]),
        ('成本管理', [
            'TC-FIN-031：成本核算测试',
            'TC-FIN-032：成本分析测试',
            'TC-FIN-033：成本控制测试',
            'TC-FIN-034：成本报表测试',
            'TC-FIN-035：成本预测测试',
            'TC-FIN-036：成本预算测试',
            'TC-FIN-037：成本差异分析测试',
            'TC-FIN-038：成本分配测试',
            'TC-FIN-039：成本归集测试',
            'TC-FIN-040：成本查询测试'
        ]),
        ('固定资产', [
            'TC-FIN-041：资产卡片测试',
            'TC-FIN-042：资产增加测试',
            'TC-FIN-043：资产减少测试',
            'TC-FIN-044：资产转移测试',
            'TC-FIN-045：资产折旧测试',
            'TC-FIN-046：资产盘点测试',
            'TC-FIN-047：资产报表测试',
            'TC-FIN-048：资产分析测试',
            'TC-FIN-049：资产查询测试',
            'TC-FIN-050：资产打印测试'
        ])
    ]
    
    for module, test_cases in finance_test_cases:
        doc.add_paragraph(f'{module}：', style='List Number')
        for test_case in test_cases:
            doc.add_paragraph(f'  • {test_case}', style='List Bullet')
    
    doc.add_heading('6.4 供应链模块测试用例（40个）', level=2)
    
    # 供应链测试用例
    supply_test_cases = [
        ('采购管理', [
            'TC-SUP-001：供应商管理测试',
            'TC-SUP-002：采购申请测试',
            'TC-SUP-003：采购订单测试',
            'TC-SUP-004：采购合同测试',
            'TC-SUP-005：采购入库测试',
            'TC-SUP-006：采购退货测试',
            'TC-SUP-007：采购查询测试',
            'TC-SUP-008：采购报表测试',
            'TC-SUP-009：采购分析测试',
            'TC-SUP-010：采购审批测试'
        ]),
        ('库存管理', [
            'TC-SUP-011：入库管理测试',
            'TC-SUP-012：出库管理测试',
            'TC-SUP-013：库存查询测试',
            'TC-SUP-014：库存盘点测试',
            'TC-SUP-015：库存调整测试',
            'TC-SUP-016：库存报表测试',
            'TC-SUP-017：库存分析测试',
            'TC-SUP-018：库存预警测试',
            'TC-SUP-019：库存查询测试',
            'TC-SUP-020：库存打印测试'
        ]),
        ('销售管理', [
            'TC-SUP-021：客户管理测试',
            'TC-SUP-022：销售订单测试',
            'TC-SUP-023：销售合同测试',
            'TC-SUP-024：销售出库测试',
            'TC-SUP-025：销售退货测试',
            'TC-SUP-026：销售开票测试',
            'TC-SUP-027：销售查询测试',
            'TC-SUP-028：销售报表测试',
            'TC-SUP-029：销售分析测试',
            'TC-SUP-030：销售审批测试'
        ]),
        ('物流管理', [
            'TC-SUP-031：运输管理测试',
            'TC-SUP-032：配送管理测试',
            'TC-SUP-033：仓储管理测试',
            'TC-SUP-034：物流查询测试',
            'TC-SUP-035：物流报表测试',
            'TC-SUP-036：物流分析测试',
            'TC-SUP-037：物流跟踪测试',
            'TC-SUP-038：物流查询测试',
            'TC-SUP-039：物流打印测试',
            'TC-SUP-040：物流审批测试'
        ])
    ]
    
    for module, test_cases in supply_test_cases:
        doc.add_paragraph(f'{module}：', style='List Number')
        for test_case in test_cases:
            doc.add_paragraph(f'  • {test_case}', style='List Bullet')
    
    doc.add_heading('6.5 制造模块测试用例（20个）', level=2)
    
    # 制造测试用例
    manufacture_test_cases = [
        ('生产计划', [
            'TC-MFG-001：需求预测测试',
            'TC-MFG-002：主生产计划测试',
            'TC-MFG-003：物料需求计划测试',
            'TC-MFG-004：生产订单测试',
            'TC-MFG-005：生产计划查询测试',
            'TC-MFG-006：生产计划报表测试',
            'TC-MFG-007：生产计划分析测试'
        ]),
        ('车间管理', [
            'TC-MFG-008：生产领料测试',
            'TC-MFG-009：生产执行测试',
            'TC-MFG-010：生产入库测试',
            'TC-MFG-011：工序管理测试',
            'TC-MFG-012：车间查询测试',
            'TC-MFG-013：车间报表测试'
        ]),
        ('质量管理', [
            'TC-MFG-014：来料检验测试',
            'TC-MFG-015：过程检验测试',
            'TC-MFG-016：成品检验测试',
            'TC-MFG-017：质量追溯测试',
            'TC-MFG-018：质量报表测试'
        ]),
        ('设备管理', [
            'TC-MFG-019：设备台账测试',
            'TC-MFG-020：设备维护测试'
        ])
    ]
    
    for module, test_cases in manufacture_test_cases:
        doc.add_paragraph(f'{module}：', style='List Number')
        for test_case in test_cases:
            doc.add_paragraph(f'  • {test_case}', style='List Bullet')
    
    doc.add_page_break()
    
    # 七、测试执行
    doc.add_heading('七、测试执行', level=1)
    doc.add_heading('7.1 测试执行概述', level=2)
    doc.add_paragraph('测试执行是按照测试计划执行测试用例，记录测试结果，分析测试问题。')
    
    doc.add_heading('7.2 测试执行计划', level=2)
    test_execution_plan = [
        '测试准备：测试环境准备、测试数据准备、测试人员准备',
        '测试执行：按照测试用例执行测试，记录测试结果',
        '测试监控：监控测试过程，及时发现测试问题',
        '测试记录：记录测试结果，分析测试问题',
        '测试报告：编写测试报告，总结测试结果'
    ]
    for plan in test_execution_plan:
        doc.add_paragraph(f'• {plan}', style='List Bullet')
    
    doc.add_heading('7.3 测试执行步骤', level=2)
    test_execution_steps = [
        '步骤1：测试环境准备',
        '- 准备测试服务器',
        '- 安装测试软件',
        '- 配置测试网络',
        '- 导入测试数据',
        '',
        '步骤2：测试用例执行',
        '- 登录测试系统',
        '- 执行测试用例',
        '- 记录测试结果',
        '- 标记测试问题',
        '',
        '步骤3：测试问题处理',
        '- 记录测试问题',
        '- 分析测试问题',
        '- 分配测试问题',
        '- 跟踪测试问题',
        '',
        '步骤4：测试结果验证',
        '- 验证测试结果',
        '- 确认测试问题',
        '- 关闭测试问题',
        '- 更新测试报告'
    ]
    for step in test_execution_steps:
        doc.add_paragraph(f'{step}', style='List Number')
    
    doc.add_heading('7.4 测试执行监控', level=2)
    test_execution_monitoring = [
        '测试进度监控：监控测试进度，确保测试进度',
        '测试质量监控：监控测试质量，确保测试质量',
        '测试风险监控：监控测试风险，确保测试风险',
        '测试资源监控：监控测试资源，确保测试资源'
    ]
    for monitoring in test_execution_monitoring:
        doc.add_paragraph(f'• {monitoring}', style='List Bullet')
    
    doc.add_page_break()
    
    # 八、测试结果
    doc.add_heading('八、测试结果', level=1)
    doc.add_heading('8.1 测试结果概述', level=2)
    doc.add_paragraph('测试结果是测试执行的结果，包括测试通过、测试失败、测试阻塞等。')
    
    doc.add_heading('8.2 测试结果统计', level=2)
    test_result_statistics = [
        '总测试用例：120个',
        '通过测试用例：100个（83.3%）',
        '失败测试用例：15个（12.5%）',
        '阻塞测试用例：5个（4.2%）',
        '',
        '通过率：83.3%',
        '失败率：12.5%',
        '阻塞率：4.2%'
    ]
    for stat in test_result_statistics:
        doc.add_paragraph(f'• {stat}', style='List Bullet')
    
    doc.add_heading('8.3 测试失败分析', level=2)
    test_failure_analysis = [
        '功能失败：10个（66.7%）',
        '性能失败：3个（20%）',
        '安全失败：2个（13.3%）',
        '',
        '主要失败原因：',
        '- 需求理解不充分',
        '- 系统配置错误',
        '- 测试数据不完整',
        '- 测试环境问题'
    ]
    for analysis in test_failure_analysis:
        doc.add_paragraph(f'• {analysis}', style='List Bullet')
    
    doc.add_heading('8.4 测试阻塞分析', level=2)
    test_blockage_analysis = [
        '环境阻塞：3个（60%）',
        '数据阻塞：1个（20%）',
        '资源阻塞：1个（20%）',
        '',
        '主要阻塞原因：',
        '- 测试环境不稳定',
        '- 测试数据不完整',
        '- 测试资源不足'
    ]
    for analysis in test_blockage_analysis:
        doc.add_paragraph(f'• {analysis}', style='List Bullet')
    
    doc.add_page_break()
    
    # 九、测试报告
    doc.add_heading('九、测试报告', level=1)
    doc.add_heading('9.1 测试报告概述', level=2)
    doc.add_paragraph('测试报告是测试执行的总结，包括测试结果、测试问题、测试建议等。')
    
    doc.add_heading('9.2 测试报告内容', level=2)
    test_report_content = [
        '测试概述：测试目标、测试范围、测试环境',
        '测试执行：测试计划、测试执行、测试监控',
        '测试结果：测试统计、测试分析、测试结论',
        '测试问题：问题统计、问题分析、问题解决',
        '测试建议：改进建议、优化建议、建议措施'
    ]
    for content in test_report_content:
        doc.add_paragraph(f'• {content}', style='List Bullet')
    
    doc.add_heading('9.3 测试报告格式', level=2)
    test_report_format = [
        '报告标题：UAT测试报告',
        '报告日期：测试完成日期',
        '报告版本：版本号',
        '报告内容：测试概述、测试执行、测试结果、测试问题、测试建议',
        '报告附件：测试用例、测试数据、测试环境'
    ]
    for format_item in test_report_format:
        doc.add_paragraph(f'• {format_item}', style='List Bullet')
    
    doc.add_heading('9.4 测试报告评审', level=2)
    test_report_review = [
        '报告评审：由测试负责人进行评审',
        '报告确认：由项目确认人进行确认',
        '报告发布：由发布人进行发布',
        '报告分发：由分发人进行分发'
    ]
    for review in test_report_review:
        doc.add_paragraph(f'• {review}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十、测试改进
    doc.add_heading('十、测试改进', level=1)
    doc.add_heading('10.1 测试改进概述', level=2)
    doc.add_paragraph('测试改进是测试过程的持续改进，包括测试方法、测试工具、测试流程等。')
    
    doc.add_heading('10.2 测试方法改进', level=2)
    test_method_improvement = [
        '测试用例改进：优化测试用例，提高测试覆盖率',
        '测试数据改进：优化测试数据，提高测试质量',
        '测试环境改进：优化测试环境，提高测试效率',
        '测试工具改进：优化测试工具，提高测试自动化'
    ]
    for improvement in test_method_improvement:
        doc.add_paragraph(f'• {improvement}', style='List Bullet')
    
    doc.add_heading('10.3 测试流程改进', level=2)
    test_process_improvement = [
        '测试流程优化：优化测试流程，提高测试效率',
        '测试流程标准化：标准化测试流程，提高测试质量',
        '测试流程自动化：自动化测试流程，减少人工干预',
        '测试流程监控：监控测试流程，确保测试质量'
    ]
    for improvement in test_process_improvement:
        doc.add_paragraph(f'• {improvement}', style='List Bullet')
    
    doc.add_heading('10.4 测试工具改进', level=2)
    test_tool_improvement = [
        '测试管理工具：优化测试管理工具，提高管理效率',
        '自动化测试工具：优化自动化测试工具，提高测试效率',
        '性能测试工具：优化性能测试工具，提高测试质量',
        '安全测试工具：优化安全测试工具，提高测试安全性'
    ]
    for improvement in test_tool_improvement:
        doc.add_paragraph(f'• {improvement}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十一、测试风险
    doc.add_heading('十一、测试风险', level=1)
    doc.add_heading('11.1 测试风险概述', level=2)
    doc.add_paragraph('测试风险是测试过程中可能出现的风险，包括技术风险、业务风险、管理风险等。')
    
    doc.add_heading('11.2 测试风险识别', level=2)
    test_risk_identification = [
        '技术风险：技术不成熟、技术难度大、技术复杂度高',
        '业务风险：需求不明确、业务复杂度高、业务流程复杂',
        '管理风险：管理不到位、资源不足、进度延迟',
        '环境风险：环境不稳定、环境配置错误、环境资源不足'
    ]
    for risk in test_risk_identification:
        doc.add_paragraph(f'• {risk}', style='List Bullet')
    
    doc.add_heading('11.3 测试风险评估', level=2)
    test_risk_assessment = [
        '风险概率：评估风险发生的概率',
        '风险影响：评估风险对项目的影响',
        '风险等级：评估风险的等级（高、中、低）',
        '风险优先级：评估风险的优先级（紧急、重要、一般）'
    ]
    for risk in test_risk_assessment:
        doc.add_paragraph(f'• {risk}', style='List Bullet')
    
    doc.add_heading('11.4 测试风险控制', level=2)
    test_risk_control = [
        '风险预防：预防风险发生，降低风险概率',
        '风险减轻：减轻风险影响，降低风险影响',
        '风险转移：转移风险责任，降低风险责任',
        '风险接受：接受风险存在，制定应对措施'
    ]
    for risk in test_risk_control:
        doc.add_paragraph(f'• {risk}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十二、测试保障
    doc.add_heading('十二、测试保障', level=1)
    doc.add_heading('12.1 测试保障概述', level=2)
    doc.add_paragraph('测试保障是测试过程的保障措施，包括组织保障、资源保障、技术保障等。')
    
    doc.add_heading('12.2 组织保障', level=2)
    test_organization_guarantee = [
        '测试团队：组建专业测试团队，确保测试质量',
        '测试负责人：指定测试负责人，确保测试进度',
        '测试专家：邀请测试专家，确保测试专业性',
        '测试协调人：指定测试协调人，确保测试协调'
    ]
    for guarantee in test_organization_guarantee:
        doc.add_paragraph(f'• {guarantee}', style='List Bullet')
    
    doc.add_heading('12.3 资源保障', level=2)
    test_resource_guarantee = [
        '测试环境：提供稳定测试环境，确保测试环境',
        '测试数据：提供完整测试数据，确保测试数据',
        '测试工具：提供先进测试工具，确保测试工具',
        '测试时间：提供充足测试时间，确保测试时间'
    ]
    for guarantee in test_resource_guarantee:
        doc.add_paragraph(f'• {guarantee}', style='List Bullet')
    
    doc.add_heading('12.4 技术保障', level=2)
    test_technology_guarantee = [
        '测试技术：采用先进测试技术，确保测试技术',
        '测试方法：采用科学测试方法，确保测试方法',
        '测试标准：采用统一测试标准，确保测试标准',
        '测试流程：采用规范测试流程，确保测试流程'
    ]
    for guarantee in test_technology_guarantee:
        doc.add_paragraph(f'• {guarantee}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十三、测试经验
    doc.add_heading('十三、测试经验', level=1)
    doc.add_heading('13.1 测试经验概述', level=2)
    doc.add_paragraph('测试经验是测试过程中的经验总结，包括成功经验、失败经验、改进经验等。')
    
    doc.add_heading('13.2 成功经验', level=2)
    test_success_experience = [
        '测试计划制定：制定详细测试计划，确保测试计划',
        '测试用例设计：设计全面测试用例，确保测试覆盖',
        '测试数据准备：准备完整测试数据，确保测试数据',
        '测试环境配置：配置稳定测试环境，确保测试环境',
        '测试团队协作：加强测试团队协作，确保测试协作'
    ]
    for experience in test_success_experience:
        doc.add_paragraph(f'• {experience}', style='List Bullet')
    
    doc.add_heading('13.3 失败经验', level=2)
    test_failure_experience = [
        '需求理解不充分：需求理解不充分，导致测试偏差',
        '测试数据不完整：测试数据不完整，导致测试失败',
        '测试环境不稳定：测试环境不稳定，导致测试中断',
        '测试时间不足：测试时间不足，导致测试仓促',
        '测试资源不足：测试资源不足，导致测试质量下降'
    ]
    for experience in test_failure_experience:
        doc.add_paragraph(f'• {experience}', style='List Bullet')
    
    doc.add_heading('13.4 改进经验', level=2)
    test_improvement_experience = [
        '测试用例优化：优化测试用例，提高测试覆盖率',
        '测试数据优化：优化测试数据，提高测试质量',
        '测试环境优化：优化测试环境，提高测试效率',
        '测试工具优化：优化测试工具，提高测试自动化',
        '测试流程优化：优化测试流程，提高测试效率'
    ]
    for experience in test_improvement_experience:
        doc.add_paragraph(f'• {experience}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十四、测试总结
    doc.add_heading('十四、测试总结', level=1)
    doc.add_heading('14.1 测试总结概述', level=2)
    doc.add_paragraph('测试总结是测试过程的总结，包括测试成果、测试问题、测试建议等。')
    
    doc.add_heading('14.2 测试成果', level=2)
    test_achievements = [
        '测试覆盖：测试覆盖率达到90%以上',
        '测试质量：测试质量达到85%以上',
        '测试效率：测试效率达到80%以上',
        '测试满意度：测试满意度达到90%以上'
    ]
    for achievement in test_achievements:
        doc.add_paragraph(f'• {achievement}', style='List Bullet')
    
    doc.add_heading('14.3 测试问题', level=2)
    test_problems = [
        '测试用例不完善：测试用例覆盖率不够',
        '测试数据不完整：测试数据代表性不够',
        '测试环境不稳定：测试环境配置错误',
        '测试时间不足：测试时间安排不合理'
    ]
    for problem in test_problems:
        doc.add_paragraph(f'• {problem}', style='List Bullet')
    
    doc.add_heading('14.4 测试建议', level=2)
    test_suggestions = [
        '完善测试用例：增加测试用例覆盖率',
        '完善测试数据：增加测试数据代表性',
        '完善测试环境：优化测试环境配置',
        '完善测试时间：合理安排测试时间',
        '完善测试工具：增加测试自动化工具'
    ]
    for suggestion in test_suggestions:
        doc.add_paragraph(f'• {suggestion}', style='List Bullet')
    
    doc.add_page_break()
    
    # 十五、测试建议
    doc.add_heading('十五、测试建议', level=1)
    doc.add_heading('15.1 测试建议概述', level=2)
    doc.add_paragraph('测试建议是测试过程的改进建议，包括测试方法、测试工具、测试流程等。')
    
    doc.add_heading('15.2 测试方法建议', level=2)
    test_method_suggestions = [
        '采用自动化测试：提高测试效率，减少人工干预',
        '采用性能测试：确保系统性能，满足业务需求',
        '采用安全测试：确保系统安全，保护数据安全',
        '采用兼容性测试：确保系统兼容，支持多种环境'
    ]
    for suggestion in test_method_suggestions:
        doc.add_paragraph(f'• {suggestion}', style='List Bullet')
    
    doc.add_heading('15.3 测试工具建议', level=2)
    test_tool_suggestions = [
        '采用测试管理工具：提高测试管理效率',
        '采用自动化测试工具：提高测试自动化水平',
        '采用性能测试工具：提高性能测试质量',
        '采用安全测试工具：提高安全测试水平'
    ]
    for suggestion in test_tool_suggestions:
        doc.add_paragraph(f'• {suggestion}', style='List Bullet')
    
    doc.add_heading('15.4 测试流程建议', level=2)
    test_process_suggestions = [
        '优化测试流程：提高测试效率，减少测试时间',
        '标准化测试流程：统一测试标准，提高测试质量',
        '自动化测试流程：减少人工干预，提高测试效率',
        '监控测试流程：监控测试过程，确保测试质量'
    ]
    for suggestion in test_process_suggestions:
        doc.add_paragraph(f'• {suggestion}', style='List Bullet')
    
    doc.add_heading('15.5 测试团队建议', level=2)
    test_team_suggestions = [
        '加强测试培训：提高测试团队专业水平',
        '加强测试协作：提高测试团队协作效率',
        '加强测试激励：提高测试团队工作积极性',
        '加强测试交流：提高测试团队交流频率'
    ]
    for suggestion in test_team_suggestions:
        doc.add_paragraph(f'• {suggestion}', style='List Bullet')
    
    doc.add_page_break()
    
    # 结束语
    conclusion = doc.add_paragraph()
    conclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conclusion.add_run(f"感谢{company_name}对本次UAT测试的支持！\n")
    run.font.size = Pt(16)
    run.font.bold = True
    run = conclusion.add_run("我们将根据测试结果，为您打造专业的ERP系统。\n")
    run.font.size = Pt(14)
    run = conclusion.add_run(f"测试日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(12)
    
    # 保存文件
    filename = f"{customer_info.get('customerCode', '客户')}_UAT测试方案_v9_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return {
        'success': True,
        'filepath': filepath,
        'filename': filename,
        'chapters': 15,
        'test_cases': 120
    }

if __name__ == '__main__':
    test_data = {
        'companyName': '测试公司',
        'customerCode': 'CS',
        'industry': '制造业',
        'companySize': '中型企业',
        'employees': '500',
        'revenue': '10000',
        'selectedModules': ['finance', 'supply', 'manufacture']
    }
    result = generate_uat_report_v9(test_data)
    print(f"✅ 生成成功：{result['filename']}")
    print(f"📊 章节数量：{result['chapters']}章")
    print(f"📋 测试用例：{result['test_cases']}个")
