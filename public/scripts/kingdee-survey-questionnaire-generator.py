#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶调研问卷生成器 - ZIP打包版
每个子系统生成一个文件，打包成ZIP：项目简称_调研提纲_日期.zip
"""

import os
import sys
import json
import shutil
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt

# 模板目录
TEMPLATE_DIR = "/mnt/d/Kingdee文档/自动化交付工具/金蝶AI星空（模板）/调研提纲"
OUTPUT_DIR = os.path.expanduser("~/.openclaw/workspace/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 子模块到模板文件的映射
SUBMODULE_TEMPLATE_MAP = {
    'gl': 'MKY_FD_调研提纲_总账_20231219_V1.0.docx',
    'smart_accounting': 'MKY_FD_调研提纲_总账_20231219_V1.0.docx',
    'ar': 'MKY_FD_调研提纲_应收管理_20231219_V1.0.docx',
    'ap': 'MKY_FD_调研提纲_应付管理_20231219_V1.0.docx',
    'fa': 'MKY_FD_调研提纲_费用报销_20231214_V1.0.docx',
    'cashier': 'MKY_FD_调研提纲__出纳+付款排程+票据_20231219_V1.0.docx',
    'fin_report': 'MKY_FD_调研提纲_总账_20231219_V1.0.docx',
    'expense_mgmt': 'MKY_FD_调研提纲_费用报销_20231214_V1.0.docx',
    'shared_finance': '星瀚V7.0交付工具包_财务云_共享任务中心_调研问卷.docx',
    'expense_everyone': '星瀚V7.0交付工具包_财务云_费用管理_调研问卷.docx',
    'bank_connect': '星瀚V7.0交付工具包_资金云_银企服务云&银企互联_操练手册.docx',
    'e_archive': 'MKY_FD_调研提纲_总账_20231219_V1.0.docx',
    'smart_audit': '星瀚V7.0交付工具包_财务云_财务报账_调研问卷.docx',
    'reconcile': '星瀚V7.0交付工具包_财务云_共享运营管理_调研问卷.docx',
    'inv_accounting': '星瀚V7.0交付工具包_财务云_存货核算_调研问卷.docx',
    'pur_mgmt': '星瀚V7.0交付工具包_供应链云_采购管理_调研问卷(4).docx',
    'sales_mgmt': '星瀚V7.0交付工具包_供应链云_销售管理_调研问卷.docx',
    'inv_mgmt': '星瀚V7.0交付工具包_供应链云_库存管理_调研问卷.docx',
    'supply_platform': '星瀚V7.0交付工具包_供应链云_采购管理_调研问卷(4).docx',
    'credit_mgmt': '星瀚V7.0交付工具包_供应链云_信用管理_调研问卷.docx',
    'qc_mgmt': '星瀚V7.0交付工具包_供应链云_质检任务中心_调研问卷.docx',
    'outsource_mgmt': '星瀚制造交付工具包_调研问卷_产品委外.docx',
    'barcode_mgmt': '星瀚V7.0交付工具包_供应链云_库存管理_调研问卷.docx',
    'budget': '星瀚V7.0交付工具包_绩效云_全面预算_调研问卷.docx',
    'merge_report': '星瀚V7.0交付工具包_绩效云_合并报表__调研问卷.docx',
    'fin_merge': '星瀚V7.0交付工具包_绩效云_合并报表__调研问卷.docx',
    'mgmt_report': '星瀚V7.0交付工具包_绩效云_全面预算_调研问卷.docx',
    'plan_analysis': '星瀚V7.0交付工具包_绩效云_全面预算_调研问卷.docx',
    'fin_analysis': '星瀚V7.0交付工具包_绩效云_全面预算_调研问卷.docx',
    'biz_analysis': '星瀚V7.0交付工具包_绩效云_全面预算_调研问卷.docx',
    'enterprise_report': '星瀚V7.0交付工具包_绩效云_全面预算_调研问卷.docx',
    'settlement': '星瀚V7.0交付工具包_司库云__账户管理&资金结算_调研问卷.docx',
    'fund_plan': '星瀚V7.0交付工具包_司库云__资金计划_调研问卷.docx',
    'bill_mgmt': '星瀚V7.0交付工具包_司库云_票据管理_调研问卷.docx',
    'lc_mgmt': '星瀚V7.0交付工具包_司库云_授信管理_调研问卷.docx',
    'internal_bank': '星瀚V7.0交付工具包_司库云_贷款管理_调研问卷.docx',
    'financing': '星瀚V7.0交付工具包_司库云_融资租赁_调研问卷.docx',
    'fx_mgmt': '星瀚V7.0交付工具包_司库云_投资理财_调研问卷.docx',
    'fund_forecast': '星瀚V7.0交付工具包_司库云__资金计划_调研问卷.docx',
    'fund_monitor': '星瀚V7.0交付工具包_司库云__资金计划_调研问卷.docx',
    'fund_analysis': '星瀚V7.0交付工具包_司库云__资金计划_调研问卷.docx',
    'bank_connect_treasury': '星瀚V7.0交付工具包_资金云_银企服务云&银企互联_操练手册.docx',
    'financial_relation': '星瀚V7.0交付工具包_司库云_授信管理_调研问卷.docx',
    'cost_mgmt': '星瀚V7.0交付工具包_管理会计云_实际成本_调研问卷.docx',
    'profit_analysis': '星瀚V7.0交付工具包_管理会计云_盈利能力分析_调研问卷.docx',
    'responsibility_acct': '星瀚V7.0交付工具包_管理会计云_标准成本_调研问卷.docx',
    'cvp_analysis': '星瀚V7.0交付工具包_管理会计云_盈利能力分析_调研问卷.docx',
    'ma_report': '星瀚V7.0交付工具包_管理会计云_盈利能力分析_调研问卷.docx',
    'transfer_price': '星瀚V7.0交付工具包_管理会计云_标准成本_调研问卷.docx',
    'accounting_engine': '星瀚V7.0交付工具包_管理会计云_标准成本_调研问卷.docx',
    'production_mgmt': '星瀚制造交付工具包_调研问卷_生产管理.docx',
    'workshop_mgmt': '星瀚制造交付工具包_调研问卷_生产管理.docx',
    'smart_scheduling': '星瀚制造交付工具包_调研问卷_计划管理.docx',
    'plan_mgmt': '星瀚制造交付工具包_调研问卷_计划管理.docx',
    'manufacturing_bigdata': '星瀚制造交付工具包_调研问卷_质量管理.docx',
    'quality_mgmt': '星瀚制造交付工具包_调研问卷_质量管理.docx',
    'equipment_mgmt': '星瀚制造交付工具包_调研问卷_质量管理.docx',
    'energy_mgmt': '星瀚制造交付工具包_调研问卷_质量管理.docx',
    'safety_mgmt': '星瀚制造交付工具包_调研问卷_高层管理.docx',
    'smart_factory': '星瀚制造交付工具包_调研问卷_生产管理.docx',
    'pdm_ecm': '星瀚制造交付工具包_调研问卷_制造产品数据管理.docx',
    'supplier_service': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'srm': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'source_mgmt': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'procurement_mall': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'pur_collab': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'rfq_mgmt': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'contract_collab': '星瀚V7.0交付工具包_供应链云_合同管理_合同智能提取_调研问卷.docx',
    'demand_plan': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'bid_mgmt': '星瀚V7.0交付工具包_供应商协同云_调研问卷.docx',
    'invoice_mgmt': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'smart_tax_calc': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'one_click_decl': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'tax_risk': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'tax_accounting': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'tax_shared': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'tax_archive': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'tax_regulation': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'export_rebate': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'deferred_tax': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'vat_mgmt': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'behavior_tax': '星瀚V7.0交付工具包_税务云_调研问卷.docx',
    'project_reg': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_budget': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_execute': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'contract_mgmt': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_cost': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_progress': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_accept': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_accounting': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_analysis': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_asset': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'project_evaluation': '星瀚V7.0交付工具包_项目云_调研问卷.docx',
    'channel_dist': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'ec_center': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'retail_mgmt': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'member_mgmt': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'crm': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'b2b_order': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'marketing_expense': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'price_rebate': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'marketing_analysis': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
    'channel_cloud_srv': '星瀚V7.0交付工具包_全渠道云_调研问卷.docx',
}

SUBMODULE_NAMES = {
    'gl': '总账', 'smart_accounting': '智能核算', 'ar': '应收管理', 'ap': '应付管理',
    'fa': '固定资产', 'cashier': '出纳管理', 'fin_report': '财务报表',
    'expense_mgmt': '费用管理', 'shared_finance': '财务共享', 'expense_everyone': '人人费用',
    'bank_connect': '银企互联', 'e_archive': '电子会计档案', 'smart_audit': '智能审单',
    'reconcile': '对账平台', 'inv_accounting': '存货核算',
    'pur_mgmt': '采购管理', 'sales_mgmt': '销售管理', 'inv_mgmt': '库存管理',
    'supply_platform': '供应链中台', 'credit_mgmt': '信用管理', 'qc_mgmt': '质检管理',
    'outsource_mgmt': '委外管理', 'barcode_mgmt': '条码管理',
    'budget': '全面预算', 'merge_report': '合并报表', 'fin_merge': '财务合并',
    'mgmt_report': '管理报表', 'plan_analysis': '计划与分析', 'fin_analysis': '财务分析',
    'biz_analysis': '经营分析', 'enterprise_report': '企业报表平台',
    'settlement': '资金结算', 'fund_plan': '资金计划', 'bill_mgmt': '票据管理',
    'lc_mgmt': '信用证管理', 'internal_bank': '内部银行', 'financing': '融资管理',
    'fx_mgmt': '外汇管理', 'fund_forecast': '资金预测', 'fund_monitor': '资金监控',
    'fund_analysis': '资金分析', 'bank_connect_treasury': '银企互联', 'financial_relation': '金融机构关系管理',
    'cost_mgmt': '成本管理', 'profit_analysis': '盈利能力分析', 'responsibility_acct': '责任会计',
    'cvp_analysis': '本量利分析', 'ma_report': '管理报表', 'transfer_price': '内部定价', 'accounting_engine': '会计引擎',
    'production_mgmt': '生产管理', 'workshop_mgmt': '车间管理', 'smart_scheduling': '智慧排程',
    'plan_mgmt': '计划管理', 'manufacturing_bigdata': '制造大数据', 'quality_mgmt': '质量管理',
    'equipment_mgmt': '设备管理', 'energy_mgmt': '能源管理', 'safety_mgmt': '安全生产',
    'smart_factory': '智能工厂', 'pdm_ecm': '产品数据与工程变更管理', 'supplier_service': '供应商协同服务',
    'srm': '供应商管理', 'source_mgmt': '寻源管理', 'procurement_mall': '采购商城',
    'pur_collab': '采购协同', 'rfq_mgmt': '询报价管理', 'contract_collab': '合同协同',
    'demand_plan': '需求计划', 'bid_mgmt': '招投标管理',
    'invoice_mgmt': '发票管理', 'smart_tax_calc': '智能算税', 'one_click_decl': '一键申报',
    'tax_risk': '税务风险管理', 'tax_accounting': '税务会计', 'tax_shared': '税务共享',
    'tax_archive': '税收档案管理', 'tax_regulation': '法规库', 'export_rebate': '出口退税',
    'deferred_tax': '递延所得税管理', 'vat_mgmt': '增值税管理', 'behavior_tax': '财务行为税管理',
    'project_reg': '项目立项', 'project_budget': '项目预算', 'project_execute': '项目执行',
    'contract_mgmt': '合同管理', 'project_cost': '项目成本', 'project_progress': '项目进度',
    'project_accept': '项目验收', 'project_accounting': '项目核算', 'project_analysis': '项目分析',
    'project_asset': '项目资产', 'project_evaluation': '项目评价',
    'channel_dist': '渠道分销管理', 'ec_center': '电商中心', 'retail_mgmt': '零售管理',
    'member_mgmt': '会员管理', 'crm': 'CRM', 'b2b_order': 'B2B订单中心',
    'marketing_expense': '营销费用管理', 'price_rebate': '价格与返利', 'marketing_analysis': '营销分析',
    'channel_cloud_srv': '渠道云服务',
}


def add_logo(doc, logo_path=None):
    if logo_path and os.path.exists(logo_path):
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            for para in header.paragraphs:
                if '【客户LOGO】' in para.text or not para.text.strip():
                    para.text = ''
                    run = para.add_run()
                    run.add_picture(logo_path, width=Inches(1.5))
                    break
    else:
        for section in doc.sections:
            for para in section.header.paragraphs:
                if '【客户LOGO】' in para.text:
                    para.text = ''
                    break


def generate_survey_zip(customer_info, selected_submodules, logo_path=None):
    short_name = customer_info.get('shortName', '项目')
    date_str = datetime.now().strftime('%Y%m%d')
    zip_filename = f"{short_name}_调研提纲_{date_str}.zip"
    zip_path = os.path.join(OUTPUT_DIR, zip_filename)

    # 创建临时目录存放docx文件
    temp_dir = os.path.join(OUTPUT_DIR, f"temp_{date_str}")
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    docx_count = 0
    for submodule_id in selected_submodules:
        template_file = SUBMODULE_TEMPLATE_MAP.get(submodule_id)
        if not template_file:
            continue
        template_path = os.path.join(TEMPLATE_DIR, template_file)
        if not os.path.exists(template_path):
            continue
        submodule_name = SUBMODULE_NAMES.get(submodule_id, submodule_id)
        docx_filename = f"{short_name}_{submodule_name}.docx"
        docx_path = os.path.join(temp_dir, docx_filename)
        try:
            shutil.copy(template_path, docx_path)
            doc = Document(docx_path)
            company_name = customer_info.get('companyName', 'XX公司')
            for para in doc.paragraphs:
                for run in para.runs:
                    if 'XX集团' in run.text:
                        run.text = run.text.replace('XX集团', company_name)
            add_logo(doc, logo_path)
            doc.save(docx_path)
            docx_count += 1
        except Exception:
            continue

    # 创建ZIP文件
    if docx_count > 0:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for f in os.listdir(temp_dir):
                if f.endswith('.docx'):
                    zf.write(os.path.join(temp_dir, f), f)
        # 清理临时目录
        shutil.rmtree(temp_dir)
        return zip_filename
    else:
        shutil.rmtree(temp_dir)
        return None


def main():
    if len(sys.argv) < 2:
        print(json.dumps({'error': '参数不足'}))
        sys.exit(1)
    try:
        data = json.loads(sys.argv[1])
        customer_info = data.get('customerInfo', {})
        selected_submodules = data.get('selectedSubModules', [])
        logo_path = data.get('logoPath')
        if not selected_submodules:
            print(json.dumps({'error': '未选择子模块'}))
            sys.exit(1)
        zip_file = generate_survey_zip(customer_info, selected_submodules, logo_path)
        if zip_file:
            print(json.dumps({'success': True, 'filename': zip_file}))
        else:
            print(json.dumps({'error': '生成失败'}))
    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
