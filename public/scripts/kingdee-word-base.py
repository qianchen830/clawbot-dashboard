#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
金蝶Word文档基础类 v1.0
统一所有Word生成器的格式和样式

包含功能：
1. 文档初始化（页边距、字体、样式）
2. 封面页生成
3. 文档控制页（修订记录）
4. 目录页
5. 页眉页脚设置
6. 标题样式统一
7. 表格样式美化
8. 段落样式设置
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ==================== 主题配色 ====================
THEME = {
    'primary': RGBColor(0x1F, 0x4E, 0x79),      # 深蓝色（金蝶主色）
    'secondary': RGBColor(0x2E, 0x75, 0xB6),    # 中蓝色
    'accent': RGBColor(0xFF, 0x99, 0x00),       # 橙色强调
    'dark': RGBColor(0x33, 0x33, 0x33),         # 深灰文字
    'medium': RGBColor(0x66, 0x66, 0x66),       # 中灰
    'light': RGBColor(0xF2, 0xF2, 0xF2),        # 浅灰背景
    'white': RGBColor(0xFF, 0xFF, 0xFF),        # 白色
    'header_bg': RGBColor(0x1F, 0x4E, 0x79),    # 表头背景色
    'row_alt': RGBColor(0xE6, 0xF0, 0xFA),      # 交替行色（浅蓝）
    'border': RGBColor(0x8E, 0xB4, 0xD7),       # 边框色
}


# ==================== 中文数字映射 ====================
CN_NUMS = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
           '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']


class KingdeeWordBase:
    """金蝶Word文档基础类"""
    
    def __init__(self, company, doc_type="文档"):
        """
        初始化文档
        
        Args:
            company: 公司名称
            doc_type: 文档类型（如"业务蓝图"、"调研报告"等）
        """
        self.company = company
        self.doc_type = doc_type
        self.doc = Document()
        self.project = f"{company}新ERP管理系统项目"
        self.created_date = datetime.now().strftime("%Y-%m-%d")
        
        # 初始化文档设置
        self._setup_page()
        self._setup_styles()
        
    def _setup_page(self):
        """设置页面格式"""
        for section in self.doc.sections:
            # 页边距：上2.0cm，下1.0cm，左右2.0cm
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            
            # 纸张大小A4
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        
        # 设置标题样式
        self._setup_heading_style(1, Pt(22), True)  # 一级标题：22pt，加粗
        self._setup_heading_style(2, Pt(18), True)  # 二级标题：18pt，加粗
        self._setup_heading_style(3, Pt(14), True)  # 三级标题：14pt，加粗
        self._setup_heading_style(4, Pt(12), True)  # 四级标题：12pt，加粗
    
    def _setup_heading_style(self, level, size, bold):
        """设置标题样式"""
        try:
            style = self.doc.styles[f'Heading {level}']
        except:
            style = self.doc.styles.add_style(f'Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
        
        style.font.name = '微软雅黑'
        style.font.size = size
        style.font.bold = bold
        style.font.color.rgb = THEME['primary']
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    
    # ==================== 封面页 ====================
    
    def add_cover(self, title=None, subtitle=None, version="V1.0", 
                  author="金蝶软件（中国）有限公司", date=None):
        """
        添加封面页
        
        Args:
            title: 主标题（默认使用文档类型）
            subtitle: 副标题（默认使用项目名）
            version: 版本号
            author: 编制单位
            date: 日期（默认今天）
        """
        title = title or self.doc_type
        subtitle = subtitle or self.project
        date = date or self.created_date
        
        # 添加空行撑高
        for _ in range(6):
            self.doc.add_paragraph()
        
        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = '微软雅黑'
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = THEME['primary']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 副标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.name = '微软雅黑'
        run.font.size = Pt(20)
        run.font.color.rgb = THEME['secondary']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 空行
        for _ in range(8):
            self.doc.add_paragraph()
        
        # 信息表格
        table = self.doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        info_items = [
            ('编制单位：', author),
            ('编制日期：', date),
            ('文档版本：', version),
            ('文档密级：', '内部资料'),
        ]
        
        for i, (label, value) in enumerate(info_items):
            row = table.rows[i]
            # 标签列
            cell = row.cells[0]
            cell.width = Cm(4)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(label)
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 值列
            cell = row.cells[1]
            cell.width = Cm(8)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 分页
        self.doc.add_page_break()
    
    # ==================== 文档控制页 ====================
    
    def add_document_control(self, revisions=None):
        """
        添加文档控制页（修订记录）
        
        Args:
            revisions: 修订记录列表，格式：
                [
                    {'version': 'V1.0', 'date': '2024-01-01', 'author': '张三', 'changes': '初稿'},
                    {'version': 'V1.1', 'date': '2024-01-15', 'author': '李四', 'changes': '修订财务章节'},
                ]
        """
        self.add_heading('文档控制', level=1)
        
        # 文档信息
        self.add_heading('文档信息', level=2)
        info_table = self.add_table(4, 4)
        self._set_table_cell(info_table, 0, 0, '文档名称')
        self._set_table_cell(info_table, 0, 1, self.doc_type)
        self._set_table_cell(info_table, 0, 2, '文档编号')
        self._set_table_cell(info_table, 0, 3, 'KD-{0}-{1}'.format(
            datetime.now().strftime('%Y%m'), self.doc_type[:2].upper()))
        
        self._set_table_cell(info_table, 1, 0, '项目名称')
        self._set_table_cell(info_table, 1, 1, self.project)
        self._set_table_cell(info_table, 1, 2, '版本号')
        self._set_table_cell(info_table, 1, 3, 'V1.0')
        
        self._set_table_cell(info_table, 2, 0, '编制单位')
        self._set_table_cell(info_table, 2, 1, '金蝶软件（中国）有限公司')
        self._set_table_cell(info_table, 2, 2, '编制日期')
        self._set_table_cell(info_table, 2, 3, self.created_date)
        
        self._set_table_cell(info_table, 3, 0, '批准人')
        self._set_table_cell(info_table, 3, 1, '')
        self._set_table_cell(info_table, 3, 2, '批准日期')
        self._set_table_cell(info_table, 3, 3, '')
        
        self._format_table(info_table, header_row=False)
        
        # 修订记录
        self.add_heading('修订记录', level=2)
        
        if not revisions:
            revisions = [
                {'version': 'V1.0', 'date': self.created_date, 'author': '编制人', 'changes': '初稿'}
            ]
        
        rev_table = self.add_table(len(revisions) + 1, 4)
        headers = ['版本', '日期', '修订人', '修订内容']
        for i, h in enumerate(headers):
            self._set_table_cell(rev_table, 0, i, h, is_header=True)
        
        for i, rev in enumerate(revisions):
            self._set_table_cell(rev_table, i+1, 0, rev.get('version', ''))
            self._set_table_cell(rev_table, i+1, 1, rev.get('date', ''))
            self._set_table_cell(rev_table, i+1, 2, rev.get('author', ''))
            self._set_table_cell(rev_table, i+1, 3, rev.get('changes', ''))
        
        self._format_table(rev_table)
        
        # 分页
        self.doc.add_page_break()
    
    # ==================== 目录页 ====================
    
    def add_toc(self):
        """添加目录页（需要Word手动刷新）"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('目  录')
        run.font.name = '微软雅黑'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = THEME['primary']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加目录域
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        run = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        run = paragraph.add_run('（请右键更新目录）')
        run.font.color.rgb = THEME['medium']
        
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        
        # 分页
        self.doc.add_page_break()
    
    # ==================== 页眉页脚 ====================
    
    def set_header_footer(self):
        """设置页眉页脚"""
        for section in self.doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = header_para.add_run(f'{self.project} {self.doc_type}')
            run.font.name = '微软雅黑'
            run.font.size = Pt(9)
            run.font.color.rgb = THEME['medium']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 添加页眉下划线
            pPr = header_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '1F4E79')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码域
            run = footer_para.add_run('第 ')
            run.font.name = '微软雅黑'
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 页码
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run = footer_para.add_run()
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            run = footer_para.add_run(' 页 共 ')
            run.font.name = '微软雅黑'
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 总页数
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'NUMPAGES'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run = footer_para.add_run()
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            run = footer_para.add_run(' 页    内部资料，注意保密')
            run.font.name = '微软雅黑'
            run.font.size = Pt(9)
            run.font.color.rgb = THEME['medium']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ==================== 标题方法 ====================
    
    def add_heading(self, text, level=1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        return heading
    
    def add_paragraph(self, text, bold=False, indent=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """
        添加段落
        
        Args:
            text: 段落文本
            bold: 是否加粗
            indent: 首行缩进（字符数）
            alignment: 对齐方式
        """
        p = self.doc.add_paragraph()
        p.alignment = alignment
        
        if indent > 0:
            p.paragraph_format.first_line_indent = Cm(indent * 0.75)  # 约2字符
        
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        return p
    
    # ==================== 表格方法 ====================
    
    def add_table(self, rows, cols):
        """添加表格"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table
    
    def _set_table_cell(self, table, row, col, text, is_header=False, width=None):
        """设置表格单元格内容"""
        cell = table.rows[row].cells[col]
        cell.text = ''
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10.5)
        run.font.bold = is_header
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        if is_header:
            run.font.color.rgb = THEME['white']
        
        if width:
            cell.width = Cm(width)
    
    def _format_table(self, table, header_row=True, alternate_rows=True):
        """
        格式化表格样式
        
        Args:
            table: 表格对象
            header_row: 是否有表头行
            alternate_rows: 是否使用交替行色
        """
        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '8EB4D7')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        
        # 格式化每一行
        for i, row in enumerate(table.rows):
            # 设置行高
            row.height = Cm(0.8)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            for j, cell in enumerate(row.cells):
                # 设置垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 表头行背景色
                if header_row and i == 0:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '1F4E79')
                    cell._tc.get_or_add_tcPr().append(shading)
                # 交替行背景色
                elif alternate_rows and i % 2 == 0:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E6F0FA')
                    cell._tc.get_or_add_tcPr().append(shading)
    
    def add_styled_table(self, headers, data, col_widths=None):
        """
        添加格式化表格
        
        Args:
            headers: 表头列表
            data: 数据列表（二维数组）
            col_widths: 列宽列表（厘米）
        
        Returns:
            table: 表格对象
        """
        rows = len(data) + 1
        cols = len(headers)
        table = self.add_table(rows, cols)
        
        # 设置表头
        for i, h in enumerate(headers):
            self._set_table_cell(table, 0, i, h, is_header=True)
        
        # 设置数据
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                self._set_table_cell(table, i+1, j, str(cell_text))
        
        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(width)
        
        # 格式化表格
        self._format_table(table)
        
        return table
    
    # ==================== 列表方法 ====================
    
    def add_bullet_list(self, items, indent=0):
        """添加无序列表"""
        for item in items:
            p = self.doc.add_paragraph(style='List Bullet')
            run = p.add_run(item)
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def add_numbered_list(self, items, start=1):
        """添加有序列表"""
        for i, item in enumerate(items, start):
            p = self.doc.add_paragraph()
            run = p.add_run(f'{i}. {item}')
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ==================== 辅助方法 ====================
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_section_break(self):
        """添加分节符"""
        self.doc.add_section()
    
    def save(self, filename=None):
        """
        保存文档
        
        Args:
            filename: 文件名（不含路径和扩展名）
        
        Returns:
            str: 文件完整路径
        """
        output_dir = os.path.expanduser("~/.openclaw/workspace/output")
        os.makedirs(output_dir, exist_ok=True)
        
        if not filename:
            filename = f"{self.company}_{self.doc_type}_{self.created_date}"
        
        filepath = os.path.join(output_dir, f"{filename}.docx")
        self.doc.save(filepath)
        
        print(f"✅ 文档已生成: {filepath}")
        return filepath


# ==================== 工具函数 ====================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def merge_cells(table, start_row, start_col, end_row, end_col):
    """合并单元格"""
    cell = table.rows[start_row].cells[start_col]
    cell.merge(table.rows[end_row].cells[end_col])
    return cell


# ==================== 测试代码 ====================

if __name__ == '__main__':
    # 测试基础类
    doc = KingdeeWordBase("测试公司", "业务蓝图")
    
    # 封面
    doc.add_cover(
        title="业务蓝图设计报告",
        subtitle="测试公司新ERP管理系统项目",
        version="V1.0",
        author="金蝶软件（中国）有限公司"
    )
    
    # 文档控制
    doc.add_document_control([
        {'version': 'V1.0', 'date': '2024-01-01', 'author': '张三', 'changes': '初稿'},
        {'version': 'V1.1', 'date': '2024-01-15', 'author': '李四', 'changes': '修订财务章节'},
    ])
    
    # 目录
    doc.add_toc()
    
    # 设置页眉页脚
    doc.set_header_footer()
    
    # 正文
    doc.add_heading('第一章 项目概述', level=1)
    doc.add_paragraph('本文档是测试公司的业务蓝图设计报告，旨在描述系统的业务架构、数据架构、应用架构和技术架构。', indent=2)
    
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph('随着企业业务的快速发展，现有的信息系统已无法满足业务需求。', indent=2)
    
    # 添加表格
    doc.add_heading('1.2 项目目标', level=2)
    headers = ['序号', '目标', '说明']
    data = [
        ['1', '提高效率', '通过系统自动化提高业务处理效率'],
        ['2', '降低成本', '减少人工操作，降低运营成本'],
        ['3', '数据整合', '实现各业务系统数据整合'],
    ]
    doc.add_styled_table(headers, data, col_widths=[2, 4, 8])
    
    # 保存

    # 保存
    filepath = doc.save()
    print(f"测试文档已保存: {filepath}")

